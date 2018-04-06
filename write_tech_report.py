from docx import Document
from docx.shared import Inches, Pt

from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd
import numpy as np

#scipy.stats gives us the cumuluative density lookup function for our z-score to percentiles
import scipy.stats as st

#gives us survival analysis
#from lifelines import KaplanMeierFitter

from scipy.stats import chi2_contingency, fisher_exact, t as t_test
from scipy import stats
from scipy.stats.contingency import expected_freq
from random import randint
from numpy import mean, std, unique, array

import plotly
import plotly.tools as tls
import plotly.plotly as py
from plotly.graph_objs import Figure, Data, Layout
#import matplotlib.pyplot as plt
from plotly.graph_objs import *
from pylab import rcParams

import datetime
import matplotlib.pyplot as plt

base_image_path = "C:\\Users\\pgilmore\\Desktop\\Development\\"

error_pic_path = base_image_path+'Koala.jpg'


#bio matchers
#Employee Candidate Data_LOL_PAID	Employee Candidate Data_LOL_Ethnicity	Employee Candidate Data_LOL_Birth Year	Employee Candidate Data_LOL_Gender	Manager Candidate Data_LOL_PAID	Manager Candidate Data_LOL_Ethnicity	Manager Candidate Data_LOL_Birth Year	Manager Candidate Data_LOL_Gender	Employee Quality Control Data_LOL_PPI_TOTAL_ITEM_MINUTES


dupe_filter_headnames = ['Removal - unreliable_duplicate_tags_mdemas = 1',\
'mhogheads_MWM_Removal - mh lowes dreamteam dupes = 1',\
'request_data_GOBBY_Remove_these_dupes']


#the following two columns are known to be of datetime in the form M/D/YYYY; e.g.,  8/29/2014
#mhogheads_MWM_Original Hire Date (Employee)	mhogheads_MWM_Original Term Date
criterion_og_datetime_headnames = ['mhogheads_MWM_Original Hire Date (Employee)',\
'mhogheads_MWM_Original Term Date']

#these represent individuals and knows nothing of team identity
pa_id_headname = ['Employee Candidate Data_LOL_PAID']

#these represent teams only and knows nothing of individual identity
team_id_headname = ['Team ID']

#these can be used to refer to individuals within their particular team
#composed of pa_id and team_id with the special bridge _dreamteam_
special_teamtag_headname = ['master_mergekey']



def write_report_study1(filepath_for_workingdata, filepath_for_output):
        
	filepath_for_output = filepath_for_output
	filepath_for_workingdata = filepath_for_workingdata
	
	try:
		agg_aata = pd.read_excel(filepath_for_workingdata)
	except:
		agg_aata = pd.DataFrame()
	
	
	print('First lets remove the source LOWES dataset until later')
	
	aata = agg_aata[agg_aata['request_data_GOBBY_Source']=='PC']
	
	document = Document()
	
	#TITLE PAGE
	#Upload the header icon, this one is an infor icon
	document.add_picture(base_image_path+"titlepage_infor_icon.png", width=Inches(1.5), height=Inches(2.25))
		
	#This should be a header, how to do a header
	#write_title(), takes a custom string and returns the proper header
	lil_title = write_title('Test TD Facet Document')
	document.add_heading(lil_title, level=1)
	h1_style = document.styles['Heading 1']
	h1_style.font.name = 'Times New Roman'
	h1_style.font.size = Pt(24)
		
	#authors
	lil_author_list = ['Phillip L. Gilmore']
	
	for i in range(0,len(lil_author_list)):
		paragraph = document.add_paragraph(lil_author_list[i])
	
	#todays date
	today = datetime.date.today()
	todays_date_nicestring = today.strftime('%b %d, %Y')
	paragraph = document.add_paragraph(todays_date_nicestring)	
	
	document.add_page_break()	
	#PAGE
	
	#TABLE OF CONTENTS
	#SKIP THIS FOR NOW
	document.add_heading('Table of Contents', level=2)
	
	document.add_page_break()	
	#PAGE
	
	
	#INTRODUCTION - Purpose
	document.add_heading('Purpose', level=2)

	#INTRODUCTION - Theory
	document.add_heading('Theory', level=2)
	
	#INTRODUCTION - Formulations of the Algorithm
	document.add_heading('Formulations of the Algorithm', level=2)
	
	document.add_page_break()	
	#PAGE
	
	
	#METHOD - Method: Sample, Norms, and Cut Scores
	document.add_heading('Method: Sample, Norms, and Cut Scores', level=2)
	document.add_heading('Sample Description and Size', level=3)
	
	#get these values from your final dataset, here is where the fun begins
	
	final_n = int()
	final_team_n = int()
	
	starting_n, starting_team_n = describe_final_sample(aata)
	
	
	starting_n_text = "The starting dataset\
 for which we could link facet and team identity data contained\
 {} job incumbents across {} unique teams.".format(starting_n, starting_team_n)
	
	
	aata = build_dupe_filter(aata)
	
	#'Removal - within team dupes = 1' is now available
	bata = aata[aata['Removal - within team dupes = 1'] == 0]
	#bata has filtered out any teams in which non unique team members were discovered
	
	
	if len(aata)==len(bata):
		removal_dupe_text = ''
	else:
		removal_dupe_text = "Within team duplicate records\
 and their associated team members (n = {}).".format(len(aata)-len(bata))
	
	final_n, final_team_n = describe_final_sample(bata)
	
	
	bata = build_partial_predictor_filter(bata)
	#'Removal - Team members missing facet data = 1' is now available
	cata = bata[bata['Removal - Team members missing facet data = 1'] == 0]
	final_n, final_team_n = describe_final_sample(cata)
	
	
	if len(bata) == len(cata):
		removal_missing_pred_text = ''
	else:
		removal_missing_pred_text = "Individuals on teams with \
 missing predictor data (n = {})".format(len(bata)-len(cata))
		
	final_n_text = "The final sample\
 included {} employees across {} teams.".format(final_n,final_team_n)
	methods_sample_size_p1 = document.add_paragraph(starting_n_text)
	document.add_heading('Removals', level=4)
	document.add_paragraph(removal_dupe_text)
	document.add_paragraph(removal_missing_pred_text)
	document.add_paragraph(final_n_text)
	
	
	document.add_heading('Team Sizes in the Norming Sample (n = {}, teams = {})'.format(final_n,final_team_n), level=4)
	
	team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std = \
	quick_teamsize_stats(cata)
	
	team_size_text = 'Team sizes varied from {} to {} \
(team size: median = {:.2f}; mean = {:.2f}; std = {:.2f}).'.format(\
	team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std)
	
	team_size_distr_frame = build_team_size_frame(cata)
	#just parsed these into lists for easier looping
	size_list = [int(i) for i in team_size_distr_frame['Team Size'].tolist()]
	freq_list = [int(i) for i in team_size_distr_frame['Frequency'].tolist()]
	
	document.add_paragraph(team_size_text)
	
	team_size_distribution_text = 'Teams in general tended to be smaller. The table below displays frequencies of team sizes.\
 The team sizes appear to shrink at a geometric rate with the vast majority of teams containing single digit numbers of team members.'
	document.add_paragraph(team_size_distribution_text)
	
	table_teamsize = document.add_table(len(size_list)+1,2)    
	heading_cells = table_teamsize.rows[0].cells
	heading_cells[0].text = 'Team Size (team members)'
	heading_cells[1].text = 'Frequency (number teams)'
	
	
	for i in range(0,len(size_list)):
		this_row = table_teamsize.rows[i+1].cells
		this_row[0].text = str(size_list[i])
		this_row[1].text = str(freq_list[i])
		
	
	document.add_page_break()
	
	#These _sqrt columns are the "Transformed" columns that we need to compute norming parameters for
	cata = build_sqrt_cols(cata)
	
	
	'''TABLE OF INDIVIDUAL FACET POP PARAMETERS'''
	#First we just display the individual facet means and sds, bc this is a new assessment and these
	#params are not well known
	document.add_heading('Individual Facet Information in the Norming Sample (n = {} Incumbents)'.format(final_n), level=4)

	
	facet_info_frame = build_facet_info_frame(cata)
	facet_list = [str(i) for i in facet_info_frame['Facets'].tolist()]
	mean_list = ['{:.2f}'.format(float(i)) for i in facet_info_frame['mean'].tolist()]
	std_list = ['{:.2f}'.format(float(i)) for i in facet_info_frame['std'].tolist()]
	
	table_facet_info = document.add_table(len(facet_info_frame)+1,3)
	heading_cells = table_facet_info.rows[0].cells
	heading_cells[0].text = 'Facet'
	heading_cells[1].text = 'Mean'
	heading_cells[2].text = 'Std Dev'
	
	for i in range(0,len(facet_list)):
		this_row = table_facet_info.rows[i+1].cells
		this_row[0].text = facet_list[i]
		this_row[1].text = mean_list[i]
		this_row[2].text = std_list[i]
		
	#print(facet_info_frame)
	document.add_page_break()	
	
	
	'''TABLE OF AVG TEAM DISTANCE TRANSFORMED POP PARAMETERS
	FOR EACH SCOPE OF THE METRIC: GLOBAL OR FACET;
	THESE BECOME THE Z-NORMS FOR THIS METRIC'''
	#Distance to Team (Transformed) Parameters Table, global and each facet
	document.add_heading('Population Parameter Estimates for Distance to Team (Transformed) (n = {} Incumbents)'.format(final_n), level=4)
	
	#The predictor label column is "Metric" here
	team_dist_info_frame = build_team_dist_info_frame(cata)	
	lil_team_dist_metric_namelist = [str(i) for i in team_dist_info_frame['Metric'].tolist()]
	mean_list = ['{:.4f}'.format(float(i)) for i in team_dist_info_frame['mean'].tolist()]
	std_list = ['{:.4f}'.format(float(i)) for i in team_dist_info_frame['std'].tolist()]
	
	table_team_dist_info = document.add_table(len(team_dist_info_frame)+1,3)
	heading_cells = table_team_dist_info.rows[0].cells
	heading_cells[0].text = 'Metric'
	heading_cells[1].text = 'Mean'
	heading_cells[2].text = 'Std Dev'
	
	for i in range(0,len(lil_team_dist_metric_namelist)):
		this_row = table_team_dist_info.rows[i+1].cells
		this_row[0].text = lil_team_dist_metric_namelist[i]
		this_row[1].text = mean_list[i]
		this_row[2].text = std_list[i]
	
	#print(team_dist_info_frame)	
	document.add_page_break()
	
	#Distance to Manager (Transformed), global and each facet
	mgr_dyad_n = int()
	mgr_dist_info_frame, mgr_dyad_n = build_mgr_dist_info_frame(cata)
	document.add_heading('Population Parameter Estimates for Distance to Manager (Transformed) (n = {} Incumbent-Manager Dyads)'.format(mgr_dyad_n), level=4)	
	
	lil_mgr_dist_metric_namelist = [str(i) for i in mgr_dist_info_frame['Metric'].tolist()]
	mean_list = ['{:.4f}'.format(float(i)) for i in mgr_dist_info_frame['mean'].tolist()]
	std_list = ['{:.4f}'.format(float(i)) for i in mgr_dist_info_frame['std'].tolist()]
	
	table_mgr_dist_info = document.add_table(len(mgr_dist_info_frame)+1,3)
	heading_cells = table_mgr_dist_info.rows[0].cells
	heading_cells[0].text = 'Metric'
	heading_cells[1].text = 'Mean'
	heading_cells[2].text = 'Std Dev'
	
	for i in range(0,len(lil_mgr_dist_metric_namelist)):
		this_row = table_mgr_dist_info.rows[i+1].cells
		this_row[0].text = lil_mgr_dist_metric_namelist[i]
		this_row[1].text = mean_list[i]
		this_row[2].text = std_list[i]
	
	document.add_page_break()
	
	#Overall Team Distance, global and each facet
	ot_teams_n = int()
	ot_dist_info_frame, ot_teams_n = build_ot_dist_info_frame(cata)
	document.add_heading('Population Parameter Estimates for Overall Team Distance (Transformed) (n = {} Teams)'.format(ot_teams_n), level=4)
	
	lil_ot_dist_metric_namelist = [str(i) for i in ot_dist_info_frame['Metric'].tolist()]
	mean_list = ['{:.4f}'.format(float(i)) for i in ot_dist_info_frame['mean'].tolist()]
	std_list = ['{:.4f}'.format(float(i)) for i in ot_dist_info_frame['std'].tolist()]
	
	table_ot_dist_info = document.add_table(len(ot_dist_info_frame)+1,3)
	heading_cells = table_ot_dist_info.rows[0].cells
	heading_cells[0].text = 'Metric'
	heading_cells[1].text = 'Mean'
	heading_cells[2].text = 'Std Dev'
	
	for i in range(0,len(lil_ot_dist_metric_namelist)):
		this_row = table_ot_dist_info.rows[i+1].cells
		this_row[0].text = lil_ot_dist_metric_namelist[i]
		this_row[1].text = mean_list[i]
		this_row[2].text = std_list[i]
	
	document.add_page_break()
	
	document.add_heading('Norms: Distance to Team, Distance to Manager, and Overall Team Distance', level=3)	
	
	#for each histo, we can want mean(), .std()	
	histo_metric_namelist = build_histo_metric_namelist()
	for i in histo_metric_namelist:
		document.add_paragraph('Figure Histogram:')
		try:
			temp_fp = hist_show(cata,i)
			document.add_picture(temp_fp, width=Inches(3.0), height=Inches(2.0))
		except:
			document.add_picture(error_pic_path, width=Inches(0.5), height=Inches(0.5))
			
		#write the descriptive tables here
	
	cata.to_excel("C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\"+"study1_working_datafile.xlsx")
	
	document.add_page_break()	
	#PAGE
	
	'''	
	document.add_picture("hist_show.png", width=Inches(5.0), height=Inches(4.0))
	'''	
		
	document.save(filepath_for_output)
		
	return



def write_report_study2(filepath_for_workingdata, filepath_for_output):
        
	filepath_for_output = filepath_for_output
	filepath_for_workingdata = filepath_for_workingdata
	
	try:
		agg_aata = pd.read_excel(filepath_for_workingdata)
	except:
		agg_aata = pd.DataFrame()
	
	
	print('First lets select the source LOWES dataset')
	
	aata = agg_aata[agg_aata['request_data_GOBBY_Source']=='LOWES']
	
	document = Document()
	
	#TITLE PAGE
	#Upload the header icon, this one is an infor icon
	document.add_picture(base_image_path+"titlepage_infor_icon.png", width=Inches(1.5), height=Inches(2.25))
		
	#This should be a header, how to do a header
	#write_title(), takes a custom string and returns the proper header
	lil_title = write_title('Test TD - Teams in the Lawn and Garden Equipment and Supplies Stores')
	document.add_heading(lil_title, level=1)
	h1_style = document.styles['Heading 1']
	h1_style.font.name = 'Times New Roman'
	h1_style.font.size = Pt(24)
		
	#authors
	lil_author_list = ['Phillip L. Gilmore']
	
	for i in range(0,len(lil_author_list)):
		paragraph = document.add_paragraph(lil_author_list[i])
	
	#todays date
	today = datetime.date.today()
	todays_date_nicestring = today.strftime('%b %d, %Y')
	paragraph = document.add_paragraph(todays_date_nicestring)	
	
	document.add_page_break()	
	#PAGE
	
	#TABLE OF CONTENTS
	#SKIP THIS FOR NOW
	document.add_heading('Table of Contents', level=2)
	
	document.add_page_break()	
	#PAGE
	
	
	#INTRODUCTION - Purpose
	document.add_heading('Purpose', level=2)

	#INTRODUCTION - Theory
	document.add_heading('Theory', level=2)
	
	#INTRODUCTION - Formulations of the Algorithm
	document.add_heading('Formulations of the Algorithm', level=2)
	
	document.add_page_break()	
	#PAGE
	
	
	#METHOD - Method: Sample, Norms, and Cut Scores
	document.add_heading('Method: Sample, Norms, and Cut Scores', level=2)
	document.add_heading('Sample Description and Size', level=3)
	
	#get these values from your final dataset, here is where the fun begins
	
	final_n = int()
	final_team_n = int()
	
	starting_n, starting_team_n = describe_final_sample(aata)
	
	
	starting_n_text = "The starting dataset\
 for which we could link facet and team identity data contained\
 {} job incumbents across {} unique teams.".format(starting_n, starting_team_n)
	
	
	aata = build_dupe_filter(aata)
	
	#'Removal - within team dupes = 1' is now available
	bata = aata[aata['Removal - within team dupes = 1'] == 0]
	#bata has filtered out any teams in which non unique team members were discovered
	
	
	if len(aata)==len(bata):
		removal_dupe_text = ''
	else:
		removal_dupe_text = "Within team duplicate records\
 and their associated team members (n = {}).".format(len(aata)-len(bata))
	
	final_n, final_team_n = describe_final_sample(bata)
	
	
	bata = build_partial_predictor_filter(bata)
	#'Removal - Team members missing facet data = 1' is now available
	cata = bata[bata['Removal - Team members missing facet data = 1'] == 0]
	
	
	if len(bata) == len(cata):
		removal_missing_pred_text = ''
	else:
		removal_missing_pred_text = "Individuals on teams with \
 missing predictor data (n = {})".format(len(bata)-len(cata))
	
	
	data = cata[cata['mhogheads_MWM_Removal (Low Confidence - Team Match)'] == 0]
	final_n, final_team_n = describe_final_sample(data)
	
	if len(cata) == len(data):
		removal_low_conf_text = ''
	else:
		removal_low_conf_text = "Individuals with low confidence \
 match to the team (n = {})".format(len(cata)-len(data))
	
	
	final_n_text = "The final sample\
 included {} employees across {} teams.".format(final_n,final_team_n)
	methods_sample_size_p1 = document.add_paragraph(starting_n_text)
	document.add_heading('Removals', level=4)
	document.add_paragraph(removal_dupe_text)
	document.add_paragraph(removal_missing_pred_text)
	document.add_paragraph(removal_low_conf_text)
	document.add_paragraph(final_n_text)
	
	
	document.add_heading('Team Sizes in the Client Sample (n = {}, teams = {})'.format(final_n,final_team_n), level=4)
	
	team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std = \
	quick_teamsize_stats(cata)
	
	team_size_text = 'Team sizes varied from {} to {} \
(team size: median = {:.2f}; mean = {:.2f}; std = {:.2f}).'.format(\
	team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std)
	
	team_size_distr_frame = build_team_size_frame(cata)
	#just parsed these into lists for easier looping
	size_list = [int(i) for i in team_size_distr_frame['Team Size'].tolist()]
	freq_list = [int(i) for i in team_size_distr_frame['Frequency'].tolist()]
	
	document.add_paragraph(team_size_text)
	
	
	table_teamsize = document.add_table(len(size_list)+1,2)    
	heading_cells = table_teamsize.rows[0].cells
	heading_cells[0].text = 'Team Size (team members)'
	heading_cells[1].text = 'Frequency (number teams)'
	
	
	for i in range(0,len(size_list)):
		this_row = table_teamsize.rows[i+1].cells
		this_row[0].text = str(size_list[i])
		this_row[1].text = str(freq_list[i])
		
	
	document.add_page_break()
	
	#These _sqrt columns are the "Transformed" columns that we need to compute norming parameters for
	data = build_client_sqrt_cols(data)
	data = build_facet_metric_cols(data)
	
	#in this section we depend on the termination columns: 'mhogheads_MWM_Status'
	#and our cohesion labels which we computed from known metric column headers
	#comparison
	data['Termed=1'] = [0 if i =='Active' else 1 for i in data['mhogheads_MWM_Status']]	
	
	data.to_excel("C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\"+"study2_working_datafile.xlsx")
		
	
	'''TABLE OF INDIVIDUAL FACET POP PARAMETERS'''
	#First we just display the individual facet means and sds, bc this is a new assessment and these
	#params are not well known
	document.add_heading('Individual Facet Information in the Client Sample (n = {} Incumbents)'.format(final_n), level=4)
	
	facet_info_frame = build_facet_info_frame(data)
	facet_list = [str(i) for i in facet_info_frame['Facets'].tolist()]
	mean_list = ['{:.2f}'.format(float(i)) for i in facet_info_frame['mean'].tolist()]
	std_list = ['{:.2f}'.format(float(i)) for i in facet_info_frame['std'].tolist()]
	
	table_facet_info = document.add_table(len(facet_info_frame)+1,3)
	heading_cells = table_facet_info.rows[0].cells
	heading_cells[0].text = 'Facet'
	heading_cells[1].text = 'Mean'
	heading_cells[2].text = 'Std Dev'
	
	for i in range(0,len(facet_list)):
		this_row = table_facet_info.rows[i+1].cells
		this_row[0].text = facet_list[i]
		this_row[1].text = mean_list[i]
		this_row[2].text = std_list[i]
		
	#print(facet_info_frame)
	
	document.add_page_break()
	
	document.add_heading('Z-Scored Avg. Distance to Team in the Client Sample', level=3)	
	
	#for each histo, we can want mean(), .std()
	document.add_paragraph('Figure Histogram:')
	
	team_dist_study2_list = [\
	'Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_sqrt_z_score',\
	'Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_Cohesion_Percentile',\
	'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_sqrt_z_score',\
	'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_Cohesion_Percentile',\
	'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_sqrt_z_score',\
	'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Percentile']
	
	for i in team_dist_study2_list:
		try:
			temp_fp = hist_show(data,i)
			document.add_picture(temp_fp, width=Inches(3.0), height=Inches(2.0))
		except:
			document.add_picture(error_pic_path, width=Inches(0.5), height=Inches(0.5))
		
		#write the descriptive tables here
		
	document.add_page_break()	
	#PAGE
	
	document.add_heading('Termination Rate by Avg. Score to Team Cohesion Category', level=3)	
	p1_blurb = 'Average Score to Team Cohesion categories were tested for their relative termination rates in the final sample (N = {} incumbents)'.format(len(data))
	p2_blurb = 'A chi-squared based hypothesis test was conducted on the termination frequencies. We expected that the impact of Average Score to Team Cohesion on termination would be negative; such that, team members classified as Very Cohesive (higher on cohesion) should display the lowest termination rates (lower termination frequency)'
	document.add_paragraph(p1_blurb)
	document.add_paragraph(p2_blurb)
	
	table_group_types = ['All Cohesion Levels', 'Very Diverse compared to all else', 'Very Cohesive compared to all else']
	metric_types = ['Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_Cohesion_Label']
	
	#actual column headers of the returned table_freqs_frame
	#N	Termed_Freq.	Termed_Pct.	Retained_Freq.	Retained_Pct.
	#'Chi_sq' 'p_value' 'dof'
	
	#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR ALL COHESION GROUPS
	lil_table_head = 'Table X. Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('All Cohesion Levels',len(data))
	
	document.add_heading(lil_table_head, level=3)
	term_freq_table_frame = build_term_table_frame(data,metric_types[0],'Termed=1','Very Diverse','Very Cohesive')
	
	#this part generalizes fairly well once you have defined your termination and cohesion columns
	cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
	cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
	term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
	term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
	ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
	ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
	chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
	p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
	dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
				
	table_term = document.add_table(len(term_freq_table_frame)+1,8)
	heading_cells = table_term.rows[0].cells
	heading_cells[0].text = 'Cohesion Group'
	heading_cells[1].text = 'n'
	heading_cells[2].text = 'Termed Freq.'
	heading_cells[3].text = 'Termed Pct.'
	heading_cells[4].text = 'Retained Freq.'
	heading_cells[5].text = 'Retained Pct.'
	heading_cells[6].text = 'Chi sq.'
	heading_cells[7].text = 'p'
	
	for i in range(0,len(cohesion_group_list)):
		this_row = table_term.rows[i+1].cells
		this_row[0].text = cohesion_group_list[i]
		this_row[1].text = cell_n_list[i]
		this_row[2].text = term_freq_list[i]
		this_row[3].text = term_pct_list[i]
		this_row[4].text = ret_freq_list[i]
		this_row[5].text = ret_pct_list[i]
		
		if i == 0:				
			this_row[6].text = chi_sq_list[i]
			this_row[7].text = p_value_list[i]
		else:
			this_row[6].text = ''
			this_row[7].text = ''
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
	document.add_paragraph(dof_blurb)
		
	
	#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR Very Diverse to others
	lil_table_head = 'Table X. Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Very Diverse vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	quick_trip_data = pd.DataFrame()
	quick_trip_data = data[[metric_types[0],'Termed=1']]
	quick_trip_data['Cohesion_VeryDiverse_vs_others'] = ['Very Diverse' if i == 'Very Diverse' else 'Diverse+'\
	for i in quick_trip_data[metric_types[0]].tolist()]
	
	term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryDiverse_vs_others','Termed=1','Very Diverse','Diverse+')
	
	#this part generalizes fairly well once you have defined your termination and cohesion columns
	cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
	cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
	term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
	term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
	ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
	ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
	chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
	p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
	dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
				
	table_term = document.add_table(len(term_freq_table_frame)+1,8)
	heading_cells = table_term.rows[0].cells
	heading_cells[0].text = 'Cohesion Group'
	heading_cells[1].text = 'n'
	heading_cells[2].text = 'Termed Freq.'
	heading_cells[3].text = 'Termed Pct.'
	heading_cells[4].text = 'Retained Freq.'
	heading_cells[5].text = 'Retained Pct.'
	heading_cells[6].text = 'Chi sq.'
	heading_cells[7].text = 'p'
	
	for i in range(0,len(cohesion_group_list)):
		this_row = table_term.rows[i+1].cells
		this_row[0].text = cohesion_group_list[i]
		this_row[1].text = cell_n_list[i]
		this_row[2].text = term_freq_list[i]
		this_row[3].text = term_pct_list[i]
		this_row[4].text = ret_freq_list[i]
		this_row[5].text = ret_pct_list[i]
		
		if i == 0:				
			this_row[6].text = chi_sq_list[i]
			this_row[7].text = p_value_list[i]
		else:
			this_row[6].text = ''
			this_row[7].text = ''
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
	document.add_paragraph(dof_blurb)
	
	
	
	#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR Very Cohesive to others
	
	lil_table_head = 'Table X. Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Very Cohesive vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	quick_trip_data = pd.DataFrame()
	quick_trip_data = data[[metric_types[0],'Termed=1']]
	quick_trip_data['Cohesion_VeryCohesive_vs_others'] = ['Very Cohesive' if i == 'Very Cohesive' else 'Cohesive+'\
	for i in quick_trip_data[metric_types[0]].tolist()]
	
	term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryCohesive_vs_others','Termed=1','Cohesive+','Very Cohesive')
	
	
	#this part generalizes fairly well once you have defined your termination and cohesion columns
	cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
	cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
	term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
	term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
	ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
	ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
	chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
	p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
	dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
				
	table_term = document.add_table(len(term_freq_table_frame)+1,8)
	heading_cells = table_term.rows[0].cells
	heading_cells[0].text = 'Cohesion Group'
	heading_cells[1].text = 'n'
	heading_cells[2].text = 'Termed Freq.'
	heading_cells[3].text = 'Termed Pct.'
	heading_cells[4].text = 'Retained Freq.'
	heading_cells[5].text = 'Retained Pct.'
	heading_cells[6].text = 'Chi sq.'
	heading_cells[7].text = 'p'
	
	for i in range(0,len(cohesion_group_list)):
		this_row = table_term.rows[i+1].cells
		this_row[0].text = cohesion_group_list[i]
		this_row[1].text = cell_n_list[i]
		this_row[2].text = term_freq_list[i]
		this_row[3].text = term_pct_list[i]
		this_row[4].text = ret_freq_list[i]
		this_row[5].text = ret_pct_list[i]
		
		if i == 0:				
			this_row[6].text = chi_sq_list[i]
			this_row[7].text = p_value_list[i]
		else:
			this_row[6].text = ''
			this_row[7].text = ''
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
	
	document.add_paragraph(dof_blurb)		
	
	
	facet_loop_blurb = 'Additionally, individual facets were aggregated to form new Average Team Scores. The Average Team Score for each facet was examined using the same method just described.'
	document.add_paragraph(facet_loop_blurb)
	
	
	#FACET LOOP
	#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM loop for facets
	#comparing Cohesive+ vs others
	
	lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Cohesive+ vs. Diverse+',len(data))
	document.add_heading(lil_table_head, level=3)
	
	facet_tps_cohesion_metric_list = [i+'_LOL_Avg. Score to Team (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_chisq_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_term_pct_list = list()
	lil_ret_pct_list = list()
	lil_support_list = list()
	quick_trip_data = pd.DataFrame()
	
	for facet_tps in facet_tps_cohesion_metric_list:		
		try:
			quick_trip_data = pd.DataFrame()
			quick_trip_data = data[[facet_tps,'Termed=1']]
			
			quick_trip_data['Cohesion_plus_vs_others'] = ['Cohesive+' if 'Cohesive' in i else 'Diverse+'\
			for i in quick_trip_data[facet_tps].tolist()]
		
			term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_plus_vs_others','Termed=1','Diverse+','Cohesive+')
		
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
			lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
			if isinstance(term_freq_table_frame['Pct Diff'][0],float):
				if term_freq_table_frame['Pct Diff'][0] > 0:
					lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				elif term_freq_table_frame['Pct Diff'][0] <= 0:
					lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('n/a')
			lil_p_list.append('n/a')
			lil_ret_pct_list = list('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
	heading_cells = table_facet_tps_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support, Pct. Diff'
	heading_cells[2].text = 'Chi-sq'
	heading_cells[3].text = 'p-value'
	
	for i in range(0,len(facet_tps_cohesion_metric_list)):
		this_row = table_facet_tps_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_chisq_list[i]
		this_row[3].text = lil_p_list[i]
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
	support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
	"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
	document.add_paragraph(dof_blurb)
	document.add_paragraph(support_blurb)
	
	
	#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM loop for facets
	#comparing Very Diverse vs others
	
	lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Very Diverse vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	facet_tps_cohesion_metric_list = [i+'_LOL_Avg. Score to Team (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_chisq_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_term_pct_list = list()
	lil_ret_pct_list = list()
	lil_support_list = list()
	quick_trip_data = pd.DataFrame()
	
	for facet_tps in facet_tps_cohesion_metric_list:		
		try:
			quick_trip_data = pd.DataFrame()
			quick_trip_data = data[[facet_tps,'Termed=1']]
			
			quick_trip_data['VeryDiverse_vs_others'] = ['Very Diverse' if 'Very Diverse' in i else 'others'\
			for i in quick_trip_data[facet_tps].tolist()]
		
			term_freq_table_frame = build_term_table_frame(quick_trip_data,'VeryDiverse_vs_others','Termed=1','Very Diverse','others')
		
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
			lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
			if isinstance(term_freq_table_frame['Pct Diff'][0],float):
				if term_freq_table_frame['Pct Diff'][0] > 0:
					lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				elif term_freq_table_frame['Pct Diff'][0] <= 0:
					lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('n/a')
			lil_p_list.append('n/a')
			lil_ret_pct_list = list('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
	heading_cells = table_facet_tps_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support, Pct. Diff'
	heading_cells[2].text = 'Chi-sq'
	heading_cells[3].text = 'p-value'
	
	for i in range(0,len(facet_tps_cohesion_metric_list)):
		this_row = table_facet_tps_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_chisq_list[i]
		this_row[3].text = lil_p_list[i]
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
	support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
	"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
	document.add_paragraph(dof_blurb)
	document.add_paragraph(support_blurb)
	
	document.add_page_break()
	
	
	document.add_heading('Termination Rate by Manager Fit Cohesion Category', level=3)	
	p1_blurb = 'Manager Fit Cohesion categories were tested for their relative termination rates in the final sample (N = {} manager-incumbent dyads)'.format(len(data))
	p2_blurb = 'A chi-squared based hypothesis test was conducted on the termination frequencies. We expected that the impact of Manager Fit Cohesion on termination would be negative; such that, team members classified as Very Cohesive (higher on cohesion) should display the lowest termination rates (lower termination frequency)'
	document.add_paragraph(p1_blurb)
	document.add_paragraph(p2_blurb)
	
	table_group_types = ['All Cohesion Levels', 'Very Diverse compared to all else', 'Very Cohesive compared to all else']
	metric_types = ['Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_Cohesion_Label']
	
	#actual column headers of the returned table_freqs_frame
	#N	Termed_Freq.	Termed_Pct.	Retained_Freq.	Retained_Pct.
	#'Chi_sq' 'p_value' 'dof'
	
	#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR ALL COHESION GROUPS
	lil_table_head = 'Table X. Termination Frequencies by Manager Fit: {} (N = {})'.format('All Cohesion Levels',len(data))
	
	document.add_heading(lil_table_head, level=3)
	term_freq_table_frame = build_term_table_frame(data,metric_types[0],'Termed=1','Very Diverse','Very Cohesive')
	
	#this part generalizes fairly well once you have defined your termination and cohesion columns
	cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
	cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
	term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
	term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
	ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
	ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
	chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
	p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
	dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
				
	table_term = document.add_table(len(term_freq_table_frame)+1,8)
	heading_cells = table_term.rows[0].cells
	heading_cells[0].text = 'Cohesion Group'
	heading_cells[1].text = 'n'
	heading_cells[2].text = 'Termed Freq.'
	heading_cells[3].text = 'Termed Pct.'
	heading_cells[4].text = 'Retained Freq.'
	heading_cells[5].text = 'Retained Pct.'
	heading_cells[6].text = 'Chi sq.'
	heading_cells[7].text = 'p'
	
	for i in range(0,len(cohesion_group_list)):
		this_row = table_term.rows[i+1].cells
		this_row[0].text = cohesion_group_list[i]
		this_row[1].text = cell_n_list[i]
		this_row[2].text = term_freq_list[i]
		this_row[3].text = term_pct_list[i]
		this_row[4].text = ret_freq_list[i]
		this_row[5].text = ret_pct_list[i]
		
		if i == 0:				
			this_row[6].text = chi_sq_list[i]
			this_row[7].text = p_value_list[i]
		else:
			this_row[6].text = ''
			this_row[7].text = ''
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
	document.add_paragraph(dof_blurb)
		
	
	#TABLE TERMINATION FREQUENCIES MANAGER FIT FOR Very Diverse to others
	lil_table_head = 'Table X. Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Diverse vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	quick_trip_data = pd.DataFrame()
	quick_trip_data = data[[metric_types[0],'Termed=1']]
	quick_trip_data['Cohesion_VeryDiverse_vs_others'] = ['Very Diverse' if i == 'Very Diverse' else 'Diverse+'\
	for i in quick_trip_data[metric_types[0]].tolist()]
	
	term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryDiverse_vs_others','Termed=1','Very Diverse','Diverse+')
	
	#this part generalizes fairly well once you have defined your termination and cohesion columns
	cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
	cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
	term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
	term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
	ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
	ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
	chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
	p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
	dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
				
	table_term = document.add_table(len(term_freq_table_frame)+1,8)
	heading_cells = table_term.rows[0].cells
	heading_cells[0].text = 'Cohesion Group'
	heading_cells[1].text = 'n'
	heading_cells[2].text = 'Termed Freq.'
	heading_cells[3].text = 'Termed Pct.'
	heading_cells[4].text = 'Retained Freq.'
	heading_cells[5].text = 'Retained Pct.'
	heading_cells[6].text = 'Chi sq.'
	heading_cells[7].text = 'p'
	
	for i in range(0,len(cohesion_group_list)):
		this_row = table_term.rows[i+1].cells
		this_row[0].text = cohesion_group_list[i]
		this_row[1].text = cell_n_list[i]
		this_row[2].text = term_freq_list[i]
		this_row[3].text = term_pct_list[i]
		this_row[4].text = ret_freq_list[i]
		this_row[5].text = ret_pct_list[i]
		
		if i == 0:				
			this_row[6].text = chi_sq_list[i]
			this_row[7].text = p_value_list[i]
		else:
			this_row[6].text = ''
			this_row[7].text = ''
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
	document.add_paragraph(dof_blurb)
	
	
	
	#TABLE TERMINATION FREQUENCIES MANAGER FIT FOR Very Cohesive to others
	
	lil_table_head = 'Table X. Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Cohesive vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	quick_trip_data = pd.DataFrame()
	quick_trip_data = data[[metric_types[0],'Termed=1']]
	quick_trip_data['Cohesion_VeryCohesive_vs_others'] = ['Very Cohesive' if i == 'Very Cohesive' else 'Cohesive+'\
	for i in quick_trip_data[metric_types[0]].tolist()]
	
	term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryCohesive_vs_others','Termed=1','Cohesive+','Very Cohesive')
	
	
	#this part generalizes fairly well once you have defined your termination and cohesion columns
	cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
	cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
	term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
	term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
	ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
	ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
	chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
	p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
	dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
				
	table_term = document.add_table(len(term_freq_table_frame)+1,8)
	heading_cells = table_term.rows[0].cells
	heading_cells[0].text = 'Cohesion Group'
	heading_cells[1].text = 'n'
	heading_cells[2].text = 'Termed Freq.'
	heading_cells[3].text = 'Termed Pct.'
	heading_cells[4].text = 'Retained Freq.'
	heading_cells[5].text = 'Retained Pct.'
	heading_cells[6].text = 'Chi sq.'
	heading_cells[7].text = 'p'
	
	for i in range(0,len(cohesion_group_list)):
		this_row = table_term.rows[i+1].cells
		this_row[0].text = cohesion_group_list[i]
		this_row[1].text = cell_n_list[i]
		this_row[2].text = term_freq_list[i]
		this_row[3].text = term_pct_list[i]
		this_row[4].text = ret_freq_list[i]
		this_row[5].text = ret_pct_list[i]
		
		if i == 0:				
			this_row[6].text = chi_sq_list[i]
			this_row[7].text = p_value_list[i]
		else:
			this_row[6].text = ''
			this_row[7].text = ''
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
	
	document.add_paragraph(dof_blurb)		
	
	
	facet_loop_blurb = 'Additionally, individual facets were aggregated to form new Manager Fit Scores. The Manager Fit Score for each facet was examined using the same method just described.'
	document.add_paragraph(facet_loop_blurb)
	
	#FACET LOOP
	#TABLE TERMINATION FREQUENCIES MANAGER FIT loop for facets
	#comparing Very Cohesive vs others
	
	lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Cohesive vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	facet_tps_cohesion_metric_list = [i+'_LOL_Manager Fit Score (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_chisq_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_term_pct_list = list()
	lil_ret_pct_list = list()
	lil_support_list = list()
	quick_trip_data = pd.DataFrame()
	
	for facet_tps in facet_tps_cohesion_metric_list:		
		try:
			quick_trip_data = pd.DataFrame()
			quick_trip_data = data[[facet_tps,'Termed=1']]
			
			quick_trip_data['VeryCohesive_vs_others'] = ['Very Cohesive' if 'Very Cohesive' in i else 'others'\
			for i in quick_trip_data[facet_tps].tolist()]
		
			term_freq_table_frame = build_term_table_frame(quick_trip_data,'VeryDiverse_vs_others','Termed=1','others','Very Cohesive')
		
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
			lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
			if isinstance(term_freq_table_frame['Pct Diff'][0],float):
				if term_freq_table_frame['Pct Diff'][0] > 0:
					lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				elif term_freq_table_frame['Pct Diff'][0] <= 0:
					lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('n/a')
			lil_p_list.append('n/a')
			lil_ret_pct_list = list('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
	heading_cells = table_facet_tps_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support, Pct. Diff'
	heading_cells[2].text = 'Chi-sq'
	heading_cells[3].text = 'p-value'
	
	for i in range(0,len(facet_tps_cohesion_metric_list)):
		this_row = table_facet_tps_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_chisq_list[i]
		this_row[3].text = lil_p_list[i]
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
	support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
	"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
	document.add_paragraph(dof_blurb)
	document.add_paragraph(support_blurb)
	
	
	#TABLE TERMINATION FREQUENCIES MANAGER FIT loop for facets
	#comparing Cohesive+ vs others
	
	lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Manager Fit: {} (N = {})'.format('Cohesive+ vs. Diverse+',len(data))
	document.add_heading(lil_table_head, level=3)
	
	facet_tps_cohesion_metric_list = [i+'_LOL_Manager Fit Score (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_chisq_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_term_pct_list = list()
	lil_ret_pct_list = list()
	lil_support_list = list()
	quick_trip_data = pd.DataFrame()
	
	for facet_tps in facet_tps_cohesion_metric_list:		
		try:
			quick_trip_data = pd.DataFrame()
			quick_trip_data = data[[facet_tps,'Termed=1']]
			
			quick_trip_data['Cohesion_plus_vs_others'] = ['Cohesive+' if 'Cohesive' in i else 'Diverse+'\
			for i in quick_trip_data[facet_tps].tolist()]
		
			term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_plus_vs_others','Termed=1','Diverse+','Cohesive+')
		
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
			lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
			if isinstance(term_freq_table_frame['Pct Diff'][0],float):
				if term_freq_table_frame['Pct Diff'][0] > 0:
					lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				elif term_freq_table_frame['Pct Diff'][0] <= 0:
					lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('n/a')
			lil_p_list.append('n/a')
			lil_ret_pct_list = list('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
	heading_cells = table_facet_tps_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support, Pct. Diff'
	heading_cells[2].text = 'Chi-sq'
	heading_cells[3].text = 'p-value'
	
	for i in range(0,len(facet_tps_cohesion_metric_list)):
		this_row = table_facet_tps_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_chisq_list[i]
		this_row[3].text = lil_p_list[i]
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
	support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
	"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
	document.add_paragraph(dof_blurb)
	document.add_paragraph(support_blurb)
	
	
	#TABLE TERMINATION FREQUENCIES MGR FIT loop for facets
	#comparing Very Diverse vs others
	
	lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Diverse vs. Others',len(data))
	document.add_heading(lil_table_head, level=3)
	
	facet_tps_cohesion_metric_list = [i+'_LOL_Manager Fit Score (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_chisq_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_term_pct_list = list()
	lil_ret_pct_list = list()
	lil_support_list = list()
	quick_trip_data = pd.DataFrame()
	
	for facet_tps in facet_tps_cohesion_metric_list:		
		try:
			quick_trip_data = pd.DataFrame()
			quick_trip_data = data[[facet_tps,'Termed=1']]
			
			quick_trip_data['VeryDiverse_vs_others'] = ['Very Diverse' if 'Very Diverse' in i else 'others'\
			for i in quick_trip_data[facet_tps].tolist()]
		
			term_freq_table_frame = build_term_table_frame(quick_trip_data,'VeryDiverse_vs_others','Termed=1','Very Diverse','others')
		
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
			lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
			if isinstance(term_freq_table_frame['Pct Diff'][0],float):
				if term_freq_table_frame['Pct Diff'][0] > 0:
					lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				elif term_freq_table_frame['Pct Diff'][0] <= 0:
					lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_tps.split('_')[0])
			lil_chisq_list.append('n/a')
			lil_p_list.append('n/a')
			lil_ret_pct_list = list('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
	heading_cells = table_facet_tps_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support, Pct. Diff'
	heading_cells[2].text = 'Chi-sq'
	heading_cells[3].text = 'p-value'
	
	for i in range(0,len(facet_tps_cohesion_metric_list)):
		this_row = table_facet_tps_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_chisq_list[i]
		this_row[3].text = lil_p_list[i]
				
	dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
	support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
	document.add_paragraph(dof_blurb)
	document.add_paragraph(support_blurb)
	
	document.add_page_break()	

	
	#STARTS THE TEAM LEVEL OTC TABLES
	#There were no Very Cohesive teams in the overall otc
	
	
	#TABLE TEAM TERMS OTC FOR Very Diverse to others
	metric_types = ['Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label']
	
	quick_team_data = pd.DataFrame()
	quick_team_data = data
	quick_team_data['Cohesion_VeryDiverse_vs_others'] = ['Very Diverse' if i == 'Very Diverse' else 'Diverse+'\
	for i in quick_team_data[metric_types[0]].tolist()]	
	
	table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'Cohesion_VeryDiverse_vs_others','Termed=1','Very Diverse','Diverse+')
	
	if len(table_otc_t_test_frame) < 2:
		pass
	else:	
	
		lil_table_head = 'Table X. Termination Rate by Overall Team Chemistry: {} (N = {} teams)'.format('Very Diverse vs. Others',int(table_otc_t_test_frame['N'].sum()))
		document.add_heading(lil_table_head, level=3)
			
		#column headers of the otc t test frame
		#'N', 'Avg. Term Pct', 'SD. Term Pct', 'Cohens_D', 'p_value', 'dof'
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		
		cohesion_group_list = [str(i) for i in table_otc_t_test_frame.index.tolist()]
		cell_n_list = [str(i) for i in table_otc_t_test_frame['N'].tolist()]
		avg_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['Avg. Term Pct'].tolist()]
		sd_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['SD. Term Pct'].tolist()]
		cohens_d_list = ['{:.2f}'.format(i) for i in table_otc_t_test_frame['Cohens_D'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in table_otc_t_test_frame['p_value'].tolist()]
					
		team_table_term = document.add_table(len(table_otc_t_test_frame)+1,6)
		heading_cells = team_table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Avg. Pct. Terms within Team'
		heading_cells[3].text = 'SD. Pct. Terms within Team'
		heading_cells[4].text = 'Cohens D'
		heading_cells[5].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = team_table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = avg_team_term_pct_list[i]
			this_row[3].text = sd_team_term_pct_list[i]
			
			if i == 0:				
				this_row[4].text = cohens_d_list[i]
				this_row[5].text = p_value_list[i]
			else:
				this_row[4].text = ''
				this_row[5].text = ''
		
		#Here the t-test and associated degrees of freedom relate the most extreme groups available for analysis (not always full sample); it is these cell sizes summed minus 1
		dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
		int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
		
		cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
	
		document.add_paragraph(dof_blurb)
		document.add_paragraph(cohens_d_blurb)
		
		document.add_paragraph(dof_blurb)
	
		
	#TABLE TERMINATION FREQUENCIES OTC FOR Cohesive+ to others
	
	metric_types = ['Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label']
	
	quick_team_data = pd.DataFrame()
	quick_team_data = data
	quick_team_data['Cohesion_VeryCohesive_vs_others'] = ['Cohesive+' if i == 'Very Cohesive' or i == 'Cohesive' else 'Diverse+'\
	for i in quick_team_data[metric_types[0]].tolist()]	
	
	table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'Cohesion_VeryCohesive_vs_others','Termed=1','Diverse+','Cohesive+')
	
	if len(table_otc_t_test_frame) < 2:
		pass
	else:
		lil_table_head = 'Table X. Termination Frequencies by Overall Team Chemistry: {} (N = {})'.format('Cohesive+ vs. Diverse+',int(table_otc_t_test_frame['N'].sum()))
		document.add_heading(lil_table_head, level=3)
			
		#column headers of the otc t test frame
		#'N', 'Avg. Term Pct', 'SD. Term Pct', 'Cohens_D', 'p_value', 'dof'
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		
		cohesion_group_list = [str(i) for i in table_otc_t_test_frame.index.tolist()]
		cell_n_list = [str(i) for i in table_otc_t_test_frame['N'].tolist()]
		avg_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['Avg. Term Pct'].tolist()]
		sd_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['SD. Term Pct'].tolist()]
		cohens_d_list = ['{:.2f}'.format(i) for i in table_otc_t_test_frame['Cohens_D'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in table_otc_t_test_frame['p_value'].tolist()]
					
		team_table_term = document.add_table(len(table_otc_t_test_frame)+1,6)
		heading_cells = team_table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Avg. Pct. Terms within Team'
		heading_cells[3].text = 'SD. Pct. Terms within Team'
		heading_cells[4].text = 'Cohens D'
		heading_cells[5].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = team_table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = avg_team_term_pct_list[i]
			this_row[3].text = sd_team_term_pct_list[i]
			
			if i == 0:				
				this_row[4].text = cohens_d_list[i]
				this_row[5].text = p_value_list[i]
			else:
				this_row[4].text = ''
				this_row[5].text = ''
		
		#Here the t-test and associated degrees of freedom relate the most extreme groups available for analysis (not always full sample); it is these cell sizes summed minus 1
		dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
		int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
		
		cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
		
		document.add_paragraph(dof_blurb)
		document.add_paragraph(cohens_d_blurb)		
		document.add_page_break()		
	
		
	#FACET LOOP
	#TABLE TEAM TERMINATION RATES OTC loop for facets
	#There were no Very Cohesive teams in this dataset
		
	#TABLE TEAM TERMINATION RATES OTC loop for facets
	#comparing Cohesive+ vs others	
	
	lil_table_head = 'Table X. Summary of Facet-level Cohens D and T-tests for Termination Rate by Overall Team Chemistry: {} (N = {})'.format('Cohesive+ vs. Diverse+',23)
	
	document.add_heading(lil_table_head, level=3)
	
	facet_otc_cohesion_metric_list = [i+'_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_cohens_d_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_support_list = list()
	quick_team_data = pd.DataFrame()
	
	#PLG YOU ARE HERE, WHAT IS WRONG WITH THIS FACET LOOP; THESE WILL BE AMONGST YOUR BEST RESULTS.
	
	for facet_otc in facet_otc_cohesion_metric_list:		
		try:			
			
			quick_team_data = data
			quick_team_data['Cohesion_plus_vs_others'] = ['Cohesive+' if str(i) == 'Very Cohesive' or str(i) == 'Cohesive' else 'Diverse+'\
			for i in quick_team_data[facet_otc].tolist()]
			
			table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'Cohesion_plus_vs_others','Termed=1','Diverse+','Cohesive+')
			
			lil_facet_name_list.append(facet_otc.split('_')[0])
			lil_cohens_d_list.append('{:.2f}'.format(table_otc_t_test_frame['Cohens_D'][0]))
			lil_p_list.append('{:.4f}'.format(table_otc_t_test_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(table_otc_t_test_frame['dof'][0]))
			if isinstance(table_otc_t_test_frame['Cohens_D'][0],float):
				if table_otc_t_test_frame['Cohens_D'][0] > 0:
					lil_support_list.append('Yes')
				elif table_otc_t_test_frame['Cohens_D'][0] <= 0:
					lil_support_list.append('No')
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_otc.split('_')[0])
			lil_cohens_d_list.append('n/a')
			lil_p_list.append('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_otc_term = document.add_table(len(facet_otc_cohesion_metric_list)+1,4)
	heading_cells = table_facet_otc_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support'
	heading_cells[2].text = 'Cohens D'
	heading_cells[3].text = 'T-test p-value'
	
	for i in range(0,len(facet_otc_cohesion_metric_list)):
		this_row = table_facet_otc_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_cohens_d_list[i]
		this_row[3].text = lil_p_list[i] 
				
	dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
	int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
	
	support_blurb = '"Support" refers to the directionality of the within team termination rates. The group representing the most diverse group\
would have a positive relative termination rate as compared to the most cohesive group.'
	cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
	t_blurb = 'The p-value refers to a two-tailed independent t-test with equal variance assumptions, and the test p-value was divided by two incorporate the expectation our unidirectional hypothesis test.'
	
	document.add_paragraph(support_blurb)	
	document.add_paragraph(dof_blurb)
	document.add_paragraph(t_blurb)	
	
	#TABLE TERMINATION Rates OTC loop for facets
	#comparing Very Diverse vs others
	
	lil_table_head = 'Table X. Summary of Facet-level Cohens D and T-tests for Termination Rate by Overall Team Chemistry: {} (N = {})'.format('Very Diverse vs. Others',23)
	
	document.add_heading(lil_table_head, level=3)
	
	facet_otc_cohesion_metric_list = [i+'_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label' for i in facet_list]
	
	lil_facet_name_list = list()
	lil_cohens_d_list = list()
	lil_p_list = list()
	lil_dof_list = list()
	lil_support_list = list()
	quick_team_data = pd.DataFrame()
	
	for facet_otc in facet_otc_cohesion_metric_list:		
		try:			
			
			quick_team_data = data
			
			quick_team_data['VeryDiverse_vs_others'] = ['Very Diverse' if 'Very Diverse' == str(i) else 'others'\
			for i in quick_team_data[facet_otc].tolist()]
		
			table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'VeryDiverse_vs_others','Termed=1','Very Diverse','others')
			
			lil_facet_name_list.append(facet_otc.split('_')[0])
			lil_cohens_d_list.append('{:.2f}'.format(table_otc_t_test_frame['Cohens_D'][0]))
			lil_p_list.append('{:.4f}'.format(table_otc_t_test_frame['p_value'][0]))
			lil_dof_list.append('{}'.format(table_otc_t_test_frame['dof'][0]))
			if isinstance(table_otc_t_test_frame['Cohens_D'][0],float):
				if table_otc_t_test_frame['Cohens_D'][0] > 0:
					lil_support_list.append('Yes')
				elif table_otc_t_test_frame['Cohens_D'][0] <= 0:
					lil_support_list.append('No')
				else:
					lil_support_list.append('Compute error')
			else:
				lil_support_list.append('Also compute error')
		
		except:
			lil_facet_name_list.append(facet_otc.split('_')[0])
			lil_cohens_d_list.append('n/a')
			lil_p_list.append('n/a')
			lil_dof_list.append('n/a')
			lil_support_list.append('n/a')
				
	table_facet_otc_term = document.add_table(len(facet_otc_cohesion_metric_list)+1,4)
	heading_cells = table_facet_otc_term.rows[0].cells
	heading_cells[0].text = 'Facet Source'
	heading_cells[1].text = 'Support, Pct. Diff'
	heading_cells[2].text = 'Cohens D'
	heading_cells[3].text = 'T-test p-value'
	
	for i in range(0,len(facet_otc_cohesion_metric_list)):
		this_row = table_facet_otc_term.rows[i+1].cells
		this_row[0].text = lil_facet_name_list[i]
		this_row[1].text = lil_support_list[i]
		this_row[2].text = lil_cohens_d_list[i]
		this_row[3].text = lil_p_list[i] 
				
	dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
	int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
	
	support_blurb = '"Support" refers to the directionality of the within team termination rates. The group representing the most diverse group\
would have a positive relative termination rate as compared to the most cohesive group.'
	cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
	t_blurb = 'The p-value refers to a two-tailed independent t-test with equal variance assumptions, and the test p-value was divided by two incorporate the expectation our unidirectional hypothesis test.'
	
	document.add_paragraph(support_blurb)	
	document.add_paragraph(dof_blurb)
	document.add_paragraph(t_blurb)	
	
	document.add_page_break()
	
	document.save(filepath_for_output)

		
	return
	


def write_report_study3(filepath_for_workingdata, filepath_for_output):
		
	employee_gender_head = "Employee Candidate Data_LOL_Gender"
	mgr_gender_head = "Manager Candidate Data_LOL_Gender"
	
	#observed values are F, M, U, and X
	#lets squash U and X bc they both represent unknown types and may ultimately be excluded
	        
	filepath_for_output = filepath_for_output
	filepath_for_workingdata = filepath_for_workingdata
	
	try:
		agg_aata = pd.read_excel(filepath_for_workingdata)
	except:
		agg_aata = pd.DataFrame()
	
	
	document = Document()
	
	#TITLE PAGE
	#Upload the header icon, this one is an infor icon
	document.add_picture(base_image_path+"titlepage_infor_icon.png", width=Inches(1.5), height=Inches(2.25))	
	
	print('First lets select the source LOWES dataset')
	
	mgr_type_list = ['F','M','Other']
	
	agg_aata[mgr_gender_head] = [str(i).strip() for i in agg_aata[mgr_gender_head].tolist()]
	
	lil_title = write_title('Test TD - Teams in the Lawn and Garden Equipment and Supplies Stores')
	
	for mgr_type in mgr_type_list:
		aata = agg_aata[agg_aata['request_data_GOBBY_Source']=='LOWES']		
		
		if mgr_type == 'F':
			aata = aata[aata[mgr_gender_head]=='F']
			document.add_page_break()
			lil_title = write_title('Female Managers Test TD - Teams in the Lawn and Garden Equipment and Supplies Stores')
		elif mgr_type == 'M':
			aata = aata[aata[mgr_gender_head]=='M']
			document.add_page_break()
			lil_title = write_title('Male Managers Test TD - Teams in the Lawn and Garden Equipment and Supplies Stores')
		else:
			aata = aata[(aata[mgr_gender_head]!='M') & (aata[mgr_gender_head]!='F')]
			document.add_page_break()
			lil_title = write_title('Unknown Gender Managers Test TD - Teams in the Lawn and Garden Equipment and Supplies Stores')
	
	#This should be a header, how to do a header
	#write_title(), takes a custom string and returns the proper header
	
		document.add_heading(lil_title, level=1)
		h1_style = document.styles['Heading 1']
		h1_style.font.name = 'Times New Roman'
		h1_style.font.size = Pt(24)
			
		#authors
		lil_author_list = ['Phillip L. Gilmore']
		
		for i in range(0,len(lil_author_list)):
			paragraph = document.add_paragraph(lil_author_list[i])
		
		#todays date
		today = datetime.date.today()
		todays_date_nicestring = today.strftime('%b %d, %Y')
		paragraph = document.add_paragraph(todays_date_nicestring)	
		
		document.add_page_break()	
		#PAGE
		
		#TABLE OF CONTENTS
		#SKIP THIS FOR NOW
		document.add_heading('Table of Contents', level=2)
		
		document.add_page_break()	
		#PAGE
		
		
		#INTRODUCTION - Purpose
		document.add_heading('Purpose', level=2)

		#INTRODUCTION - Theory
		document.add_heading('Theory', level=2)
		
		#INTRODUCTION - Formulations of the Algorithm
		document.add_heading('Formulations of the Algorithm', level=2)
		
		document.add_page_break()	
		#PAGE
		
		
		#METHOD - Method: Sample, Norms, and Cut Scores
		document.add_heading('Method: Sample, Norms, and Cut Scores', level=2)
		document.add_heading('Sample Description and Size', level=3)
		
		#get these values from your final dataset, here is where the fun begins
		
		final_n = int()
		final_team_n = int()
		
		starting_n, starting_team_n = describe_final_sample(aata)
		
		
		starting_n_text = "The starting dataset\
	 for which we could link facet and team identity data contained\
	 {} job incumbents across {} unique teams.".format(starting_n, starting_team_n)
		
		
		aata = build_dupe_filter(aata)
		
		#'Removal - within team dupes = 1' is now available
		bata = aata[aata['Removal - within team dupes = 1'] == 0]
		#bata has filtered out any teams in which non unique team members were discovered
		
		
		if len(aata)==len(bata):
			removal_dupe_text = ''
		else:
			removal_dupe_text = "Within team duplicate records\
	 and their associated team members (n = {}).".format(len(aata)-len(bata))
		
		final_n, final_team_n = describe_final_sample(bata)
		
		
		bata = build_partial_predictor_filter(bata)
		#'Removal - Team members missing facet data = 1' is now available
		cata = bata[bata['Removal - Team members missing facet data = 1'] == 0]
		
		
		if len(bata) == len(cata):
			removal_missing_pred_text = ''
		else:
			removal_missing_pred_text = "Individuals on teams with \
	 missing predictor data (n = {})".format(len(bata)-len(cata))
		
		
		data = cata[cata['mhogheads_MWM_Removal (Low Confidence - Team Match)'] == 0]
		final_n, final_team_n = describe_final_sample(data)
		
		if len(cata) == len(data):
			removal_low_conf_text = ''
		else:
			removal_low_conf_text = "Individuals with low confidence \
	 match to the team (n = {})".format(len(cata)-len(data))
		
		
		final_n_text = "The final sample\
	 included {} employees across {} teams.".format(final_n,final_team_n)
		methods_sample_size_p1 = document.add_paragraph(starting_n_text)
		document.add_heading('Removals', level=4)
		document.add_paragraph(removal_dupe_text)
		document.add_paragraph(removal_missing_pred_text)
		document.add_paragraph(removal_low_conf_text)
		document.add_paragraph(final_n_text)
		
		
		document.add_heading('Team Sizes in the Client Sample (n = {}, teams = {})'.format(final_n,final_team_n), level=4)
		
		team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std = \
		quick_teamsize_stats(cata)
		
		team_size_text = 'Team sizes varied from {} to {} \
	(team size: median = {:.2f}; mean = {:.2f}; std = {:.2f}).'.format(\
		team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std)
		
		team_size_distr_frame = build_team_size_frame(cata)
		#just parsed these into lists for easier looping
		size_list = [int(i) for i in team_size_distr_frame['Team Size'].tolist()]
		freq_list = [int(i) for i in team_size_distr_frame['Frequency'].tolist()]
		
		document.add_paragraph(team_size_text)
		
		
		table_teamsize = document.add_table(len(size_list)+1,2)    
		heading_cells = table_teamsize.rows[0].cells
		heading_cells[0].text = 'Team Size (team members)'
		heading_cells[1].text = 'Frequency (number teams)'
		
		
		for i in range(0,len(size_list)):
			this_row = table_teamsize.rows[i+1].cells
			this_row[0].text = str(size_list[i])
			this_row[1].text = str(freq_list[i])
			
		
		document.add_page_break()
		
		#These _sqrt columns are the "Transformed" columns that we need to compute norming parameters for
		data = build_client_sqrt_cols(data)
		data = build_facet_metric_cols(data)
		
		#in this section we depend on the termination columns: 'mhogheads_MWM_Status'
		#and our cohesion labels which we computed from known metric column headers
		#comparison
		data['Termed=1'] = [0 if i =='Active' else 1 for i in data['mhogheads_MWM_Status']]	
		
		#data.to_excel("C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\"+"study2_working_datafile.xlsx")
			
		
		'''TABLE OF INDIVIDUAL FACET POP PARAMETERS'''
		#First we just display the individual facet means and sds, bc this is a new assessment and these
		#params are not well known
		document.add_heading('Individual Facet Information in the Client Sample (n = {} Incumbents)'.format(final_n), level=4)
		
		facet_info_frame = build_facet_info_frame(data)
		facet_list = [str(i) for i in facet_info_frame['Facets'].tolist()]
		mean_list = ['{:.2f}'.format(float(i)) for i in facet_info_frame['mean'].tolist()]
		std_list = ['{:.2f}'.format(float(i)) for i in facet_info_frame['std'].tolist()]
		
		table_facet_info = document.add_table(len(facet_info_frame)+1,3)
		heading_cells = table_facet_info.rows[0].cells
		heading_cells[0].text = 'Facet'
		heading_cells[1].text = 'Mean'
		heading_cells[2].text = 'Std Dev'
		
		for i in range(0,len(facet_list)):
			this_row = table_facet_info.rows[i+1].cells
			this_row[0].text = facet_list[i]
			this_row[1].text = mean_list[i]
			this_row[2].text = std_list[i]
			
		#print(facet_info_frame)
		
		document.add_page_break()
		
		document.add_heading('Z-Scored Avg. Distance to Team in the Client Sample', level=3)	
		
		#for each histo, we can want mean(), .std()
		document.add_paragraph('Figure Histogram:')
		
		team_dist_study2_list = [\
		'Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_sqrt_z_score',\
		'Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_Cohesion_Percentile',\
		'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_sqrt_z_score',\
		'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_Cohesion_Percentile',\
		'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_sqrt_z_score',\
		'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Percentile']
		
		for i in team_dist_study2_list:
			try:
				temp_fp = hist_show(data,i)
				document.add_picture(temp_fp, width=Inches(3.0), height=Inches(2.0))
			except:
				document.add_picture(error_pic_path, width=Inches(0.5), height=Inches(0.5))
			
			#write the descriptive tables here
			
		document.add_page_break()	
		#PAGE
		
		document.add_heading('Termination Rate by Avg. Score to Team Cohesion Category', level=3)	
		p1_blurb = 'Average Score to Team Cohesion categories were tested for their relative termination rates in the final sample (N = {} incumbents)'.format(len(data))
		p2_blurb = 'A chi-squared based hypothesis test was conducted on the termination frequencies. We expected that the impact of Average Score to Team Cohesion on termination would be negative; such that, team members classified as Very Cohesive (higher on cohesion) should display the lowest termination rates (lower termination frequency)'
		document.add_paragraph(p1_blurb)
		document.add_paragraph(p2_blurb)
		
		table_group_types = ['All Cohesion Levels', 'Very Diverse compared to all else', 'Very Cohesive compared to all else']
		metric_types = ['Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_Cohesion_Label']
		
		#actual column headers of the returned table_freqs_frame
		#N	Termed_Freq.	Termed_Pct.	Retained_Freq.	Retained_Pct.
		#'Chi_sq' 'p_value' 'dof'
		
		#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR ALL COHESION GROUPS
		lil_table_head = 'Table X. Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('All Cohesion Levels',len(data))
		
		document.add_heading(lil_table_head, level=3)
		term_freq_table_frame = build_term_table_frame(data,metric_types[0],'Termed=1','Very Diverse','Very Cohesive')
		
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
		cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
		term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
		term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
		ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
		ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
		chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
		dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
					
		table_term = document.add_table(len(term_freq_table_frame)+1,8)
		heading_cells = table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Termed Freq.'
		heading_cells[3].text = 'Termed Pct.'
		heading_cells[4].text = 'Retained Freq.'
		heading_cells[5].text = 'Retained Pct.'
		heading_cells[6].text = 'Chi sq.'
		heading_cells[7].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = term_freq_list[i]
			this_row[3].text = term_pct_list[i]
			this_row[4].text = ret_freq_list[i]
			this_row[5].text = ret_pct_list[i]
			
			if i == 0:				
				this_row[6].text = chi_sq_list[i]
				this_row[7].text = p_value_list[i]
			else:
				this_row[6].text = ''
				this_row[7].text = ''
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
		document.add_paragraph(dof_blurb)
			
		
		#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR Very Diverse to others
		lil_table_head = 'Table X. Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Very Diverse vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		quick_trip_data = pd.DataFrame()
		quick_trip_data = data[[metric_types[0],'Termed=1']]
		quick_trip_data['Cohesion_VeryDiverse_vs_others'] = ['Very Diverse' if i == 'Very Diverse' else 'Diverse+'\
		for i in quick_trip_data[metric_types[0]].tolist()]
		
		term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryDiverse_vs_others','Termed=1','Very Diverse','Diverse+')
		
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
		cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
		term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
		term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
		ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
		ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
		chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
		dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
					
		table_term = document.add_table(len(term_freq_table_frame)+1,8)
		heading_cells = table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Termed Freq.'
		heading_cells[3].text = 'Termed Pct.'
		heading_cells[4].text = 'Retained Freq.'
		heading_cells[5].text = 'Retained Pct.'
		heading_cells[6].text = 'Chi sq.'
		heading_cells[7].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = term_freq_list[i]
			this_row[3].text = term_pct_list[i]
			this_row[4].text = ret_freq_list[i]
			this_row[5].text = ret_pct_list[i]
			
			if i == 0:				
				this_row[6].text = chi_sq_list[i]
				this_row[7].text = p_value_list[i]
			else:
				this_row[6].text = ''
				this_row[7].text = ''
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
		document.add_paragraph(dof_blurb)
		
		
		
		#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR Very Cohesive to others
		
		lil_table_head = 'Table X. Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Very Cohesive vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		quick_trip_data = pd.DataFrame()
		quick_trip_data = data[[metric_types[0],'Termed=1']]
		quick_trip_data['Cohesion_VeryCohesive_vs_others'] = ['Very Cohesive' if i == 'Very Cohesive' else 'Cohesive+'\
		for i in quick_trip_data[metric_types[0]].tolist()]
		
		term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryCohesive_vs_others','Termed=1','Cohesive+','Very Cohesive')
		
		
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
		cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
		term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
		term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
		ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
		ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
		chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
		dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
					
		table_term = document.add_table(len(term_freq_table_frame)+1,8)
		heading_cells = table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Termed Freq.'
		heading_cells[3].text = 'Termed Pct.'
		heading_cells[4].text = 'Retained Freq.'
		heading_cells[5].text = 'Retained Pct.'
		heading_cells[6].text = 'Chi sq.'
		heading_cells[7].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = term_freq_list[i]
			this_row[3].text = term_pct_list[i]
			this_row[4].text = ret_freq_list[i]
			this_row[5].text = ret_pct_list[i]
			
			if i == 0:				
				this_row[6].text = chi_sq_list[i]
				this_row[7].text = p_value_list[i]
			else:
				this_row[6].text = ''
				this_row[7].text = ''
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
		
		document.add_paragraph(dof_blurb)		
		
		
		facet_loop_blurb = 'Additionally, individual facets were aggregated to form new Average Team Scores. The Average Team Score for each facet was examined using the same method just described.'
		document.add_paragraph(facet_loop_blurb)
		
		
		#FACET LOOP
		#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM loop for facets
		#comparing Cohesive+ vs others
		
		lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Cohesive+ vs. Diverse+',len(data))
		document.add_heading(lil_table_head, level=3)
		
		facet_tps_cohesion_metric_list = [i+'_LOL_Avg. Score to Team (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_chisq_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_term_pct_list = list()
		lil_ret_pct_list = list()
		lil_support_list = list()
		quick_trip_data = pd.DataFrame()
		
		for facet_tps in facet_tps_cohesion_metric_list:		
			try:
				quick_trip_data = pd.DataFrame()
				quick_trip_data = data[[facet_tps,'Termed=1']]
				
				quick_trip_data['Cohesion_plus_vs_others'] = ['Cohesive+' if 'Cohesive' in i else 'Diverse+'\
				for i in quick_trip_data[facet_tps].tolist()]
			
				term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_plus_vs_others','Termed=1','Diverse+','Cohesive+')
			
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
				lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
				if isinstance(term_freq_table_frame['Pct Diff'][0],float):
					if term_freq_table_frame['Pct Diff'][0] > 0:
						lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					elif term_freq_table_frame['Pct Diff'][0] <= 0:
						lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('n/a')
				lil_p_list.append('n/a')
				lil_ret_pct_list = list('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
		heading_cells = table_facet_tps_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support, Pct. Diff'
		heading_cells[2].text = 'Chi-sq'
		heading_cells[3].text = 'p-value'
		
		for i in range(0,len(facet_tps_cohesion_metric_list)):
			this_row = table_facet_tps_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_chisq_list[i]
			this_row[3].text = lil_p_list[i]
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
		support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
		would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
		"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
		document.add_paragraph(dof_blurb)
		document.add_paragraph(support_blurb)
		
		
		#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM loop for facets
		#comparing Very Diverse vs others
		
		lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Avg. Score to Team: {} (N = {})'.format('Very Diverse vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		facet_tps_cohesion_metric_list = [i+'_LOL_Avg. Score to Team (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_chisq_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_term_pct_list = list()
		lil_ret_pct_list = list()
		lil_support_list = list()
		quick_trip_data = pd.DataFrame()
		
		for facet_tps in facet_tps_cohesion_metric_list:		
			try:
				quick_trip_data = pd.DataFrame()
				quick_trip_data = data[[facet_tps,'Termed=1']]
				
				quick_trip_data['VeryDiverse_vs_others'] = ['Very Diverse' if 'Very Diverse' in i else 'others'\
				for i in quick_trip_data[facet_tps].tolist()]
			
				term_freq_table_frame = build_term_table_frame(quick_trip_data,'VeryDiverse_vs_others','Termed=1','Very Diverse','others')
			
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
				lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
				if isinstance(term_freq_table_frame['Pct Diff'][0],float):
					if term_freq_table_frame['Pct Diff'][0] > 0:
						lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					elif term_freq_table_frame['Pct Diff'][0] <= 0:
						lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('n/a')
				lil_p_list.append('n/a')
				lil_ret_pct_list = list('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
		heading_cells = table_facet_tps_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support, Pct. Diff'
		heading_cells[2].text = 'Chi-sq'
		heading_cells[3].text = 'p-value'
		
		for i in range(0,len(facet_tps_cohesion_metric_list)):
			this_row = table_facet_tps_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_chisq_list[i]
			this_row[3].text = lil_p_list[i]
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
		support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
		would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
		"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
		document.add_paragraph(dof_blurb)
		document.add_paragraph(support_blurb)
		
		document.add_page_break()
		
		
		document.add_heading('Termination Rate by Manager Fit Cohesion Category', level=3)	
		p1_blurb = 'Manager Fit Cohesion categories were tested for their relative termination rates in the final sample (N = {} manager-incumbent dyads)'.format(len(data))
		p2_blurb = 'A chi-squared based hypothesis test was conducted on the termination frequencies. We expected that the impact of Manager Fit Cohesion on termination would be negative; such that, team members classified as Very Cohesive (higher on cohesion) should display the lowest termination rates (lower termination frequency)'
		document.add_paragraph(p1_blurb)
		document.add_paragraph(p2_blurb)
		
		table_group_types = ['All Cohesion Levels', 'Very Diverse compared to all else', 'Very Cohesive compared to all else']
		metric_types = ['Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_Cohesion_Label']
		
		#actual column headers of the returned table_freqs_frame
		#N	Termed_Freq.	Termed_Pct.	Retained_Freq.	Retained_Pct.
		#'Chi_sq' 'p_value' 'dof'
		
		#TABLE TERMINATION FREQUENCIES AVG SCORE TO TEAM FOR ALL COHESION GROUPS
		lil_table_head = 'Table X. Termination Frequencies by Manager Fit: {} (N = {})'.format('All Cohesion Levels',len(data))
		
		document.add_heading(lil_table_head, level=3)
		term_freq_table_frame = build_term_table_frame(data,metric_types[0],'Termed=1','Very Diverse','Very Cohesive')
		
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
		cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
		term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
		term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
		ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
		ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
		chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
		dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
					
		table_term = document.add_table(len(term_freq_table_frame)+1,8)
		heading_cells = table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Termed Freq.'
		heading_cells[3].text = 'Termed Pct.'
		heading_cells[4].text = 'Retained Freq.'
		heading_cells[5].text = 'Retained Pct.'
		heading_cells[6].text = 'Chi sq.'
		heading_cells[7].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = term_freq_list[i]
			this_row[3].text = term_pct_list[i]
			this_row[4].text = ret_freq_list[i]
			this_row[5].text = ret_pct_list[i]
			
			if i == 0:				
				this_row[6].text = chi_sq_list[i]
				this_row[7].text = p_value_list[i]
			else:
				this_row[6].text = ''
				this_row[7].text = ''
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
		document.add_paragraph(dof_blurb)
			
		
		#TABLE TERMINATION FREQUENCIES MANAGER FIT FOR Very Diverse to others
		lil_table_head = 'Table X. Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Diverse vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		quick_trip_data = pd.DataFrame()
		quick_trip_data = data[[metric_types[0],'Termed=1']]
		quick_trip_data['Cohesion_VeryDiverse_vs_others'] = ['Very Diverse' if i == 'Very Diverse' else 'Diverse+'\
		for i in quick_trip_data[metric_types[0]].tolist()]
		
		term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryDiverse_vs_others','Termed=1','Very Diverse','Diverse+')
		
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
		cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
		term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
		term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
		ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
		ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
		chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
		dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
					
		table_term = document.add_table(len(term_freq_table_frame)+1,8)
		heading_cells = table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Termed Freq.'
		heading_cells[3].text = 'Termed Pct.'
		heading_cells[4].text = 'Retained Freq.'
		heading_cells[5].text = 'Retained Pct.'
		heading_cells[6].text = 'Chi sq.'
		heading_cells[7].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = term_freq_list[i]
			this_row[3].text = term_pct_list[i]
			this_row[4].text = ret_freq_list[i]
			this_row[5].text = ret_pct_list[i]
			
			if i == 0:				
				this_row[6].text = chi_sq_list[i]
				this_row[7].text = p_value_list[i]
			else:
				this_row[6].text = ''
				this_row[7].text = ''
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
		document.add_paragraph(dof_blurb)
		
		
		
		#TABLE TERMINATION FREQUENCIES MANAGER FIT FOR Very Cohesive to others
		
		lil_table_head = 'Table X. Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Cohesive vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		quick_trip_data = pd.DataFrame()
		quick_trip_data = data[[metric_types[0],'Termed=1']]
		quick_trip_data['Cohesion_VeryCohesive_vs_others'] = ['Very Cohesive' if i == 'Very Cohesive' else 'Cohesive+'\
		for i in quick_trip_data[metric_types[0]].tolist()]
		
		term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_VeryCohesive_vs_others','Termed=1','Cohesive+','Very Cohesive')
		
		
		#this part generalizes fairly well once you have defined your termination and cohesion columns
		cohesion_group_list = [str(i) for i in term_freq_table_frame.index.tolist()]
		cell_n_list = [str(i) for i in term_freq_table_frame['N'].tolist()]
		term_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Termed_Freq.'].tolist()]
		term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Termed_Pct.'].tolist()]
		ret_freq_list = ['{}'.format(int(i)) for i in term_freq_table_frame['Retained_Freq.'].tolist()]
		ret_pct_list = ['{:.2f}%'.format(float(i)*100) for i in term_freq_table_frame['Retained_Pct.'].tolist()]
		chi_sq_list = ['{:.2f}'.format(i) for i in term_freq_table_frame['Chi_sq'].tolist()]
		p_value_list = ['{:.4f}'.format(i) for i in term_freq_table_frame['p_value'].tolist()]
		dof_list = ['{}'.format(i) for i in term_freq_table_frame['dof'].tolist()]
					
		table_term = document.add_table(len(term_freq_table_frame)+1,8)
		heading_cells = table_term.rows[0].cells
		heading_cells[0].text = 'Cohesion Group'
		heading_cells[1].text = 'n'
		heading_cells[2].text = 'Termed Freq.'
		heading_cells[3].text = 'Termed Pct.'
		heading_cells[4].text = 'Retained Freq.'
		heading_cells[5].text = 'Retained Pct.'
		heading_cells[6].text = 'Chi sq.'
		heading_cells[7].text = 'p'
		
		for i in range(0,len(cohesion_group_list)):
			this_row = table_term.rows[i+1].cells
			this_row[0].text = cohesion_group_list[i]
			this_row[1].text = cell_n_list[i]
			this_row[2].text = term_freq_list[i]
			this_row[3].text = term_pct_list[i]
			this_row[4].text = ret_freq_list[i]
			this_row[5].text = ret_pct_list[i]
			
			if i == 0:				
				this_row[6].text = chi_sq_list[i]
				this_row[7].text = p_value_list[i]
			else:
				this_row[6].text = ''
				this_row[7].text = ''
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),dof_list[0])
		
		document.add_paragraph(dof_blurb)		
		
		
		facet_loop_blurb = 'Additionally, individual facets were aggregated to form new Manager Fit Scores. The Manager Fit Score for each facet was examined using the same method just described.'
		document.add_paragraph(facet_loop_blurb)
		
		#FACET LOOP
		#TABLE TERMINATION FREQUENCIES MANAGER FIT loop for facets
		#comparing Very Cohesive vs others
		
		lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Cohesive vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		facet_tps_cohesion_metric_list = [i+'_LOL_Manager Fit Score (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_chisq_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_term_pct_list = list()
		lil_ret_pct_list = list()
		lil_support_list = list()
		quick_trip_data = pd.DataFrame()
		
		for facet_tps in facet_tps_cohesion_metric_list:		
			try:
				quick_trip_data = pd.DataFrame()
				quick_trip_data = data[[facet_tps,'Termed=1']]
				
				quick_trip_data['VeryCohesive_vs_others'] = ['Very Cohesive' if 'Very Cohesive' in i else 'others'\
				for i in quick_trip_data[facet_tps].tolist()]
			
				term_freq_table_frame = build_term_table_frame(quick_trip_data,'VeryDiverse_vs_others','Termed=1','others','Very Cohesive')
			
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
				lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
				if isinstance(term_freq_table_frame['Pct Diff'][0],float):
					if term_freq_table_frame['Pct Diff'][0] > 0:
						lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					elif term_freq_table_frame['Pct Diff'][0] <= 0:
						lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('n/a')
				lil_p_list.append('n/a')
				lil_ret_pct_list = list('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
		heading_cells = table_facet_tps_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support, Pct. Diff'
		heading_cells[2].text = 'Chi-sq'
		heading_cells[3].text = 'p-value'
		
		for i in range(0,len(facet_tps_cohesion_metric_list)):
			this_row = table_facet_tps_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_chisq_list[i]
			this_row[3].text = lil_p_list[i]
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
		support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
		would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
		"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
		document.add_paragraph(dof_blurb)
		document.add_paragraph(support_blurb)
		
		
		#TABLE TERMINATION FREQUENCIES MANAGER FIT loop for facets
		#comparing Cohesive+ vs others
		
		lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Manager Fit: {} (N = {})'.format('Cohesive+ vs. Diverse+',len(data))
		document.add_heading(lil_table_head, level=3)
		
		facet_tps_cohesion_metric_list = [i+'_LOL_Manager Fit Score (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_chisq_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_term_pct_list = list()
		lil_ret_pct_list = list()
		lil_support_list = list()
		quick_trip_data = pd.DataFrame()
		
		for facet_tps in facet_tps_cohesion_metric_list:		
			try:
				quick_trip_data = pd.DataFrame()
				quick_trip_data = data[[facet_tps,'Termed=1']]
				
				quick_trip_data['Cohesion_plus_vs_others'] = ['Cohesive+' if 'Cohesive' in i else 'Diverse+'\
				for i in quick_trip_data[facet_tps].tolist()]
			
				term_freq_table_frame = build_term_table_frame(quick_trip_data,'Cohesion_plus_vs_others','Termed=1','Diverse+','Cohesive+')
			
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
				lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
				if isinstance(term_freq_table_frame['Pct Diff'][0],float):
					if term_freq_table_frame['Pct Diff'][0] > 0:
						lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					elif term_freq_table_frame['Pct Diff'][0] <= 0:
						lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('n/a')
				lil_p_list.append('n/a')
				lil_ret_pct_list = list('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
		heading_cells = table_facet_tps_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support, Pct. Diff'
		heading_cells[2].text = 'Chi-sq'
		heading_cells[3].text = 'p-value'
		
		for i in range(0,len(facet_tps_cohesion_metric_list)):
			this_row = table_facet_tps_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_chisq_list[i]
			this_row[3].text = lil_p_list[i]
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
		support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
		would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
		"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
		document.add_paragraph(dof_blurb)
		document.add_paragraph(support_blurb)
		
		
		#TABLE TERMINATION FREQUENCIES MGR FIT loop for facets
		#comparing Very Diverse vs others
		
		lil_table_head = 'Table X. Summary of Facet-level Chi-Square Tests for Termination Frequencies by Manager Fit: {} (N = {})'.format('Very Diverse vs. Others',len(data))
		document.add_heading(lil_table_head, level=3)
		
		facet_tps_cohesion_metric_list = [i+'_LOL_Manager Fit Score (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_chisq_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_term_pct_list = list()
		lil_ret_pct_list = list()
		lil_support_list = list()
		quick_trip_data = pd.DataFrame()
		
		for facet_tps in facet_tps_cohesion_metric_list:		
			try:
				quick_trip_data = pd.DataFrame()
				quick_trip_data = data[[facet_tps,'Termed=1']]
				
				quick_trip_data['VeryDiverse_vs_others'] = ['Very Diverse' if 'Very Diverse' in i else 'others'\
				for i in quick_trip_data[facet_tps].tolist()]
			
				term_freq_table_frame = build_term_table_frame(quick_trip_data,'VeryDiverse_vs_others','Termed=1','Very Diverse','others')
			
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('{:.2f}'.format(term_freq_table_frame['Chi_sq'][0]))
				lil_p_list.append('{:.4f}'.format(term_freq_table_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(term_freq_table_frame['dof'][0]))
				if isinstance(term_freq_table_frame['Pct Diff'][0],float):
					if term_freq_table_frame['Pct Diff'][0] > 0:
						lil_support_list.append('Yes, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					elif term_freq_table_frame['Pct Diff'][0] <= 0:
						lil_support_list.append('No, {:.2f}%'.format(term_freq_table_frame['Pct Diff'][0]*100))
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_tps.split('_')[0])
				lil_chisq_list.append('n/a')
				lil_p_list.append('n/a')
				lil_ret_pct_list = list('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_tps_term = document.add_table(len(facet_tps_cohesion_metric_list)+1,4)
		heading_cells = table_facet_tps_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support, Pct. Diff'
		heading_cells[2].text = 'Chi-sq'
		heading_cells[3].text = 'p-value'
		
		for i in range(0,len(facet_tps_cohesion_metric_list)):
			this_row = table_facet_tps_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_chisq_list[i]
			this_row[3].text = lil_p_list[i]
					
		dof_blurb = 'Pearson chi-square test from scipy.stats chi2_contingency method. N = {}, df = {}'.format(len(data),lil_dof_list[0])
		support_blurb = '"Support" refers to the directionality of the termination frequencies. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group (i.e., Pct Diff = (Term Pct Most Diverse - Term Pct Most Cohesive)/Term Pct Most Diverse).\
	"Yes" refers to a positive value which indicates directional support for our expectations. "No" is zero or negative Pct. Diff.'
		document.add_paragraph(dof_blurb)
		document.add_paragraph(support_blurb)
		
		document.add_page_break()	

		
		#STARTS THE TEAM LEVEL OTC TABLES
		#There were no Very Cohesive teams in the overall otc
		
		
		#TABLE TEAM TERMS OTC FOR Very Diverse to others
		metric_types = ['Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label']
		
		quick_team_data = pd.DataFrame()
		quick_team_data = data
		quick_team_data['Cohesion_VeryDiverse_vs_others'] = ['Very Diverse' if i == 'Very Diverse' else 'Diverse+'\
		for i in quick_team_data[metric_types[0]].tolist()]	
		
		table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'Cohesion_VeryDiverse_vs_others','Termed=1','Very Diverse','Diverse+')
		
		if len(table_otc_t_test_frame) < 2:
			pass
		else:	
		
			lil_table_head = 'Table X. Termination Rate by Overall Team Chemistry: {} (N = {} teams)'.format('Very Diverse vs. Others',int(table_otc_t_test_frame['N'].sum()))
			document.add_heading(lil_table_head, level=3)
				
			#column headers of the otc t test frame
			#'N', 'Avg. Term Pct', 'SD. Term Pct', 'Cohens_D', 'p_value', 'dof'
			#this part generalizes fairly well once you have defined your termination and cohesion columns
			
			cohesion_group_list = [str(i) for i in table_otc_t_test_frame.index.tolist()]
			cell_n_list = [str(i) for i in table_otc_t_test_frame['N'].tolist()]
			avg_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['Avg. Term Pct'].tolist()]
			sd_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['SD. Term Pct'].tolist()]
			cohens_d_list = ['{:.2f}'.format(i) for i in table_otc_t_test_frame['Cohens_D'].tolist()]
			p_value_list = ['{:.4f}'.format(i) for i in table_otc_t_test_frame['p_value'].tolist()]
						
			team_table_term = document.add_table(len(table_otc_t_test_frame)+1,6)
			heading_cells = team_table_term.rows[0].cells
			heading_cells[0].text = 'Cohesion Group'
			heading_cells[1].text = 'n'
			heading_cells[2].text = 'Avg. Pct. Terms within Team'
			heading_cells[3].text = 'SD. Pct. Terms within Team'
			heading_cells[4].text = 'Cohens D'
			heading_cells[5].text = 'p'
			
			for i in range(0,len(cohesion_group_list)):
				this_row = team_table_term.rows[i+1].cells
				this_row[0].text = cohesion_group_list[i]
				this_row[1].text = cell_n_list[i]
				this_row[2].text = avg_team_term_pct_list[i]
				this_row[3].text = sd_team_term_pct_list[i]
				
				if i == 0:				
					this_row[4].text = cohens_d_list[i]
					this_row[5].text = p_value_list[i]
				else:
					this_row[4].text = ''
					this_row[5].text = ''
			
			#Here the t-test and associated degrees of freedom relate the most extreme groups available for analysis (not always full sample); it is these cell sizes summed minus 1
			dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
			int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
			
			cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
		
			document.add_paragraph(dof_blurb)
			document.add_paragraph(cohens_d_blurb)
			
			document.add_paragraph(dof_blurb)
		
			
		#TABLE TERMINATION FREQUENCIES OTC FOR Cohesive+ to others
		
		metric_types = ['Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label']
		
		quick_team_data = pd.DataFrame()
		quick_team_data = data
		quick_team_data['Cohesion_VeryCohesive_vs_others'] = ['Cohesive+' if i == 'Very Cohesive' or i == 'Cohesive' else 'Diverse+'\
		for i in quick_team_data[metric_types[0]].tolist()]	
		
		table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'Cohesion_VeryCohesive_vs_others','Termed=1','Diverse+','Cohesive+')
		
		if len(table_otc_t_test_frame) < 2:
			pass
		else:
			lil_table_head = 'Table X. Termination Frequencies by Overall Team Chemistry: {} (N = {})'.format('Cohesive+ vs. Diverse+',int(table_otc_t_test_frame['N'].sum()))
			document.add_heading(lil_table_head, level=3)
				
			#column headers of the otc t test frame
			#'N', 'Avg. Term Pct', 'SD. Term Pct', 'Cohens_D', 'p_value', 'dof'
			#this part generalizes fairly well once you have defined your termination and cohesion columns
			
			cohesion_group_list = [str(i) for i in table_otc_t_test_frame.index.tolist()]
			cell_n_list = [str(i) for i in table_otc_t_test_frame['N'].tolist()]
			avg_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['Avg. Term Pct'].tolist()]
			sd_team_term_pct_list = ['{:.2f}%'.format(float(i)*100) for i in table_otc_t_test_frame['SD. Term Pct'].tolist()]
			cohens_d_list = ['{:.2f}'.format(i) for i in table_otc_t_test_frame['Cohens_D'].tolist()]
			p_value_list = ['{:.4f}'.format(i) for i in table_otc_t_test_frame['p_value'].tolist()]
						
			team_table_term = document.add_table(len(table_otc_t_test_frame)+1,6)
			heading_cells = team_table_term.rows[0].cells
			heading_cells[0].text = 'Cohesion Group'
			heading_cells[1].text = 'n'
			heading_cells[2].text = 'Avg. Pct. Terms within Team'
			heading_cells[3].text = 'SD. Pct. Terms within Team'
			heading_cells[4].text = 'Cohens D'
			heading_cells[5].text = 'p'
			
			for i in range(0,len(cohesion_group_list)):
				this_row = team_table_term.rows[i+1].cells
				this_row[0].text = cohesion_group_list[i]
				this_row[1].text = cell_n_list[i]
				this_row[2].text = avg_team_term_pct_list[i]
				this_row[3].text = sd_team_term_pct_list[i]
				
				if i == 0:				
					this_row[4].text = cohens_d_list[i]
					this_row[5].text = p_value_list[i]
				else:
					this_row[4].text = ''
					this_row[5].text = ''
			
			#Here the t-test and associated degrees of freedom relate the most extreme groups available for analysis (not always full sample); it is these cell sizes summed minus 1
			dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
			int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
			
			cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
			
			document.add_paragraph(dof_blurb)
			document.add_paragraph(cohens_d_blurb)		
			document.add_page_break()		
		
			
		#FACET LOOP
		#TABLE TEAM TERMINATION RATES OTC loop for facets
		#There were no Very Cohesive teams in this dataset
			
		#TABLE TEAM TERMINATION RATES OTC loop for facets
		#comparing Cohesive+ vs others	
		
		lil_table_head = 'Table X. Summary of Facet-level Cohens D and T-tests for Termination Rate by Overall Team Chemistry: {} (N = {})'.format('Cohesive+ vs. Diverse+',23)
		
		document.add_heading(lil_table_head, level=3)
		
		facet_otc_cohesion_metric_list = [i+'_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_cohens_d_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_support_list = list()
		quick_team_data = pd.DataFrame()
		
		#PLG YOU ARE HERE, WHAT IS WRONG WITH THIS FACET LOOP; THESE WILL BE AMONGST YOUR BEST RESULTS.
		
		for facet_otc in facet_otc_cohesion_metric_list:		
			try:			
				
				quick_team_data = data
				quick_team_data['Cohesion_plus_vs_others'] = ['Cohesive+' if str(i) == 'Very Cohesive' or str(i) == 'Cohesive' else 'Diverse+'\
				for i in quick_team_data[facet_otc].tolist()]
				
				table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'Cohesion_plus_vs_others','Termed=1','Diverse+','Cohesive+')
				
				lil_facet_name_list.append(facet_otc.split('_')[0])
				lil_cohens_d_list.append('{:.2f}'.format(table_otc_t_test_frame['Cohens_D'][0]))
				lil_p_list.append('{:.4f}'.format(table_otc_t_test_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(table_otc_t_test_frame['dof'][0]))
				if isinstance(table_otc_t_test_frame['Cohens_D'][0],float):
					if table_otc_t_test_frame['Cohens_D'][0] > 0:
						lil_support_list.append('Yes')
					elif table_otc_t_test_frame['Cohens_D'][0] <= 0:
						lil_support_list.append('No')
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_otc.split('_')[0])
				lil_cohens_d_list.append('n/a')
				lil_p_list.append('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_otc_term = document.add_table(len(facet_otc_cohesion_metric_list)+1,4)
		heading_cells = table_facet_otc_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support'
		heading_cells[2].text = 'Cohens D'
		heading_cells[3].text = 'T-test p-value'
		
		for i in range(0,len(facet_otc_cohesion_metric_list)):
			this_row = table_facet_otc_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_cohens_d_list[i]
			this_row[3].text = lil_p_list[i] 
					
		dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
		int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
		
		support_blurb = '"Support" refers to the directionality of the within team termination rates. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group.'
		cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
		t_blurb = 'The p-value refers to a two-tailed independent t-test with equal variance assumptions, and the test p-value was divided by two incorporate the expectation our unidirectional hypothesis test.'
		
		document.add_paragraph(support_blurb)	
		document.add_paragraph(dof_blurb)
		document.add_paragraph(t_blurb)	
		
		#TABLE TERMINATION Rates OTC loop for facets
		#comparing Very Diverse vs others
		
		lil_table_head = 'Table X. Summary of Facet-level Cohens D and T-tests for Termination Rate by Overall Team Chemistry: {} (N = {})'.format('Very Diverse vs. Others',23)
		
		document.add_heading(lil_table_head, level=3)
		
		facet_otc_cohesion_metric_list = [i+'_LOL_Overall Team Avg. Score (Standardized)_Cohesion_Label' for i in facet_list]
		
		lil_facet_name_list = list()
		lil_cohens_d_list = list()
		lil_p_list = list()
		lil_dof_list = list()
		lil_support_list = list()
		quick_team_data = pd.DataFrame()
		
		for facet_otc in facet_otc_cohesion_metric_list:		
			try:			
				
				quick_team_data = data
				
				quick_team_data['VeryDiverse_vs_others'] = ['Very Diverse' if 'Very Diverse' == str(i) else 'others'\
				for i in quick_team_data[facet_otc].tolist()]
			
				table_otc_t_test_frame = build_otc_t_table_frame(quick_team_data,'VeryDiverse_vs_others','Termed=1','Very Diverse','others')
				
				lil_facet_name_list.append(facet_otc.split('_')[0])
				lil_cohens_d_list.append('{:.2f}'.format(table_otc_t_test_frame['Cohens_D'][0]))
				lil_p_list.append('{:.4f}'.format(table_otc_t_test_frame['p_value'][0]))
				lil_dof_list.append('{}'.format(table_otc_t_test_frame['dof'][0]))
				if isinstance(table_otc_t_test_frame['Cohens_D'][0],float):
					if table_otc_t_test_frame['Cohens_D'][0] > 0:
						lil_support_list.append('Yes')
					elif table_otc_t_test_frame['Cohens_D'][0] <= 0:
						lil_support_list.append('No')
					else:
						lil_support_list.append('Compute error')
				else:
					lil_support_list.append('Also compute error')
			
			except:
				lil_facet_name_list.append(facet_otc.split('_')[0])
				lil_cohens_d_list.append('n/a')
				lil_p_list.append('n/a')
				lil_dof_list.append('n/a')
				lil_support_list.append('n/a')
					
		table_facet_otc_term = document.add_table(len(facet_otc_cohesion_metric_list)+1,4)
		heading_cells = table_facet_otc_term.rows[0].cells
		heading_cells[0].text = 'Facet Source'
		heading_cells[1].text = 'Support, Pct. Diff'
		heading_cells[2].text = 'Cohens D'
		heading_cells[3].text = 'T-test p-value'
		
		for i in range(0,len(facet_otc_cohesion_metric_list)):
			this_row = table_facet_otc_term.rows[i+1].cells
			this_row[0].text = lil_facet_name_list[i]
			this_row[1].text = lil_support_list[i]
			this_row[2].text = lil_cohens_d_list[i]
			this_row[3].text = lil_p_list[i] 
					
		dof_blurb = 'Independent t-test with equal variance assumptions test was used from scipy.stats.ttest_ind. N = {}, df = {},{}'.format(\
		int(table_otc_t_test_frame['N'].sum()),table_otc_t_test_frame['dof'].unique()[0],1)
		
		support_blurb = '"Support" refers to the directionality of the within team termination rates. The group representing the most diverse group\
	would have a positive relative termination rate as compared to the most cohesive group.'
		cohens_d_blurb = 'Cohens D represents a sample-weighted effect size. A positive Cohens D indicates directional support of our hypotheses.'
		t_blurb = 'The p-value refers to a two-tailed independent t-test with equal variance assumptions, and the test p-value was divided by two incorporate the expectation our unidirectional hypothesis test.'
		
		document.add_paragraph(support_blurb)	
		document.add_paragraph(dof_blurb)
		document.add_paragraph(t_blurb)	
		
		document.add_page_break()
		
		document.save(filepath_for_output)	

	return
	
	
def write_title(title_request):

	title_request = title_request
	title_base = 'Science/Technical Specification:'
	
	
	try:
		if len(title_request) == 0:
			title_string = 'Science/Technical Specification:'
		else:
			title_string = '{} {}'.format(title_base,title_request)
		
	except:
		title_string = '{} Team Dynamics Formulation Using the IPA Facets'.format(title_base)
	

	return title_string
		

def describe_final_sample(working_data_frame):
	
	working_data_frame = working_data_frame
	
	team_frame = working_data_frame['Team ID'].value_counts()
	
	final_n, final_team_n = int(), int()
	
	final_n = len(working_data_frame)
	final_team_n = len(team_frame)
	
	return final_n, final_team_n
	

def build_dupe_filter(working_data_frame):
	working_data_frame = working_data_frame
	
	#this will extract the team ids with non unique within team members
	dupe_special_tag = working_data_frame['master_mergekey'].value_counts()
	team_id_dupe_filter_list = dupe_special_tag[dupe_special_tag > 1].index.tolist()
	team_id_dupe_filter_list = [int(str(i.split('_')[-1])) for i in team_id_dupe_filter_list]
	
	#[ unicode(x.strip()) if x is not None else '' for x in row ]
	quick_list = [1 if int(i) in team_id_dupe_filter_list else 0 for i in working_data_frame['Team ID'].tolist()]
	
	working_data_frame['Removal - within team dupes = 1'] = quick_list

	return working_data_frame
	

def build_team_size_frame(working_data_frame):
	
	working_data_frame = working_data_frame
	
	try:
		team_size_distr_frame = working_data_frame['Team ID'].value_counts().value_counts()
		team_size_distr_frame = team_size_distr_frame.to_frame()
		team_size_distr_frame['Team Size'] = team_size_distr_frame.index
		team_size_distr_frame['Frequency'] = team_size_distr_frame['Team ID']
		del team_size_distr_frame['Team ID']
		team_size_distr_frame['inv_Frequency'] = team_size_distr_frame['Team Size'] * -1
		team_size_distr_frame = team_size_distr_frame.sort_values(by=['inv_Frequency','Team Size'],\
		ascending=False)
		del team_size_distr_frame['inv_Frequency']
		
	except:
		team_size_distr_frame = pd.DataFrame({'Team Size': [], 'Frequency': []})
	
	return team_size_distr_frame

	
def build_facet_info_frame(working_data_frame):
	
	
	'''Takes the known individual facet level columns for each candidate ID and 
	returns a frame of their observed statistical characteristics'''
	working_data_frame = working_data_frame
	facet_namelist = build_facet_namelist()
	
	facet_columns = ["Employee Facet Score_LOL_{}".format(i) for i in facet_namelist]
	
	facet_info_frame = working_data_frame[facet_columns].describe().T
	facet_info_frame['Facets'] = [str(i.split('_')[-1]) for i in facet_info_frame.index]
	
	return facet_info_frame

	
def build_team_dist_info_frame(working_data_frame):
	
	'''Takes the known average team distance transformed metric columns for each candidate ID and 
	returns a frame of their observed statistical characteristics'''
	working_data_frame = working_data_frame
	team_dist_metric_namelist = build_team_dist_metric_namelist()
	
	team_dist_info_frame = working_data_frame[team_dist_metric_namelist].describe().T
	team_dist_info_frame['Metric'] = [str(i.split('_')[0]) for i in team_dist_info_frame.index]

	return team_dist_info_frame
	

def build_mgr_dist_info_frame(working_data_frame):
	
	'''Takes the known mgr distance transformed metric columns for each candidate ID and 
	returns a frame of their observed statistical characteristics'''
	working_data_frame = working_data_frame
	mgr_dist_metric_namelist = build_mgr_dist_metric_namelist()
	
	dropped_wframe = working_data_frame[mgr_dist_metric_namelist].dropna()
	
	mgr_dist_info_frame = dropped_wframe.describe().T
	mgr_dist_info_frame['Metric'] = [str(i.split('_')[0]) for i in mgr_dist_info_frame.index]
	
	return mgr_dist_info_frame, len(dropped_wframe)

	
def  build_ot_dist_info_frame(working_data_frame):
	
	'''Takes the known overall team distance transformed metric columns for each team and 
	returns a frame of their observed statistical characteristics'''
	
	working_data_frame = working_data_frame
	ot_dist_metric_namelist = build_ot_dist_metric_namelist()
	lil_mean_list = list()
	lil_std_list = list()
	lil_compute_frame = pd.DataFrame()
	
	for i in range(0,len(ot_dist_metric_namelist)):
		
		lil_compute_frame = pd.pivot_table(working_data_frame,values=ot_dist_metric_namelist[i],index='Team ID',aggfunc=np.max)
		lil_ot_dist_frame = lil_compute_frame.describe().T
		lil_mean_list.append(float(lil_ot_dist_frame['mean']))
		lil_std_list.append(float(lil_ot_dist_frame['std']))
		ot_teams_n = len(lil_compute_frame)
	
	#build this frame, the 'Metric' column, the "mean" and "std"
	ot_dist_info_frame = pd.DataFrame({"Metric_full_label": ot_dist_metric_namelist, "mean": \
	lil_mean_list,"std":lil_std_list})
	
	ot_dist_info_frame['Metric'] = [str(i.split('_')[0]) for i in ot_dist_info_frame['Metric_full_label'].tolist()]
	
	return	ot_dist_info_frame, ot_teams_n

	
def build_facet_namelist():
	
	'''Employee Facet Score_LOL_ precedes all of the individual facet information
	such as, Employee Facet Score_LOL_Appeasement
	then the facet names are given in a list below
	'''

	facet_namelist = ['Appeasement','Approachability','Camaraderie','Control of Negativity','Egotism',\
	'Empathy','Familiarity','Focus','General Business Practices','Haste','Hustle','Imaginative',\
	'Intuition-Based Decision-Making','Need for Mental Challenge','Need for Simplicity','Need to Obtain Power',\
	'Perfectionism','Positivity','Preference for Group Work','Reward Focus','Righteousness','Solace','Vexation']

	return facet_namelist
	

def build_partial_predictor_filter(working_data_frame):
	
	working_data_frame = working_data_frame
	quick_list = list()
	
	facet_namelist = build_facet_namelist()
	facet_columns = ["Employee Facet Score_LOL_{}".format(i) for i in facet_namelist]
	
	working_data_frame['EmployeeFacetCount_23max'] = \
	working_data_frame[facet_columns].count(axis=1)
	
	working_data_frame['Removal - Partially missing facet data = 1'] = \
    pd.Series((working_data_frame['EmployeeFacetCount_23max'] < 23),dtype=np.int32)
	
	#data['Item SD.s'] = data.ix[:,start_item_column:(start_item_column+k)].std(axis=1,skipna=True)	
	#this will extract the team ids with non unique within team members
	
	team_id_missingfacets_filter_list = \
	working_data_frame[working_data_frame\
	['Removal - Partially missing facet data = 1'] == 1]['Team ID'].value_counts().index.tolist()
	#team_id_dupe_filter_list = [str(i.split('_')[-1]) for i in team_id_dupe_filter_list]
	
	quick_list = [1 if int(i) in team_id_missingfacets_filter_list\
	else 0 for i in working_data_frame['Team ID'].tolist()]
	
	working_data_frame['Removal - Team members missing facet data = 1'] = quick_list

	return working_data_frame  

	
def hist_show(working_data_frame,col_tag):
    
	working_data_frame = working_data_frame
	team_level_frame = pd.DataFrame()
	base_image_path = "C:\\Users\\pgilmore\\Desktop\\Development\\"

	#temp local rig, dont let image name confuse the graph labels
	img_name = 'tempe_image.png'
	temp_fp = base_image_path+img_name
	
	#direct reference to the df column header
	col_tag = col_tag
	clean_metric_name = str()
	fig_n = int()
	fig_mean = float()
	fig_std = float()
	fig_y_units = str('Incumbents')
	data_vector = None
	
	
	#first lets handle global IPA for the metric set
	
	#We will want to capture the transformed columns as well; which is all the metrics with suffix _sqrt
	if 'Overall Team Avg. Score' in col_tag.split('_LOL_')[-1]:
		#do team level histos and stats
		team_level_frame = pd.pivot_table(working_data_frame,values=col_tag,index='Team ID',aggfunc=np.max)
		data_vector = team_level_frame[col_tag]
		clean_metric_name = '{} - {}'.format(col_tag.split('_LOL_')[0],col_tag.split('_LOL_')[-1])
		fig_y_units = str('Teams')
			
	elif 'Manager Fit Score' in col_tag.split('_LOL_')[-1]:
		data_vector = working_data_frame[col_tag].dropna()
		clean_metric_name = '{} - {}'.format(col_tag.split('_LOL_')[0],col_tag.split('_LOL_')[-1])
		fig_y_units = str('Incumbent-Manager Dyads')
	
	else:
		data_vector = working_data_frame[col_tag]
		clean_metric_name = '{} - {}'.format(col_tag.split('_LOL_')[0],col_tag.split('_LOL_')[-1])	
		
		
	if '_sqrt' in clean_metric_name:
		clean_metric_name = clean_metric_name.replace('_sqrt',' (Transformed)')
	
	fig_n = len(data_vector)
	fig_mean = data_vector.mean()
	fig_std = data_vector.std()
	
	
	if 'Employee Facet' in col_tag.split('_LOL_')[-1]:		
		fig_title_tag = '{}\n\
		(n = {}, mean = {:.2f}, std = {:.2f})'.format(clean_metric_name,fig_n,\
		fig_mean, fig_std)	
	
	else:
		fig_title_tag = '{}\n\
		(n = {}, mean = {:.4f}, std = {:.4f})'.format(clean_metric_name,fig_n,\
		fig_mean, fig_std)
	
	
	fig, ax = plt.subplots()
	
	a_heights, a_bins = np.histogram(data_vector,bins=30)
	width = ((a_bins[1] - a_bins[0])/3)
	ax.bar(a_bins[:-1], a_heights, width=width, facecolor='purple')
	ax.set_title(fig_title_tag)
	ax.set_xlabel(clean_metric_name)
	
	#this will vary whether we are looking at team or incumbs
	ax.set_ylabel('Number {}'.format(fig_y_units))
	
	#title_tag = 'Figure fig_num: {}'.format(clean_metric_name)
	#plt.legend(['Does this work'], loc='best', ncol=1, fancybox=True, shadow=True)
	#plt.title(fig_title_tag)
	
	plt.savefig(temp_fp)
	#plt.show()
	plt.close()
	return temp_fp
	

def build_td_metric_namelist():

	'''"Team Dynamic Metrics_LOL_Manager Fit Distance (Raw)	
	Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)	
	Team Dynamic Metrics_LOL_Avg. Distance to Team (Raw)	
	Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)	
	Team Dynamic Metrics_LOL_Overall Team Avg. Distance (Raw)	
	Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)	
	
	Team Dynamics Metrics are our major interest here, and their dedicated columns
	appear as above
	
	for each of the 23 facets, a similar column set is created with prefix facet name
	example for appeasement below
	
	Appeasement_LOL_Manager Fit Distance (Raw)	
	Appeasement_LOL_Manager Fit Score (Standardized)	
	Appeasement_LOL_Avg. Distance to Team (Raw)	
	Appeasement_LOL_Avg. Score to Team (Standardized)	
	Appeasement_LOL_Overall Team Avg. Distance (Raw)	
	Appeasement_LOL_Overall Team Avg. Score (Standardized)"
	'''
	
	#for this investigation our primary interests are each of the (Standardized) versions
	#this will yield 3 "metrics" for team dynamics and also 3 "metrics" per 23 facets
	td_metric_namelist = ['Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)',\
	'Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)',\
	'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)']
	
	facet_namelist = build_facet_namelist()
	
	for i in facet_namelist:
		td_metric_namelist.append('Employee Facet Score_LOL_{}'.format(i))
		td_metric_namelist.append('{}_LOL_Manager Fit Score (Standardized)'.format(i))
		td_metric_namelist.append('{}_LOL_Avg. Score to Team (Standardized)'.format(i))
		td_metric_namelist.append('{}_LOL_Overall Team Avg. Score (Standardized)'.format(i))
		
	
	return td_metric_namelist

	
def build_team_dist_metric_namelist():

	'''"
	
	Team Distance Metrics are our major interest here, and their dedicated columns
	
	for each of the 23 facets, a similar column set is created with prefix facet name
	example for appeasement below
	
	Appeasement_LOL_Manager Fit Distance (Raw)	
	Appeasement_LOL_Manager Fit Score (Standardized)	
	Appeasement_LOL_Avg. Distance to Team (Raw)	
	Appeasement_LOL_Avg. Score to Team (Standardized)	
	Appeasement_LOL_Overall Team Avg. Distance (Raw)	
	Appeasement_LOL_Overall Team Avg. Score (Standardized)"
	'''
	
	#for this investigation our primary interests are each of the (Standardized) versions
	#this will yield 3 "metrics" for team dynamics and also 3 "metrics" per 23 facets
	team_dist_metric_namelist = ['Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_sqrt']
	
	facet_namelist = build_facet_namelist()
	
	for i in facet_namelist:
		team_dist_metric_namelist.append('{}_LOL_Avg. Score to Team (Standardized)_sqrt'.format(i))
	
	return team_dist_metric_namelist

	
def	build_mgr_dist_metric_namelist():

	mgr_dist_metric_namelist = ['Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_sqrt']
	facet_namelist = build_facet_namelist()
	for i in facet_namelist:
		mgr_dist_metric_namelist.append('{}_LOL_Manager Fit Score (Standardized)_sqrt'.format(i))	
	
	return mgr_dist_metric_namelist

	
def build_ot_dist_metric_namelist():
	
	ot_dist_metric_namelist = ['Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_sqrt']
	facet_namelist = build_facet_namelist()
	for i in facet_namelist:
		ot_dist_metric_namelist.append('{}_LOL_Overall Team Avg. Score (Standardized)_sqrt'.format(i))	
	
	return ot_dist_metric_namelist
	
	
def build_histo_metric_namelist():

	'''"
	
	Team Dynamics Metrics are our major interest here, and their dedicated columns
	appear as above
	
	fig_tags can split
	
	for each of the 23 facets, a similar column set is created with prefix facet name
	example for appeasement below
	
	Appeasement_LOL_Manager Fit Distance (Raw)	
	Appeasement_LOL_Manager Fit Score (Standardized)	
	Appeasement_LOL_Avg. Distance to Team (Raw)	
	Appeasement_LOL_Avg. Score to Team (Standardized)	
	Appeasement_LOL_Overall Team Avg. Distance (Raw)	
	Appeasement_LOL_Overall Team Avg. Score (Standardized)"
	'''
	
	#for this investigation our primary interests are each of the (Standardized) versions
	#this will yield 3 "metrics" for team dynamics and also 3 "metrics" per 23 facets
	histo_metric_namelist = ['Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)',\
	'Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)_sqrt',\
	'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)',\
	'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)_sqrt',\
	'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)',\
	'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)_sqrt']
	
	facet_namelist = build_facet_namelist()
	
	for i in facet_namelist:
		histo_metric_namelist.append('Employee Facet Score_LOL_{}'.format(i))
		histo_metric_namelist.append('{}_LOL_Avg. Score to Team (Standardized)'.format(i))
		histo_metric_namelist.append('{}_LOL_Avg. Score to Team (Standardized)_sqrt'.format(i))
		histo_metric_namelist.append('{}_LOL_Manager Fit Score (Standardized)'.format(i))
		histo_metric_namelist.append('{}_LOL_Manager Fit Score (Standardized)_sqrt'.format(i))
		histo_metric_namelist.append('{}_LOL_Overall Team Avg. Score (Standardized)'.format(i))
		histo_metric_namelist.append('{}_LOL_Overall Team Avg. Score (Standardized)_sqrt'.format(i))
		
	
	return histo_metric_namelist

	
def	quick_teamsize_stats(working_data_frame):
	
	working_data_frame = working_data_frame
	team_size_min = float()
	team_size_max = float()
	team_size_mean = float()
	team_size_median = float()
	team_size_std = float()
	try:
		team_size_min = working_data_frame['Team ID'].value_counts().min()
		team_size_max = working_data_frame['Team ID'].value_counts().max()
		team_size_median = working_data_frame['Team ID'].value_counts().median()
		team_size_mean = working_data_frame['Team ID'].value_counts().mean()
		team_size_std = working_data_frame['Team ID'].value_counts().std()
	except:
		pass
	
	return team_size_min, team_size_max, team_size_median, team_size_mean, team_size_std


def build_sqrt_cols(working_data_frame):
	
	working_data_frame = working_data_frame
	
	td_metric_namelist = build_td_metric_namelist()
	
	lil_mean = float()
	lil_std = float()
	team_level_frame = pd.DataFrame()
	
	for head in td_metric_namelist:
		if 'Avg. Score to Team' in head or 'Manager Fit Score' in head:
			working_data_frame[head+'_sqrt'] = working_data_frame[head]**.5
			working_data_frame[head+"_sqrt_mean"] = working_data_frame[head+'_sqrt'].mean()
			working_data_frame[head+"_sqrt_std"] = working_data_frame[head+'_sqrt'].std()
			working_data_frame[head+"_sqrt_z_score"] = \
			(working_data_frame[head+'_sqrt'] - working_data_frame[head+'_sqrt_mean'])*1.0/	working_data_frame[head+"_sqrt_std"]
			
		elif 'Overall Team Avg. Score' in head:
			team_level_frame = pd.pivot_table(working_data_frame,values=head,index='Team ID',aggfunc=np.max)
			team_level_frame[head+'_sqrt'] = team_level_frame[head]**.5
			team_level_frame[head+'_sqrt_mean'] = team_level_frame[head+'_sqrt'].mean()
			team_level_frame[head+'_sqrt_std'] = team_level_frame[head+'_sqrt'].std()
			team_level_frame[head+'_sqrt_z_score'] = \
			(team_level_frame[head+'_sqrt'] - team_level_frame[head+'_sqrt_mean'])*1.0/	team_level_frame[head+"_sqrt_std"]
			del team_level_frame[head]
			working_data_frame = working_data_frame.merge(\
			team_level_frame,how='outer',left_on='Team ID',right_index=True)
			
		else:
			pass
						
	return working_data_frame

	
def build_client_sqrt_cols(working_data_frame):
	
	#this gives you all the transformations, z scores, percentiles and labels
	#as new headers in the frame with known suffices
	
	working_data_frame = working_data_frame
	
	tps_pop_params_frame = build_tps_pop_params_frame()
	mgr_fit_pops_frame = build_mgrfit_pop_params_frame()
	otc_pops_frame = build_otc_pop_params_frame()
	
	td_avg_dist_pop_mean = tps_pop_params_frame.iloc[0,1]
	td_avg_dist_pop_std = tps_pop_params_frame.iloc[0,2]
	
	mgr_fit_pop_mean = mgr_fit_pops_frame.iloc[0,1]
	mgr_fit_pop_std = mgr_fit_pops_frame.iloc[0,2]
	
	otc_pop_mean = otc_pops_frame.iloc[0,1]
	otc_pop_std = otc_pops_frame.iloc[0,2]
	
	td_metric_namelist = ['Team Dynamic Metrics_LOL_Avg. Score to Team (Standardized)',\
	'Team Dynamic Metrics_LOL_Manager Fit Score (Standardized)',\
	'Team Dynamic Metrics_LOL_Overall Team Avg. Score (Standardized)']
	
	lil_mean = float()
	lil_std = float()
	team_level_frame = pd.DataFrame()
	
	for head in td_metric_namelist:
		if 'Avg. Score to Team' in head:
			working_data_frame[head+'_sqrt'] = working_data_frame[head]**.5			
			working_data_frame[head+"_sqrt_z_score"] = \
			(working_data_frame[head+'_sqrt'] - td_avg_dist_pop_mean)*1.0/td_avg_dist_pop_std			
			working_data_frame[head+"_Cohesion_Percentile"] = st.norm.cdf(working_data_frame[head+"_sqrt_z_score"])
			working_data_frame[head+"_Cohesion_Label"] = \
			write_cohesion_label(working_data_frame[head+"_Cohesion_Percentile"].tolist())
		
		elif 'Manager Fit Score' in head:
			
		#	data_vector = working_data_frame[col_tag].dropna()
			
			working_data_frame[head+'_sqrt'] = working_data_frame[head]**.5
			working_data_frame[head+"_sqrt_z_score"] = \
			(working_data_frame[head+'_sqrt'] - mgr_fit_pop_mean)*1.0/mgr_fit_pop_std			
			working_data_frame[head+"_Cohesion_Percentile"] = st.norm.cdf(working_data_frame[head+"_sqrt_z_score"])
			working_data_frame[head+"_Cohesion_Label"] = \
			write_cohesion_label(working_data_frame[head+"_Cohesion_Percentile"].tolist())
			
		elif 'Overall Team Avg. Score' in head:
			team_level_frame = pd.pivot_table(working_data_frame,values=head,index='Team ID',aggfunc=np.max)
			team_level_frame[head+'_sqrt'] = team_level_frame[head]**.5
			#custom z params
			team_level_frame[head+'_sqrt_z_score'] = \
			(team_level_frame[head+'_sqrt'] - otc_pop_mean)*1.0/otc_pop_std
			del team_level_frame[head]
			working_data_frame = working_data_frame.merge(\
			team_level_frame,how='outer',left_on='Team ID',right_index=True)
			working_data_frame[head+"_Cohesion_Percentile"] = st.norm.cdf(working_data_frame[head+"_sqrt_z_score"])
			working_data_frame[head+"_Cohesion_Label"] = \
			write_cohesion_label(working_data_frame[head+"_Cohesion_Percentile"].tolist())

			
		else:
			pass
						
	return working_data_frame


def build_facet_metric_cols(working_data_frame):
	
	#this gives you all the transformations, z scores, percentiles and labels
	#as new headers in the frame with known suffices
	
	#header structure is
	#Facet_LOL_Avg. Score to Team (Standardized)
	#Facet_LOL_Manager Fit Score (Standardized)
	#Facet_LOL_Overall Team Avg. Score (Standardized)
	
	#header names of the tps_pops_frame:
	#Metric: "Team Dynamics Metric TPS", all the 23 facets by name
	#Pop_Mean, Pop_Std_Dev, Norming_sample_size_teammembers
	lil_mean = float()
	lil_std = float()
	team_level_frame = pd.DataFrame()
	
	tps_pop_params_frame = build_tps_pop_params_frame()
	mgr_fit_pops_frame = build_mgrfit_pop_params_frame()
	otc_pops_frame = build_otc_pop_params_frame()
	
	working_data_frame = working_data_frame
	td_metric_namerootlist = ['_LOL_Avg. Score to Team (Standardized)',\
	'_LOL_Manager Fit Score (Standardized)',\
	'_LOL_Overall Team Avg. Score (Standardized)']
	
	
	lil_head = str()
	
	for i in range(1,len(tps_pop_params_frame)):
		for head in td_metric_namerootlist:
			lil_facet_head = tps_pop_params_frame.iloc[i,0]+head
		
			if 'Avg. Score to Team (Standardized)' in lil_facet_head.split('_'):
				working_data_frame[lil_facet_head+'_sqrt'] = working_data_frame[lil_facet_head]**.5			
				working_data_frame[lil_facet_head+"_sqrt_z_score"] = \
				(working_data_frame[lil_facet_head+'_sqrt'] - tps_pop_params_frame.iloc[i,1])*1.0/tps_pop_params_frame.iloc[i,2]			
				working_data_frame[lil_facet_head+"_Cohesion_Percentile"] = st.norm.cdf(working_data_frame[lil_facet_head+"_sqrt_z_score"])
				working_data_frame[lil_facet_head+"_Cohesion_Label"] = \
				write_cohesion_label(working_data_frame[lil_facet_head+"_Cohesion_Percentile"].tolist())
			
			elif 'Manager Fit Score (Standardized)' in lil_facet_head.split('_'):	
				working_data_frame[lil_facet_head+'_sqrt'] = working_data_frame[lil_facet_head]**.5
				working_data_frame[lil_facet_head+"_sqrt_z_score"] = \
				(working_data_frame[lil_facet_head+'_sqrt'] - mgr_fit_pops_frame.iloc[i,1])*1.0/mgr_fit_pops_frame.iloc[i,2]			
				working_data_frame[lil_facet_head+"_Cohesion_Percentile"] = st.norm.cdf(working_data_frame[lil_facet_head+"_sqrt_z_score"])
				working_data_frame[lil_facet_head+"_Cohesion_Label"] = \
				write_cohesion_label(working_data_frame[lil_facet_head+"_Cohesion_Percentile"].tolist())
				
			elif 'Overall Team Avg. Score (Standardized)' in lil_facet_head.split('_'):
				team_level_frame = pd.pivot_table(working_data_frame,values=lil_facet_head,index='Team ID',aggfunc=np.max)
				team_level_frame[lil_facet_head+'_sqrt'] = team_level_frame[lil_facet_head]**.5
				#custom z params
				team_level_frame[lil_facet_head+'_sqrt_z_score'] = \
				(team_level_frame[lil_facet_head+'_sqrt'] - otc_pops_frame.iloc[i,1])*1.0/otc_pops_frame.iloc[i,2]
				del team_level_frame[lil_facet_head]
				working_data_frame = working_data_frame.merge(\
				team_level_frame,how='outer',left_on='Team ID',right_index=True)
				working_data_frame[lil_facet_head+"_Cohesion_Percentile"] = st.norm.cdf(working_data_frame[lil_facet_head+"_sqrt_z_score"])
				working_data_frame[lil_facet_head+"_Cohesion_Label"] = \
				write_cohesion_label(working_data_frame[lil_facet_head+"_Cohesion_Percentile"].tolist())
				
			else:
				pass
	
	return working_data_frame


def write_cohesion_label(cohesion_percentile_list):
	
	cohesion_percentile_list = cohesion_percentile_list
	cohesion_label_list = list()
	
	for i in cohesion_percentile_list:
		if float(i) < 0.25:
			cohesion_label_list.append('Very Diverse')
		elif 0.25 <= float(i) < 0.50:
			cohesion_label_list.append('Diverse')
		elif 0.50 <= float(i) < 0.75:
			cohesion_label_list.append('Cohesive')
		else:
			cohesion_label_list.append('Very Cohesive')	
	
	return cohesion_label_list


def compute_chi_sq(working_array):
	chi2_data = working_array
	min_exp_freq = expected_freq(chi2_data).min()
	try:
		chi2, p, dof, expected = chi2_contingency(chi2_data, correction=False)
		dof = int(dof)
	except:
		chi2, p, dof = ''
	return chi2, p, dof


def build_term_table_frame(working_data_frame,labels_head,term_head,hypo_hi_term_classlabel,hypo_lo_term_classlabel):
	
	#represents a list of cohesion labels which can have four values
	#We depend on the users labeling system to extract the correction
	#labels column and termination column for analysis
	
	working_data_frame = working_data_frame
	labels_head = labels_head
	term_head = term_head
	
	
	#here we handle custom hypothesis directionality by having the function
	#receive theh manual classlabels
	#for example Very Diverse may be the classlabel expected for highest termination, comparable to NR
	#because all comparison classes are not standard, we just pass the labels
	#and compute our relative turnover percentages, aka, Pct Diff, for the given hypothesized groups,
	#which may or may not be all the groups used in the chi squared test, this is purely to interpret directionality and support
	
	hi_terms = hypo_hi_term_classlabel
	low_terms = hypo_lo_term_classlabel
	
	
	term_freqs = working_data_frame.groupby([labels_head,term_head]).size()
	table_freqs_frame = term_freqs.unstack(term_head)
	table_freqs_frame = table_freqs_frame.fillna(0)
	
	support_note = str()
	
	#table_freqs_frame is a dataframe with the label classes on the index
	#the 0 and 1 as columns with respective frequency counts in each cell
	#here we know that 0 represents Not Termed, 1 = Termed
	
	try:
	
		table_freqs_frame['N'] = table_freqs_frame.sum(axis=1)
		table_freqs_frame['Termed_Freq.'] = table_freqs_frame[1]
		table_freqs_frame['Termed_Pct.'] = table_freqs_frame['Termed_Freq.']*1.0 / table_freqs_frame['N']
		
		table_freqs_frame['Retained_Freq.'] = table_freqs_frame[0]
		table_freqs_frame['Retained_Pct.'] = table_freqs_frame['Retained_Freq.']*1.0 / table_freqs_frame['N']
		
		local_chi2, local_p, local_dof = compute_chi_sq(table_freqs_frame[[0,1]].T)	
		
		table_freqs_frame['Chi_sq'] = local_chi2
		table_freqs_frame['p_value'] = local_p
		table_freqs_frame['dof'] = local_dof
		
		if table_freqs_frame['Termed_Pct.'][hi_terms] == 0:
			table_freqs_frame['Pct Diff'] = 'No Terms in Expected Group'
		else:
			try:
				table_freqs_frame['Pct Diff'] = \
				(table_freqs_frame['Termed_Pct.'][hi_terms] - table_freqs_frame['Termed_Pct.'][low_terms])/table_freqs_frame['Termed_Pct.'][hi_terms]
			except:
				table_freqs_frame['Pct Diff'] = 'Error computing Pct Diff'
			
	except:
		table_freqs_frame['N'] = ''
		table_freqs_frame['Termed_Freq.'] = ''
		table_freqs_frame['Termed_Pct.'] = ''
		table_freqs_frame['Retained_Freq.'] = ''
		table_freqs_frame['Retained_Pct.'] = ''
		table_freqs_frame['Chi_sq'] = ''
		table_freqs_frame['p_value'] = ''
		table_freqs_frame['dof'] = ''
		table_freqs_frame['Pct Diff'] = ''
		
	return table_freqs_frame


def build_otc_t_table_frame(working_data_frame,labels_head,term_head,hypo_hi_term_classlabel,hypo_lo_term_classlabel):
	
	#represents a list of cohesion labels which can have four values
	#We depend on the users labeling system to extract the correction
	#labels column and termination column for analysis
	
	working_data_frame = working_data_frame
	labels_head = labels_head
	term_head = term_head
	
	#here we handle custom hypothesis directionality by having the function
	#receive theh manual classlabels
	#for example Very Diverse may be the classlabel expected for highest termination, comparable to NR
	#because all comparison classes are not standard, we just pass the labels
	#and compute our relative turnover percentages, aka, Pct Diff, for the given hypothesized groups,
	
	#In this function, the focus is on team-level termination rates which first requires team-level aggregation
	#Then, we apply ANOVA, or t-test on two focus groups, and the metric reported is a D-value (effect size)
	#and reported p value reflects the available extreme group comparison
	
	hi_terms = hypo_hi_term_classlabel
	low_terms = hypo_lo_term_classlabel
	
	table_otc_t_test_frame = pd.DataFrame()
	
	try:
		#START YOUR TRY EXCEPT HERE
		print("inner otc 1")
		#table_anova_frame	
		team_freqs = working_data_frame.groupby('Team ID').agg({'Team ID': ['count'], term_head: ['sum'], labels_head: ['max']})
		team_freqs.columns = team_freqs.columns.get_level_values(0)
		
		#here the column names are produced by the above function
		#so term_head is the sum of terms in the team, and Team ID is the count of team size
		#we can divide term pct by team size to compute team term pct
		team_freqs['Team Term Pct'] = team_freqs[term_head] / team_freqs['Team ID']
		
		#Team Term Pct, then is our new team-level metric which will be analyzed within cohesion labels
		#in this frame, the cohesion labels are under the column represented by labels_head
		
		#Here we pull out the specific cells for t-testing
		hi_terms_array = team_freqs[team_freqs[labels_head] == hi_terms]['Team Term Pct'].values
		low_terms_array = team_freqs[team_freqs[labels_head] == low_terms]['Team Term Pct'].values
		
		#Here we are assuming equal variances in line with the DA workbook
		t_test_returns = stats.ttest_ind(hi_terms_array,low_terms_array,equal_var=True)
		t_stat = t_test_returns[0]
		p_stat = t_test_returns[1]
		
		
		table_otc_t_test_frame = team_freqs.groupby([labels_head]).agg({'Team Term Pct': ['count','mean','std']})
		#this returns two columns of same name and a hierarchical index. next few lines just cleans this up
		#to fit our formatting expectations
		
		table_otc_t_test_frame.columns = table_otc_t_test_frame.columns.get_level_values(0)
		table_otc_t_test_frame = table_otc_t_test_frame.fillna(0)
		table_otc_t_test_frame['N'] = table_otc_t_test_frame.iloc[:,0]
		table_otc_t_test_frame['Avg. Term Pct'] = table_otc_t_test_frame.iloc[:,1]
		table_otc_t_test_frame['SD. Term Pct'] = table_otc_t_test_frame.iloc[:,2]
		del table_otc_t_test_frame['Team Term Pct']
		
		this_d = convert_t_to_cohensd(low_terms_array,hi_terms_array)
		
		table_otc_t_test_frame['Cohens_D'] = this_d
		
		try:
			table_otc_t_test_frame['p_value'] = p_stat/2.0
		except:
			table_otc_t_test_frame['p_value'] = 1.0
		
		table_otc_t_test_frame['dof'] = len(hi_terms_array)+len(low_terms_array)-1
		
		print("inner otc 2")
		#the division by 2 is used to have scipy produce a 1-tailed interpretation
		#this reproduces the excel calc, but not sure if this is really accurate approach
		table_otc_t_test_frame = table_otc_t_test_frame.fillna({'p_value':1.0})
		#table_otc_t_test_frame is a dataframe with the label classes on the index
		#the team term rate avg and stdev within each team type, will also produce a d (effect size) based on the t-test
		#and the p-value for that t-test, maybe put the t-stat in the footnote
		#remember the t-test just compares the extreme hypthosized groups even if more than 2 groups are displayed
		#we set the t-test to false independence assumption, and note that the computed p is the two-tailed value
	
	except:
		table_otc_t_test_frame = pd.DataFrame({'N': [0], 'Avg. Term Pct': [0], \
		'SD. Term Pct': [0], 'Cohens_D': [-99.0], 'p_value': [1.0], 'dof':[0]})
		print("inner otc 3")
		
	return table_otc_t_test_frame

	
def build_tps_pop_params_frame():
	
	#header names:
	#Metric: "Team Dynamics Metric TPS", all the 23 facets by name
	#Pop_Mean	
	#Pop_Std_Dev	
	#Norming_sample_size_teammembers
	
	working_path = "C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\PA42647_TD_FacetLevel_Quarterly\\"
	tps_master_csv_file = "teamplayerscore_pop_params.csv"	
	tps_pop_params_frame = pd.read_csv(working_path+tps_master_csv_file)
	
	return tps_pop_params_frame

	
def build_mgrfit_pop_params_frame():
	
	#header names:
	#Metric: "Team Dynamics Metric MgrFit", all the 23 facets by name
	#Pop_Mean	
	#Pop_Std_Dev	
	#Norming_sample_size_incumbentmgrdyads
	
	working_path = "C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\PA42647_TD_FacetLevel_Quarterly\\"
	mgrfit_master_csv_file = "mgrfitscore_pop_params.csv"	
	mgrfit_pop_params_frame = pd.read_csv(working_path+mgrfit_master_csv_file)
	
	return mgrfit_pop_params_frame

		
def build_otc_pop_params_frame():
	
	#header names:
	#Metric: "Team Dynamics Metric OTC", all the 23 facets by name
	#Pop_Mean	
	#Pop_Std_Dev	
	#Norming_sample_size_teams
	
	working_path = "C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\PA42647_TD_FacetLevel_Quarterly\\"
	otc_master_csv_file = "overallteamchem_pop_params.csv"
	otc_pop_params_frame = pd.read_csv(working_path+otc_master_csv_file)
	
	return otc_pop_params_frame	

	
def convert_t_to_cohensd(low_terms_array,hi_terms_array):
	
	#Here we follow the cohens d conversion formula found in the DA "Workbook for Analyses.xlsx"
	#For directionality, we take hypothetical hi_terms_array as group2
	#then low_terms_array represents group1
	
	#we do a sample-weighted mean diff as the numerator
	#and we do a sample weighted standard deviation as denominator
	low_terms_array = low_terms_array
	hi_terms_array = hi_terms_array
	
	try:
	
		cohens_d = float()
			
		n1 = len(low_terms_array)
		mean1 = mean(low_terms_array)
		
		#ddof = 1 reflects the sample estimate and matches the pandas agg function 'std'
		std1 = std(low_terms_array,ddof=1)
			
		n2 = len(hi_terms_array)
		mean2 = mean(hi_terms_array)
		std2 = std(hi_terms_array,ddof=1)
		
		numerator = (mean2-mean1)*1.0*(n1+n2)
		denominator = (((n1-1)*1.0*(std1**2) + (n2-1)*1.0*(std2**2))*(n1+n2))**0.5
		
		cohens_d = numerator / denominator
		
	except:
		cohens_d = float(-99)
	
	return cohens_d
	
	
def plot_survival_curve(working_frame,header):
	
	working_frame = working_frame
	
	kmf = KaplanMeierFitter()
	kmf.fit(working_frame['career_length'], event_observed=working_frame['censor'])
	p = kmf.plot(ci_force_lines=True, title='Career Lengths of Team players')
	kmf1 = plt.gcf()
	pyplot(kmf1)
	return

	
def pyplot(fig, ci=True, legend=True):
    # Convert mpl fig obj to plotly fig obj, resize to plotly's default
    py_fig = tls.mpl_to_plotly(fig, resize=True)
    
    # Add fill property to lower limit line
    if ci == True:
        style1 = dict(fill='tonexty')
        # apply style
        py_fig['data'][2].update(style1)
        
        # Change color scheme to black
        py_fig['data'].update(dict(line=Line(color='black')))
    
    # change the default line type to 'step'
    py_fig['data'].update(dict(line=Line(shape='hv')))
    # Delete misplaced legend annotations 
    py_fig['layout'].pop('annotations', None)
    
    if legend == True:
        # Add legend, place it at the top right corner of the plot
        py_fig['layout'].update(
            showlegend=True,
            legend=Legend(
                x=1.05,
                y=1
            )
        )
        
    # Send updated figure object to Plotly, show result in notebook
    return py.iplot(py_fig)
	
	
def main():

	fp = "C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\"
	#fname_list = ["ex_td_facets_report_study1.docx","ex_td_facets_report_study2.docx"]
	
	study1_write_out_fp = fp+"ex_td_facets_report_study1.docx"
	study2_write_out_fp = fp+"ex_td_facets_report_study2.docx"
	study3_write_out_fp = fp+"ex_td_facets_report_study3_genderfit.docx"
	
	
	workdata_fp = "C:\\Users\\pgilmore\\Desktop\\Development\\Team_Dynamics\\PA42647_TD_FacetLevel_Quarterly\\"
	workdata_fp = workdata_fp+"td_facet_masterbook_plg4.xlsx"
	
	#study 1 is population estimation
	#write_report_study1(workdata_fp, study1_write_out_fp)
	
	#study 2 is local criterion evaluation on lowes turnover data
	#write_report_study2(workdata_fp, study2_write_out_fp)
	
	#study 3 is same as study 2, but split based on gender fit
	write_report_study3(workdata_fp, study3_write_out_fp)
	
	
	print("\n\nCheck your filepath for {} and {}".format(study1_write_out_fp,study2_write_out_fp))
	
	return


	
if __name__ == '__main__':
	main()
