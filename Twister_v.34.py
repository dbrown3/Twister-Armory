"""
Originally Created on Tue Sep 02 13:25:55 2014
@author: GILMORE and dbrown3
Two Heavens Style
"""

'''
Originally, Spin

an app designed to process, format and report analyses
on the quality and statistical characteristics of performance evaluation
datasets in .csv format with known format (survey management exports, 'Answers'
format, uses numerical 1 - 7 values for item response scales, 
with optional final confidence item using letter-based
values "CA", SA", etc.)

Evolved to, Twister 

added the capability to merge personality data (predictors), thus
completing data preparation and preliminary analyses for modelling and
classification tasks.
personality data in known .csv format (ppi2+pma), exported from science matching tool (profile creation match),
the matching session dependent upon information from 
the first stage output of Spin.

Outputs in .xlsx workbooks and .docx report

Twister is developed and maintained by Derek Brown and Phillip Gilmore
with alpha community support from the Behavioral Analyst team 
at Infor Talent Science
Dallas, TX
'''


#import your libraries
import numpy as np
import matplotlib.pyplot as plt
from scipy import stats as stats
#import random

#from scipy import stats as stats
from pandas import Series, DataFrame
import pandas as pd
from pandas import ExcelWriter

#import xlwt

#string library will help me parse the file path string for easier use
#in the output filename
import string

#Tkinter is a GUI library from standard python
#it has a lot of options, this bare-bones was found online just to get
#the users file location
from Tkinter import Tk
from tkFileDialog import askopenfilename

#use the datetime library
#from datetime import datetime
from datetime import datetime
#from datetime import timedelta



'''SYS CONFIG'''
#need to make it robust to variable encoding settings
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

#related to chained operators and setting to copies of slices
#this is a workaround to errors thrown at runtime
pd.set_option('mode.chained_assignment',None)


#******************************************************************************
"""Notes for TWISTER

#DATAFRAME NAMES
#data - from spin and is the complete data frame
#data_match - everyone in the assessment match along with QC removals and it's Dummy code
#Clean_data_match - data_match data frame with the QC flags removed
#Dimension_descriptives2 - clean_data_match sample with sample and population assessment values and z scores
#Assessment_outlier_dummy_code_columns - clean_data_match sample with z scores, dummy codes, and removal
#spin_master - is Spin's output with Random ID as the index and as a column - first columns are in matching order
#full_data - master file output with the Random ID as the index and both spin and twister package results
#Proceed1 - used to create user input prior to TK interp coming up for eval
#Proceed 2 - allows user to check job titles in the sample prior to program completion
#Proceed3 - used to create user input prior to TK interp coming up for match
#start_item_column - the column number where the items begin according to Ratee Status
"""
#******************************************************************************

#Version dependent instructions
print "\n\n\n\nWelcome to Twister\
\n\nA program for Performance Evaluation Analysis\
\n\nTo get started, let's read in the performance evaluation file\
\nin .csv format. This program is designed to work with\
\nevaluation exports in 'Answers' format.\
\nPlease find the evaluation export and 'save as' into the .csv format\
\nin a location that is easy to remember and access.\
\nNow open the file so that you can view it."

Proceed1 = raw_input("\nIf you saved your evaluation as a csv into a file and \
opened it, press Enter.")

print "\nNow, open the eval and please ensure the following column labels\
\nmatch these 'Ratee Unique ID' and 'Rater Unique ID'\
\n(capitalization and spacing are important)\n\
\nPlease clean out any duplicate Ratees before using this tool.\
\n\nAlso, ensure your items start on the standard column.\
\nThat should be Column 'S'.\
\n\nEnsure the file has not been misaligned or otherwise\
\ntampered with. If unsure, you can export the 'Answers'\
\nfile yourself from Survey Management.\n\
\nIn a few moments, after you select your data file\n\
and check job titles, the program will ask you for the number of items.\n\
\nDo not count the confidence item nor\
\nsubsequent items (i.e., tenure category) at this stage.\
Open the Profile Checklist as it will provide necessary information.\
\n\nNow, find the file explorer window to access your eval file\
\nthat you saved into the .csv format.\
\nIt may be behind this window, so please minimize to find it."




'''GETTING USER INPUT'''

Tk().withdraw()
filename = askopenfilename()

#NOTES
#TKinter produced the filename in unicode format
#a string version of the input file name is stripped from 
#the full unicode filepath
#string library is used

in_name_list = string.split(filename,sep='/')
in_name = str(in_name_list[-1].rstrip('.csv'))

print "\n>>>>>\n>>>>>"
print "\nThanks!"
print "So we are working with '{}.csv'.".format(in_name)



'''READING THE CSV FILE INTO A PANDAS DATAFRAME'''
#the encoding options were specified to trouble-shoot the output file
#that sometimes comes out as .csv and sometimes as .xlsx, and sometimes throws
#an error.
try:
    data = pd.read_csv(filename, index_col=['Ratee Unique ID', 'Rater Unique ID'], \
    na_values=['0','NULL','-99', 'N/A'], encoding='cp1252')    
    
except:
    print 'WARNING: Encoding defaulted to cp1252'    
    data = pd.read_csv(filename, index_col=['Ratee Unique ID', 'Rater Unique ID'], \
    na_values=['0','NULL','-99', 'N/A'], encoding='utf-8')




'''CLARIFYING JOB TITLES FOR ANALYSES'''

#All_titles =  set([x for x in data['Job Title'] if str(x) != "nan"])
#Job_titles = ", ".join(str(x) for x in All_titles)

print "\nThe given job title(s) in this data file\n\
before any removals are as follows: \n"
print "\n", data['Job Title'].value_counts()

Proceed2 = raw_input("\nIf any titles are incorrect, you can remove those\n\
cases in the original data file and re-run Twister.\n\
Any job titles remaining in the data file will be kept for later analyses.\n\
\nIf you are ready to continue,\n\
please press Enter.")



'''RETRIEVE THE NUMBER OF ITEMS FOR MANY IMPORTANT OPERATIONS'''
k = int(raw_input("\nHow many items are in the scale? \n\
(do not count the final confidence item): "))



'''COMPUTING STANDARD ROW-WISE STATISTICS'''

start_item_column = len(data.ix[:,:'Ratee Status'].columns)

data['Item Avg'] = data.ix[:,start_item_column:(start_item_column+k)].mean(axis=1,skipna=True)
data['Item Count'] = data.ix[:,start_item_column:(start_item_column+k)].count(axis=1)
data['Item Median'] = data.ix[:,start_item_column:(start_item_column+k)].median(axis=1,skipna=True)
data['Item SD.s'] = data.ix[:,start_item_column:(start_item_column+k)].std(axis=1,skipna=True)
data['Item Min'] = data.ix[:,start_item_column:(start_item_column+k)].min(axis=1,skipna=True)
data['Item Max'] = data.ix[:,start_item_column:(start_item_column+k)].max(axis=1,skipna=True)



'''INSERTING A RANDOM ID AND SETTING IT AS THE INDEX'''

data['Random ID'] = range(1,len(data)+1)
data['Random ID_index'] = data['Random ID']

data = data.reset_index()
data.set_index('Random ID_index',inplace=True)

#At this point the column number indeces have changed due to resetting the index
#previous locator variables are no longer valid.



'''REMOVAL DUMMY CODES - COMPLETELY MISSING PERFORMANCE DATA'''

data['Removal - Completely missing performance data = 1'] = \
Series((data['Item Count'] == 0),dtype=np.int32)



'''REMOVAL DUMMY CODES - PARTIALLY MISSING PERFORMANCE DATA

The lambda x function is used to pass
 variable "f" to the "apply" method which allows you
to pass more complex functions in to elements of a DataFrame.
'''

f = lambda x: x in range(1,k)

data['Removal - Partially missing performance data = 1'] = \
Series(data['Item Count'].apply(f),dtype=np.int32)



'''CLARIFYING CONFIDENCE LEVELS ANALYSES'''

print "\nNext you will enter the confidence rating threshold."
print "If the confidence item is listed as numerical values in the evaluation \
 \nthan you will have to check for rater error yourself in the final output.\n"
print "Twister recommends including confidence of 'A' and above.\n\
You may need to consider your sample size before deciding.\n\
Below are the counts of each confidence level \n\
for anyone with complete performance data:\n"

#Find the column with rater confidence item

#sub-select the valid sample as of this point in the logic for display
#of confidence level value counts
Rater_confidence_column_index = pd.Series(data.columns).str.contains("Based on")
Rater_confidence_column = Rater_confidence_column_index[Rater_confidence_column_index == True].index[0]



rata = data[(data["Removal - Completely missing performance data = 1"] == 0) \
& (data["Removal - Partially missing performance data = 1"] == 0)]

print rata.ix[:,Rater_confidence_column].value_counts()



'''RETRIEVE THE MINIMAL CONFIDENCE LEVEL'''

Rating_given = raw_input("\nEnter the lowest acceptable confidence rating. Either - A or N: ")

if Rating_given == "N":
    Rater_confidence_input = ["CD", "SD", "D"]
else:
    Rater_confidence_input = ["CD", "SD", "D", "N"]

g = lambda x: x in Rater_confidence_input



'''REMOVAL DUMMY CODES - LACK OF RATER CONFIDENCE'''

data['Removal - Lack of Rater Confidence = 1'] = Series(data.ix[:,Rater_confidence_column].apply(g), dtype = np.int32)



'''RETRIEVE THE DATE OF DATA COLLECTION FOR TENURE COMPUTATION'''

print "\nReferencing the Profile Checklist now -"

collection_date = raw_input("Enter the date of data collection. \nThis will be \
used for the tenure calculations. \n\
Use the format mm/dd/yyyy: ")

#coverts user inputted collection_date into datetime for later use in tenure calc
removal_date = datetime.strptime(collection_date,"%m/%d/%Y")


tenure_label = "Tenure (from Evaluation Hire Date to {})".format(collection_date)
#im not sure if the formatting option is necessary, it is an attempt
#to fix the tenure calculation bug.

#needed below to convert the tenure label to a generic datatype in day units.
#converted from nanoseconds.
ns_per_day = 8.64*10**13
def convert_tenure_type(x):
    try:
        return int(x)/ns_per_day
    except:
        data[tenure_label] = "Missing hire date" 



'''CLARIFYING TENURE-BASED ANALYSES'''

#This first block of input modulates whether tenure should be considered at all.
#if userq_rem_hiredate is 'Y', then all hatanalysis should be skipped and appropriate
#removal codes statically imputed.
#If it's 'N', then hatanalysis will be necessary as in old twister.
#the goal of this new block is to allow missing hire dates that don't crash the program
#or create erroneous analytical functions.

if data['Hire Date'].nunique() < 3: #a clever way to determine unusual hire date values; there may be better ways
    print "\nUnusually low number of hire dates found."
    print "\nNormally, missing hire dates are treated as removals."
    print "\nSometimes all incumbents already meet tenure requirements, \n\
    in those cases, we can skip tenure analyses."
    
    while True:   
        try:
            userq_rem_nohiredate = raw_input("\nDoes everyone in this sample already meet tenure requiremements? (Type Y or N): ").upper()
            if  userq_rem_nohiredate == 'Y' or userq_rem_nohiredate == 'N':
                break
            else:
                print "Not a valid response"
        except:
            print "Please try again."

else:
    userq_rem_nohiredate = 'N'


if userq_rem_nohiredate == 'Y':
    data[tenure_label] = "All met tenure"
    tenure_req = 0
    data['Removal - Missing tenure data = 1'] = 0
    tenure_rem_label = 'Removal - Insufficient tenure (less than {} days) = 1'.format(tenure_req)
    data[tenure_rem_label] = 0

    print "\nThanks! Hire Dates and Tenure will not be considered"
    
else:
    data[tenure_label] = removal_date - pd.to_datetime(data['Hire Date'],format="%m/%d/%Y")
    data[tenure_label] = data[tenure_label].apply(convert_tenure_type)
    
    
    
    print "\nPerfect! Now let's look at tenure.\n\
In a moment, Twister will ask for the tenure requirement.\n\
You may need to consider your sample size before deciding.\n\
To help you, Twister will display some tenure\n\
statistics.\n\
The display shows how many cases are excluded at\n\
a few common tenure cutoffs.\n\
Next to exclusion counts are the correlations between tenure\n\
and valid performance data if that tenure cutoff\n\
is chosen.\n\
Please note that these estimates have not been cleaned for\n\
assessment-side removals.\n"
    
    
    hata = data[(data["Removal - Completely missing performance data = 1"] == 0) \
    & (data["Removal - Partially missing performance data = 1"] == 0) \
    & (data["Removal - Lack of Rater Confidence = 1"] == 0)]
    
    #sub-select the data relevant to each cutoff logic
    hatanalysis_30d = hata[(hata[tenure_label] > 29)].ix[:,['Item Avg',tenure_label]]
    hatanalysis_60d = hata[(hata[tenure_label] > 59)].ix[:,['Item Avg',tenure_label]]
    hatanalysis_90d = hata[(hata[tenure_label] > 89)].ix[:,['Item Avg',tenure_label]]
    hatanalysis_180d = hata[(hata[tenure_label] > 179)].ix[:,['Item Avg',tenure_label]]
    hatanalysis_270d = hata[(hata[tenure_label] > 269)].ix[:,['Item Avg',tenure_label]]
    hatanalysis_365d = hata[(hata[tenure_label] > 364)].ix[:,['Item Avg',tenure_label]]
    
    #build a dataframe for displaying the cutoff logic and associated correlations
    df_tenure_report = DataFrame([\
    [(len(hata)-len(hatanalysis_30d)),\
    ("%.2f" % hatanalysis_30d.ix[:,0:2].corr().ix[0,1])],\
    [(len(hata)-len(hatanalysis_60d)),\
    ("%.2f" % hatanalysis_60d.ix[:,0:2].corr().ix[0,1])],\
    [(len(hata)-len(hatanalysis_90d)),\
    ("%.2f" % hatanalysis_90d.ix[:,0:2].corr().ix[0,1])],\
    [(len(hata)-len(hatanalysis_180d)),\
    ("%.2f" % hatanalysis_180d.ix[:,0:2].corr().ix[0,1])],\
    [(len(hata)-len(hatanalysis_270d)),\
    ("%.2f" % hatanalysis_270d.ix[:,0:2].corr().ix[0,1])],\
    [(len(hata)-len(hatanalysis_365d)),\
    ("%.2f" % hatanalysis_365d.ix[:,0:2].corr().ix[0,1])]],\
    index = [30,60,90,180,270,365],\
    columns=['Cases Excluded','r_(perf,tenure)'])
    
    df_tenure_report.index.name = 'Tenure cut (days)'
    
    print df_tenure_report
    
    del hatanalysis_30d,hatanalysis_60d,hatanalysis_90d,hatanalysis_180d,hatanalysis_270d,hatanalysis_365d
    #deleting these to clean up the memory
        
    
    
    '''RETRIEVE THE TENURE REQUIREMENT'''
    
    tenure_req = int(raw_input("What is the tenure requirement (in days)?: "))
        
    
    
    '''REMOVAL DUMMY CODES - TENURE'''
    
    data['Removal - Missing tenure data = 1'] = \
    Series(data[tenure_label].isnull(),dtype=np.int32)
      
    tenure_rem_label = \
    'Removal - Insufficient tenure (less than {} days) = 1'.format(tenure_req)
    
    t = lambda x: x < tenure_req
    
    if data['Removal - Missing tenure data = 1'].all() == 1:
        data[tenure_rem_label] = 0
    else:
        data[tenure_rem_label] = Series(data[tenure_label].apply(t),dtype=np.int32)



headers_perf_removals = ['Removal - Missing tenure data = 1',\
tenure_rem_label, 'Removal - Completely missing performance data = 1',\
'Removal - Partially missing performance data = 1', 'Removal - Lack of Rater Confidence = 1']


'''VISUALIZING PERFORMANCE VARIABLE'''


def hist_show(profile_frame,name): #takes a df assuming 'Item Avg' is perf column
    fig, ax = plt.subplots()
    
    
    #assumes indexed 10th column is for visualization (perf variable)    
    a_heights, a_bins = np.histogram(profile_frame.ix[:,'Z_Item Avg'])
    
    width = (a_bins[1] - a_bins[0])/3
    
    ax.bar(a_bins[:-1], a_heights, width=width, facecolor='purple')
    
    ax.set_title('Z_Performance Distribution({})'.format(name))
    ax.set_xlabel('Z_Observed Performance Score')
    ax.set_ylabel('Number Incumbents')    
    
    
    print '\nThe pop-up displays a histogram\
    \nrepresenting observed performance data for this profile\
    \nwith only performance-side removals\
    \nand z_transformed.\
    \n\nRemember to note performance outliers and\
    \nremove those cases after Twister is finished.\
    \nThis version of Twister does not support automated\
    \nperformance outlier removal.\
    \n\nYou can save the image or just close it to proceed.'
    
    plt.savefig("hist_show.png")    
    plt.show()
    


zata = data[(data[headers_perf_removals].sum(axis=1)==0)]
zata_mean = zata['Item Avg'].mean()
zata_sd = zata['Item Avg'].std()

zata.loc[:,'Z_Item Avg'] = \
    (zata['Item Avg'] - zata_mean)/zata_sd

hist_show(zata,in_name)

'''PERFORMANCE OUTLIER REMOVALS'''
#Not yet supported; may require restructuring of later column
#lookups and header lists



'''PREPARING DATAFRAMES FOR OUTPUT TO XCEL TABS'''

headers_list = []
for i in range(len(data.columns)):
    headers_list = headers_list + [str(data.columns[i])]

#First Name = headers_list[6]
#Last = 7
#SSN = 8
#email = 10
#Random ID = headers_list[-7] #
out_headers = headers_list[6:9]
out_headers.append(headers_list[10])
out_headers.append(headers_list[-7])

for i in range(0,7):
    out_headers.append(headers_list[i])

for i in range(11,len(headers_list)-7):
    out_headers.append(headers_list[i])

for i in range(len(headers_list)-5,len(headers_list)):
    out_headers.append(headers_list[i])

#MASTER TAB
#the below structure is simply formatting
#to re-order the columns for the spin_master
spin_master = data[out_headers[0:2]]
spin_master = spin_master.join(data[out_headers[3]],how='outer')
spin_master = spin_master.join(data[out_headers[2]],how='outer')
spin_master = spin_master.join(data[out_headers[4]],how='outer')

spin_master = pd.merge(spin_master,data.ix[:,0:6],how='outer',right_index =True, left_index=True)
spin_master = spin_master.join(data['Job Title'],how='outer')

#fix this so that bogus columns following the items don't end up overwriting your item statistics
#you need to seek for the column location of 'Item Avg' and use this rather than +1+6

item_avg_loc = data.columns.get_loc('Item Avg')

spin_master = pd.merge(spin_master,data.ix[:,11:(item_avg_loc+6)],\
how='outer',right_index =True, left_index=True)
spin_master = pd.merge(spin_master,data.ix[:,-6:],\
how='outer',right_index =True, left_index=True)

#the out_headers have problem matching dataframe columns headers 
#when you get to the items; probably due to special character decoding
#this is why the excessive merging just to re-order columns;
#if i could pass the out_headers list, this would be more elegant


#creating a df that can be analyzed on performance-side
data_spin_temp = data
data_spin_temp['Sum Dummies'] = data_spin_temp[headers_perf_removals].sum(axis=1)
data_spinalysis = data_spin_temp[data_spin_temp['Sum Dummies']==0]

data_spinalysis = data_spinalysis.drop('Sum Dummies',1)

#final_removal_headers = ['Removal - Not Matched = 1'] + master_tab_dummyheaders
#data11['Final Sample Code'] = data11.ix[:,final_removal_headers].sum(axis= 1)
#Final_Data = data11[data11['Final Sample Code']==0]
#Final_Data = Final_Data.drop('Final Sample Code', 1)



#SUMMARY TAB
summary = data.describe()


#are these indeces right?
data_spinalysis.ix[:,18:(18+k)].columns
item_summary = data_spinalysis.ix[:,18:(18+k)].describe()
corrs_items = data_spinalysis.ix[:,18:(18+k)].corr()



#SOURCE TABS
df_final_titles = data_spinalysis['Job Title'].value_counts()
df_final_raters = data_spinalysis['Rater Unique ID'].value_counts()
df_final_geo1 = data_spinalysis['Geo Level 1'].value_counts()
df_final_geo2 = data_spinalysis['Geo Level 2'].value_counts()
df_final_geo3 = data_spinalysis['Geo Level 3'].value_counts()
df_final_geo4 = data_spinalysis['Geo Level 4'].value_counts()


#Floats indicating ratio of observations to error sources
#useful in logic related to the appropriateness of ANOVA
#Note that these are all based on the cleaned sample 'data_spinalysis'
subselect = data_spinalysis[['Item Avg', 'Job Title', 'Rater Unique ID',\
'Geo Level 1', 'Geo Level 2', 'Geo Level 3', 'Geo Level 4']]




'''FOR JOB TITLES'''

try:
    ratio_titles = float(len(data_spinalysis) / len(df_final_titles))
    if len(df_final_titles) > 1:
        arrDict = {}        
        #arrList = [[] for i in range(0,len(df_final_titles))]        
        for i in range(0,len(df_final_titles)):
            arrDict[i] = subselect[subselect['Job Title'] == \
            str(df_final_titles.index.values[i])]['Item Avg'].values
        
        #using the asterisk then method somehow allows the values to pass
        #Larry has seen this concept in the world of C++ programming
        f_titles, p_titles =  stats.f_oneway(*arrDict.values())
        crit_p_titles = 0.05/(len(df_final_titles)-1)
        
        '''JOB TITLES CONCLUSIONS'''
        if p_titles < crit_p_titles and ratio_titles > 10 :
            conclusion_titles = 'Significant difference here. Dig in deeper.'
        elif p_titles < crit_p_titles:
            conclusion_titles = 'There might be a difference, but confidence is low.'
        elif p_titles > crit_p_titles and ratio_titles > 10 :
            conclusion_titles = 'Evidence indicates no significant difference.'
        else:
            conclusion_titles = 'Little confidence in this, but no apparent differences observed.'
        
    else:
        f_titles, p_titles, crit_p_titles, conclusion_titles = (float(),float(),float(),'')

    
except:
    ratio_titles = "Empty titles"
    f_titles, p_titles, crit_p_titles, conclusion_titles = (float(),float(),float(),'')
    


'''FOR RATERS'''

try:
    ratio_raters = float(len(data_spinalysis) / len(df_final_raters))
    if len(df_final_raters) > 1:
        arrDict = {}        
        #arrList = [[] for i in range.(0,len(df_final_titles))]        
        for i in range(0,len(df_final_raters)):
            try:
                arrDict[i] = subselect[subselect['Rater Unique ID'] == \
                str(df_final_raters.index.values[i])]['Item Avg'].values
            except:
                arrDict[i] = subselect[subselect['Rater Unique ID'] == \
                df_final_raters.index.values[i]]['Item Avg'].values
        
        #using the asterisk then method somehow allows the values to pass
        #Larry has seen this concept in the world of C++ programming
        f_raters, p_raters =  stats.f_oneway(*arrDict.values())
        crit_p_raters = 0.05/(len(df_final_raters)-1)
        
        '''RATERS CONCLUSIONS'''
        if p_raters < crit_p_raters and ratio_raters > 10 :
            conclusion_raters = 'Significant difference here. Dig in deeper.'
        elif p_raters < crit_p_raters:
            conclusion_raters = 'There might be a difference, but confidence is low.'
        elif p_raters > crit_p_raters and ratio_raters > 10 :
            conclusion_raters = 'Evidence indicates no significant difference.'
        else:
            conclusion_raters = 'Little confidence in this, but no apparent differences observed.'
        
    else:
        f_raters, p_raters, crit_p_raters, conclusion_raters = (float(),float(),float(),'')
    
except:
    ratio_raters = "Empty raters"
    f_raters, p_raters, crit_p_raters, conclusion_raters = (float(),float(),float(),'')



'''FOR GEO 1'''

try:
    ratio_final_geo1 = float(len(data_spinalysis) / len(df_final_geo1))
    if len(df_final_geo1) > 1:
        arrDict = {}        
        #arrList = [[] for i in range(0,len(df_final_titles))]        
        for i in range(0,len(df_final_geo1)):
            try:
                arrDict[i] = subselect[subselect['Geo Level 1'] == \
                str(df_final_geo1.index.values[i])]['Item Avg'].values
            except:
                arrDict[i] = subselect[subselect['Geo Level 1'] == \
                df_final_geo1.index.values[i]]['Item Avg'].values
                   
        
        #using the asterisk then method somehow allows the values to pass
        #Larry has seen this concept in the world of C++ programming
        f_geo1, p_geo1 =  stats.f_oneway(*arrDict.values())
        crit_p_geo1 = 0.05/(len(df_final_geo1)-1)
        
        '''GEO1 CONCLUSIONS'''
        if p_geo1 < crit_p_geo1 and ratio_final_geo1 > 10 :
            conclusion_geo1 = 'Significant difference here. Dig in deeper.'
        elif p_geo1 < crit_p_geo1:
            conclusion_geo1 = 'There might be a difference, but confidence is low.'
        elif p_geo1 > crit_p_geo1 and ratio_final_geo1 > 10 :
            conclusion_geo1 = 'Evidence indicates no significant difference.'
        else:
            conclusion_geo1 = 'Little confidence in this, but no apparent differences observed.'
        
    else:
        f_geo1, p_geo1, crit_p_geo1, conclusion_geo1 = (float(),float(),float(),'')
    
except:
    ratio_final_geo1 = "Empty geo1"
    f_geo1, p_geo1, crit_p_geo1, conclusion_geo1 = (float(),float(),float(),'')



'''FOR GEO 2'''

try:
    ratio_final_geo2 = float(len(data_spinalysis) / len(df_final_geo2))
    if len(df_final_geo2) > 1:
        arrDict = {}        
        #arrList = [[] for i in range(0,len(df_final_titles))]        
        for i in range(0,len(df_final_geo2)):
            try:
                arrDict[i] = subselect[subselect['Geo Level 2'] == \
                str(df_final_geo2.index.values[i])]['Item Avg'].values
            except:
                arrDict[i] = subselect[subselect['Geo Level 2'] == \
                df_final_geo2.index.values[i]]['Item Avg'].values

        #using the asterisk then method somehow allows the values to pass
        #Larry has seen this concept in the world of C++ programming
        f_geo2, p_geo2 =  stats.f_oneway(*arrDict.values())
        crit_p_geo2 = 0.05/(len(df_final_geo2)-1)
        
        '''GEO2 CONCLUSIONS'''
        if p_geo2 < crit_p_geo2 and ratio_final_geo2 > 10 :
            conclusion_geo2 = 'Significant difference here. Dig in deeper.'
        elif p_geo2 < crit_p_geo2:
            conclusion_geo2 = 'There might be a difference, but confidence is low.'
        elif p_geo2 > crit_p_geo2 and ratio_final_geo2 > 10 :
            conclusion_geo2 = 'Evidence indicates no significant difference.'
        else:
            conclusion_geo2 = 'Little confidence in this, but no apparent differences observed.'
        
    else:
        f_geo2, p_geo2, crit_p_geo2, conclusion_geo2 = (float(),float(),float(),'')
    
except:
    ratio_final_geo2 = "Empty geo2"
    f_geo2, p_geo2, crit_p_geo2, conclusion_geo2 = (float(),float(),float(),'')



'''GEO LEVEL 3'''

try:
    ratio_final_geo3 = float(len(data_spinalysis) / len(df_final_geo3))
    if len(df_final_geo3) > 1:
        arrDict = {}        
        #arrList = [[] for i in range(0,len(df_final_titles))]        
        for i in range(0,len(df_final_geo3)):
            try:
                arrDict[i] = subselect[subselect['Geo Level 3'] == \
                str(df_final_geo3.index.values[i])]['Item Avg'].values
            except:
                arrDict[i] = subselect[subselect['Geo Level 3'] == \
                df_final_geo3.index.values[i]]['Item Avg'].values

        #using the asterisk then method somehow allows the values to pass
        #Larry has seen this concept in the world of C++ programming
        f_geo3, p_geo3 =  stats.f_oneway(*arrDict.values())
        crit_p_geo3 = 0.05/(len(df_final_geo3)-1)
        
        '''GEO3 CONCLUSIONS'''
        if p_geo3 < crit_p_geo3 and ratio_final_geo3 > 10 :
            conclusion_geo3 = 'Significant difference here. Dig in deeper.'
        elif p_geo3 < crit_p_geo3:
            conclusion_geo3 = 'There might be a difference, but confidence is low.'
        elif p_geo3 > crit_p_geo3 and ratio_final_geo3 > 10 :
            conclusion_geo3 = 'Evidence indicates no significant difference.'
        else:
            conclusion_geo3 = 'Little confidence in this, but no apparent differences observed.'
        
    else:
        f_geo3, p_geo3, crit_p_geo3, conclusion_geo3 = (float(),float(),float(),'')  
except:
    ratio_final_geo3 = "Empty geo3"
    f_geo3, p_geo3, crit_p_geo3, conclusion_geo3 = (float(),float(),float(),'')



'''GEO LEVEL 4'''

try:
    ratio_final_geo4 = float(len(data_spinalysis) / len(df_final_geo4))
    if len(df_final_geo4) > 1:
        arrDict = {}        
        #arrList = [[] for i in range(0,len(df_final_titles))]        
        for i in range(0,len(df_final_geo4)):
            try:
                arrDict[i] = subselect[subselect['Geo Level 4'] == \
                str(df_final_geo4.index.values[i])]['Item Avg'].values
            except:
                arrDict[i] = subselect[subselect['Geo Level 4'] == \
                df_final_geo4.index.values[i]]['Item Avg'].values

        #using the asterisk then method somehow allows the values to pass
        #Larry has seen this concept in the world of C++ programming
        f_geo4, p_geo4 =  stats.f_oneway(*arrDict.values())
        crit_p_geo4 = 0.05/(len(df_final_geo4)-1)
        
        '''GEO4 CONCLUSIONS'''
        if p_geo4 < crit_p_geo4 and ratio_final_geo4 > 10 :
            conclusion_geo4 = 'Significant difference here. Dig in deeper.'
        elif p_geo4 < crit_p_geo4:
            conclusion_geo4 = 'There might be a difference, but confidence is low.'
        elif p_geo4 > crit_p_geo4 and ratio_final_geo4 > 10 :
            conclusion_geo4 = 'Evidence indicates no significant difference.'
        else:
            conclusion_geo4 = 'Little confidence in this, but no apparent differences observed.'
        
    else:
        f_geo4, p_geo4, crit_p_geo4, conclusion_geo4 = (float(),float(),float(),'')
    
except:
    ratio_final_geo4 = "Empty geo4"
    f_geo4, p_geo4, crit_p_geo4, conclusion_geo4 = (float(),float(),float(),'')
  
    

#REPORTS TAB

#Computation of standardized Cronbach's alphs (scale reliability)
#requires k = number of items, and r_bar_lower which is the 
#average of all unique split-half correlations between items
#I compute this from the return of the corr function, which is full matrix

sum_corrs = corrs_items.sum().sum()
r_bar_lower = (sum_corrs-k) / (k*(k-1))
alpha_standard = k*r_bar_lower / (1+(k-1)*r_bar_lower)



df_report = DataFrame([['Scale Items', k], ['Sample size',len(data_spinalysis)],\
['r-bar', ("%.2f" % r_bar_lower)],\
['alpha_standard', ("%.2f" % alpha_standard)],\
['COVARIATES','(Ratio values over 10 \n can be analyzed with confidence)'], \
['Job titles','ratio = {}'.format(ratio_titles), \
("%.4f" % f_titles), ("%.4f" % p_titles), ("%.4f" % crit_p_titles), conclusion_titles], \
['Raters','ratio = {}'.format(ratio_raters), \
("%.4f" % f_raters), ("%.4f" % p_raters), ("%.4f" % crit_p_raters), conclusion_raters], \
['Geo Level 1','ratio = {}'.format(ratio_final_geo1), \
("%.4f" % f_geo1), ("%.4f" % p_geo1), ("%.4f" % crit_p_geo1), conclusion_geo1], \
['Geo Level 2','ratio = {}'.format(ratio_final_geo2), \
("%.4f" % f_geo2), ("%.4f" % p_geo2), ("%.4f" % crit_p_geo2), conclusion_geo2], \
['Geo Level 3','ratio = {}'.format(ratio_final_geo3), \
("%.4f" % f_geo3), ("%.4f" % p_geo3), ("%.4f" % crit_p_geo3), conclusion_geo3], \
['Geo Level 4','ratio = {}'.format(ratio_final_geo4), \
("%.4f" % f_geo4), ("%.4f" % p_geo4), ("%.4f" % crit_p_geo4), conclusion_geo4]], \
columns=['Label','Values','ANOVA F_test', 'p_value', 'p_cutoff', 'Spinsight(c)'])


'''
Printed out to a csv format. File name adjusted manually. NA values
specified otherwise will be whitespace
'''

outfile = "Spin_Results"

#in_name was parsed when Tkinter first brings in the data files common name
outfilename = outfile+"("+in_name+")"


#custom tab labels
item_summary_tab = "Item Summary (k = {})".format(k)
corrs_items_tab = "Correls Items (n = {})".format(len(data_spinalysis))
titles_tab = "Job Titles (uniq = {})".format(len(df_final_titles))
raters_tab = "Raters (uniq = {})".format(len(df_final_raters))
geo1_tab = "Geo 1 (uniq = {})".format(len(df_final_geo1))
geo2_tab = "Geo 2 (uniq = {})".format(len(df_final_geo2))
geo3_tab = "Geo 3 (uniq = {})".format(len(df_final_geo3))
geo4_tab = "Geo 4 (uniq = {})".format(len(df_final_geo4))


try:
    #this needed to be in .xlsx to align with the default excel writer engine on this computer
    xl_out = outfilename+".xlsx"
    spin_master.to_excel(xl_out,sheet_name='Master',index=False,merge_cells=False,na_rep='')
    
    with ExcelWriter(xl_out) as writer:
        spin_master.to_excel(writer,sheet_name='Master',index=False,merge_cells=False,na_rep='')
        #summary.to_excel(writer,sheet_name='Summary Stats',index=True,merge_cells=False,na_rep='blank')
        df_report.to_excel(writer,sheet_name='Scale Report',index=False)        
        item_summary.to_excel(writer,sheet_name=item_summary_tab,index=True,merge_cells=False,na_rep='blank')
        corrs_items.to_excel(writer,sheet_name=corrs_items_tab)
        if len(df_final_titles) > 0:
            df_final_titles.to_frame(name='Titles Count').to_excel(writer,sheet_name=titles_tab)
        else:
            next
        if len(df_final_raters) > 0:
            df_final_raters.to_frame(name='Raters Count').to_excel(writer,sheet_name=raters_tab)
        else:
            next
        if len(df_final_geo1) > 0:
            df_final_geo1.to_frame(name='Geo1 Count').to_excel(writer,sheet_name=geo1_tab)
        else:
            next
        if len(df_final_geo2) > 0:
            df_final_geo2.to_frame(name='Geo2 Count').to_excel(writer,sheet_name=geo2_tab)
        else:
            next 
        if len(df_final_geo3) > 0:
            df_final_geo3.to_frame(name='Geo3 Count').to_excel(writer,sheet_name=geo3_tab)
        else:
            next       
        if len(df_final_geo4) > 0:
            df_final_geo4.to_frame(name='Geo4 Count').to_excel(writer,sheet_name=geo4_tab)
        else:
            next  
except:
    spin_master.to_csv(outfilename+".csv",na_rep='',index=False)

#outfile = "RaterError_Results.pd1.csv"
#data.to_csv(outfile,na_rep='blank')

print "\n\n\n>>>>>\n>>>>>"
print "\nAll done.\n"
print "Look in your working directory for a file named '{}.xlsx'".format(outfilename)
print "\nOpen this file and use the following columns from the "
print "file to conduct your matching session:"
print "'Ratee First Name, Ratee Last Name, Email, Last 4 SSN, Random ID'"
print "Email and SSN may not always be available."

print "\nWhen you export your matched incumbents,"
print "use the 'Profile Creation Match' button."

print "\nAfter downloading your matched file, "
print "open the file and 'Save As' a .csv file and the close the file."
print "It might look like the file is already in .csv format,"
print "but the file encoding options may not be correct."
print "This is why you need to manually 'Save As' on your own machine."
print "Using the 'Save as type' dropdown box \nselect 'CSV (Comma delimited) (*.csv)'"

print "\nThis second .csv file should be in the format required by Twister."

print "\nIt is recommended to save the file in the same location you saved\
\n the performance evaluation, using a name that you can remember."

Proceed3 = raw_input("\nIf you saved your profile creation match\n\
as a csv into a file and then closed it, press Enter.")

print "\n>>>>>\n>>>>>"

print "\nNow, use the file explorer window to access your match file\
\n that you saved as a .csv."
print "The file explorer may be behind this window so \
you may have to minimize this program."
print "\nIf your program is crashing after you pick your file, "
print "this may be because the file is still open"
print "or it was not saved in the correct format."


'''GETTING USER INPUT'''

Tk().withdraw()
file_name_match_full = askopenfilename()

match_in_name_list = string.split(file_name_match_full,sep='/')
file_name_match = str(match_in_name_list[-1].rstrip('.csv'))

print "\n>>>>>\n>>>>>"
print "\nThanks!"
print "So we are working with the matched file '{}.csv'.".format(file_name_match)


try:
    data_match = pd.read_csv(file_name_match_full, na_values=['NULL'], encoding = 'utf-8')
except:
    data_match = pd.read_csv(file_name_match_full, na_values=['NULL'], encoding = 'cp1252')

#I took out the set Random ID as index command here so that we can copy it and initialize the copy as the index
#index_col=['Random ID'], \

#The alternative would have been to reset it and then copy then initialize the copy
#data_match = data_match.reset_index()

data_match['Random ID as Index'] = data_match['Random ID']
data_match = data_match.set_index('Random ID as Index')

#Create Quality Control(QC) by first reformatting 3 of the QC metrics into float types

data_match['FALSIFICATION'] = data_match['FALSIFICATION'].replace('%','',regex=True).astype('float')/100
data_match['CONSISTENCY'] = data_match['CONSISTENCY'].replace('%','',regex=True).astype('float')/100
data_match['CLICK_THROUGH'] = data_match['CLICK_THROUGH'].replace('%','',regex=True).astype('float')/100


#Let's Rename Email to Match Email so that it is unique
#data_match['Email'] = data_match['Match Email']
data_match =  data_match.rename(columns={'EMAIL': 'Match Email'})

Ptim_Dummy = lambda x: x != 0 and x <=11.12
Pim_Dummy = lambda x: x <= 12.97
Falsification_Dummy = lambda x: x >= 0.75
Con_Dummy = lambda x: x <= 0.67
Clik_Dummy = lambda x: x >= 0.6641
Seq_Dummy = lambda x: x >= 2
Lcs_Dummy = lambda x: x >= 17
Quality_column = lambda x: x != 0 


#Apply the Lambdas that were created
data_match['QC Ptim'] = Series(data_match['PPI_TOTAL_ITEM_MINUTES'].apply(Ptim_Dummy),dtype=np.int32)
data_match['QC Pim'] = Series(data_match['PPI_MINUTES'].apply(Pim_Dummy),dtype=np.int32)
data_match['QC Falsification'] = Series(data_match['FALSIFICATION'].apply(Falsification_Dummy),dtype=np.int32)
data_match['QC Con'] = Series(data_match['CONSISTENCY'].apply(Con_Dummy),dtype=np.int32)
data_match['QC Clik'] = Series(data_match['CLICK_THROUGH'].apply(Clik_Dummy),dtype=np.int32)
data_match['QC Seq'] = Series(data_match['SEQUENCE'].apply(Seq_Dummy),dtype=np.int32)
data_match['QC Lcs'] = Series(data_match['LCS_LENGTH'].apply(Lcs_Dummy),dtype=np.int32)

Quality_control_Dummy_headers = ['QC Ptim', 'QC Pim', 'QC Falsification', 'QC Con', 'QC Clik', 'QC Seq', 'QC Lcs']
Quality_control_headers = ['PPI_TOTAL_ITEM_MINUTES', 'PPI_MINUTES', 'FALSIFICATION', 'CONSISTENCY', 'CLICK_THROUGH', 'SEQUENCE', 'LCS_LENGTH']
data_match['Removal - Quality Control = 1'] = Series(data_match.ix[:,Quality_control_Dummy_headers].sum(axis = 1).apply(Quality_column),dtype=np.int32)

#Create new data frame excluding Quality Control removals
Clean_data_match = data_match[data_match['Removal - Quality Control = 1']  != 1]

#Determine mean and std of each dimension

Headers_PPI2 = ['ACCEPTANCE_OF_AUTHORITY', 'AMBITION', 'ANALYTICAL', 'ASSERTIVENESS',\
 'ATTENTION_TO_DETAILS', 'BUSINESS_ATTITUDE', 'CHANGE_ORIENTATION', \
 'COMPETITIVE_FIERCENESS', 'CONFIDENCE', 'CONSCIENTIOUSNESS', 'COOPERATIVENESS', \
 'CREATIVITY', 'DISCIPLINE', 'EMOTIONAL_CONSISTENCY', 'ENERGY', 'FLEXIBILITY', \
 'INSIGHT_INTO_OTHERS', 'JOB_ATMOSPHERE', 'LEADERSHIP_IMPACT', 'MENTAL_FLEXIBILITY', \
 'NEED_FOR_RECOGNITION', 'NUMERICAL_REASONING', 'OBJECTIVITY', 'OPTIMISM', \
 'ORGANIZATIONAL_STRUCTURE', 'ORGANIZATIONAL_TENDENCY', 'PACE', 'PEOPLE_ORIENTATION', \
 'PRACTICAL', 'REALISTIC_THINKING', 'REFLECTIVE', 'RISK_TAKER', 'SELF_RELIANCE', \
 'SOCIABILITY', 'SOCIAL_CONTACT', 'STRESS_TOLERANCE', 'TEAM_ORIENTATION', \
 'TOUGH_MINDEDNESS', 'VERBAL_REASONING']

#Creating Z Scores for each dimension
Dimension_descriptives1 = Clean_data_match.ix[:,Headers_PPI2].describe()
Dimension_descriptives2 = Clean_data_match.ix[:,Headers_PPI2]
#Dimension_Upper_Lower = Dimension_descriptives.ix[1:3,:]

for col in Dimension_descriptives2:
    col_zscore = col + '_zscore'
    Dimension_descriptives2[col_zscore] = (Dimension_descriptives2[col] - Dimension_descriptives2[col].mean())/Dimension_descriptives2[col].std(ddof=0)

'''Population parameters for PPI2+PMA are read in from 
a prepared .csv file saved on a folder in my desktop. Only PPI2+PMA at this point.

Below, the try statement reads in the .csv.
In case this file is dislocated, the except statement uses the hardcoded values.
I like the read_csv better because it gives an independent way to manage
the parameter estimates. But the values Are essential to calculations, so 
a backup hardcode may be necessary.
'''


#try:
#    ppi2_pop_params = pd.read_csv('C:\Users\dbrown3\Desktop\PPI2_PMA_PopMeanStd.csv')
#except:
ppi2_pop_params_dict = {'Dimension Name': ['ACCEPTANCE_OF_AUTHORITY', 'AMBITION', 'ANALYTICAL', 'ASSERTIVENESS',\
 'ATTENTION_TO_DETAILS', 'BUSINESS_ATTITUDE', 'CHANGE_ORIENTATION', \
 'COMPETITIVE_FIERCENESS', 'CONFIDENCE', 'CONSCIENTIOUSNESS', 'COOPERATIVENESS', \
 'CREATIVITY', 'DISCIPLINE', 'EMOTIONAL_CONSISTENCY', 'ENERGY', 'FLEXIBILITY', \
 'INSIGHT_INTO_OTHERS', 'JOB_ATMOSPHERE', 'LEADERSHIP_IMPACT', 'MENTAL_FLEXIBILITY', \
 'NEED_FOR_RECOGNITION', 'NUMERICAL_REASONING', 'OBJECTIVITY', 'OPTIMISM', \
 'ORGANIZATIONAL_STRUCTURE', 'ORGANIZATIONAL_TENDENCY', 'PACE', 'PEOPLE_ORIENTATION', \
 'PRACTICAL', 'REALISTIC_THINKING', 'REFLECTIVE', 'RISK_TAKER', 'SELF_RELIANCE', \
 'SOCIABILITY', 'SOCIAL_CONTACT', 'STRESS_TOLERANCE', 'TEAM_ORIENTATION', \
 'TOUGH_MINDEDNESS', 'VERBAL_REASONING'],\
'Mean.p': [63.96,54.13,53.8,55.15,60.02,58.44,48.96,51.1,51.09,61.39,57.06,53.46,\
60.54,60.52,57.04,54.62,60.8,59.35,52.18,52.47,50.5,51.25,56.82,57.64,47.57,63.99,55.06,\
62.33,60.12,45.89,52.86,52.02,43.68,61.53,53.47,54.77,53.84,48.59,53.69],\
'Std.p': [18.83,16.87,12.33,13.47,15.6,15.84,13.26,13.82,11.92,16.06,14.3,16.66,\
17.45,18.34,15.48,13,16.09,13.88,12.52,13.76,13.84,15.53,14.87,14.49,11.93,16.69,14.88,\
14.55,14.05,12.86,12.79,14.28,12.39,15.51,13.26,13.21,12.36,13.12,15.09]}    

ppi2_pop_params = pd.DataFrame(ppi2_pop_params_dict, columns=['Dimension Name',\
    'Mean.p','Std.p'])


#The full version would need to filter out non-relevant assessments in order
#for the standardization values to be accurate.

#if matched_first starts with "Zzt", that is a test account


#Because the length of the parameters dataframe is known a priori
#(39 dimensions, each with a mean and std),
#this list can be pre-written to loop through and standardize the values in 
#my master file called 'data'

for i in range(0,39):
    z_name = '_'.join(('z',ppi2_pop_params.ix[i,'Dimension Name']))
    Dimension_descriptives2[z_name] = (Dimension_descriptives2[ppi2_pop_params.ix[i,'Dimension Name']] \
- ppi2_pop_params[ppi2_pop_params['Dimension Name']==ppi2_pop_params.ix[i,'Dimension Name']].ix[:,'Mean.p'].values)\
/ ppi2_pop_params[ppi2_pop_params['Dimension Name']==ppi2_pop_params.ix[i,'Dimension Name']].ix[:,'Std.p'].values


Sample_Zscored_Headers_PPI2 = ['ACCEPTANCE_OF_AUTHORITY_zscore', 'AMBITION_zscore', \
'ANALYTICAL_zscore', 'ASSERTIVENESS_zscore','ATTENTION_TO_DETAILS_zscore', \
'BUSINESS_ATTITUDE_zscore', 'CHANGE_ORIENTATION_zscore', 'COMPETITIVE_FIERCENESS_zscore', \
'CONFIDENCE_zscore', 'CONSCIENTIOUSNESS_zscore', 'COOPERATIVENESS_zscore', \
'CREATIVITY_zscore', 'DISCIPLINE_zscore', 'EMOTIONAL_CONSISTENCY_zscore', 'ENERGY_zscore', \
'FLEXIBILITY_zscore', 'INSIGHT_INTO_OTHERS_zscore', 'JOB_ATMOSPHERE_zscore', 'LEADERSHIP_IMPACT_zscore', \
'MENTAL_FLEXIBILITY_zscore', 'NEED_FOR_RECOGNITION_zscore', 'NUMERICAL_REASONING_zscore', \
'OBJECTIVITY_zscore', 'OPTIMISM_zscore', 'ORGANIZATIONAL_STRUCTURE_zscore', \
'ORGANIZATIONAL_TENDENCY_zscore', 'PACE_zscore', 'PEOPLE_ORIENTATION_zscore', \
'PRACTICAL_zscore', 'REALISTIC_THINKING_zscore', 'REFLECTIVE_zscore', 'RISK_TAKER_zscore', \
'SELF_RELIANCE_zscore', 'SOCIABILITY_zscore', 'SOCIAL_CONTACT_zscore', 'STRESS_TOLERANCE_zscore', \
'TEAM_ORIENTATION_zscore', 'TOUGH_MINDEDNESS_zscore', 'VERBAL_REASONING_zscore']


Population_dummy_coded_Zscored_Headers_PPI2 = ['z_ACCEPTANCE_OF_AUTHORITY', 'z_AMBITION',\
'z_ANALYTICAL', 'z_ASSERTIVENESS', 'z_ATTENTION_TO_DETAILS', 'z_BUSINESS_ATTITUDE', 'z_CHANGE_ORIENTATION','z_COMPETITIVE_FIERCENESS',\
'z_CONFIDENCE', 'z_CONSCIENTIOUSNESS', 'z_COOPERATIVENESS', 'z_CREATIVITY', 'z_DISCIPLINE', \
 'z_EMOTIONAL_CONSISTENCY', 'z_ENERGY', 'z_FLEXIBILITY', 'z_INSIGHT_INTO_OTHERS', 'z_JOB_ATMOSPHERE', \
 'z_LEADERSHIP_IMPACT', 'z_MENTAL_FLEXIBILITY', 'z_NEED_FOR_RECOGNITION', 'z_NUMERICAL_REASONING', 'z_OBJECTIVITY', \
 'z_OPTIMISM', 'z_ORGANIZATIONAL_STRUCTURE', 'z_ORGANIZATIONAL_TENDENCY', 'z_PACE', 'z_PEOPLE_ORIENTATION', \
 'z_PRACTICAL', 'z_REALISTIC_THINKING', 'z_REFLECTIVE', 'z_RISK_TAKER', 'z_SELF_RELIANCE', 'z_SOCIABILITY', \
 'z_SOCIAL_CONTACT', 'z_STRESS_TOLERANCE', 'z_TEAM_ORIENTATION', 'z_TOUGH_MINDEDNESS', 'z_VERBAL_REASONING']

#Trying to find a code creates dummy code column for assessment outliers after evaluating whether there are 5 or more assessment outliers across all dimensions for the people in this data frame who have passed quality control.
Sample_Assessment_outlier_dummy_code_columns = Dimension_descriptives2.ix[:,Sample_Zscored_Headers_PPI2]
Population_Assessment_outlier_dummy_code_columns = Dimension_descriptives2.ix[:,Population_dummy_coded_Zscored_Headers_PPI2]


Assessment_Outlier_Column = lambda x: np.abs(x) > 3

for cols in Sample_Assessment_outlier_dummy_code_columns:
    cols_dummies = cols + '_dummy'
    Sample_Assessment_outlier_dummy_code_columns[cols_dummies] = Series(Sample_Assessment_outlier_dummy_code_columns[cols].apply(Assessment_Outlier_Column),dtype=np.int32)

for cols in Population_Assessment_outlier_dummy_code_columns:
    pop_cols_dummies = cols + '_dummy'
    Population_Assessment_outlier_dummy_code_columns[pop_cols_dummies] = Series(Population_Assessment_outlier_dummy_code_columns[cols].apply(Assessment_Outlier_Column),dtype=np.int32)

Sample_Dummy_coded_Zscored_Headers_PPI2 = ['ACCEPTANCE_OF_AUTHORITY_zscore_dummy', 'AMBITION_zscore_dummy', \
'ANALYTICAL_zscore_dummy', 'ASSERTIVENESS_zscore_dummy','ATTENTION_TO_DETAILS_zscore_dummy', \
'BUSINESS_ATTITUDE_zscore_dummy', 'CHANGE_ORIENTATION_zscore_dummy', 'COMPETITIVE_FIERCENESS_zscore_dummy', \
'CONFIDENCE_zscore_dummy', 'CONSCIENTIOUSNESS_zscore_dummy', 'COOPERATIVENESS_zscore_dummy', \
'CREATIVITY_zscore_dummy', 'DISCIPLINE_zscore_dummy', 'EMOTIONAL_CONSISTENCY_zscore_dummy', 'ENERGY_zscore_dummy', \
'FLEXIBILITY_zscore_dummy', 'INSIGHT_INTO_OTHERS_zscore_dummy', 'JOB_ATMOSPHERE_zscore_dummy', 'LEADERSHIP_IMPACT_zscore_dummy', \
'MENTAL_FLEXIBILITY_zscore_dummy', 'NEED_FOR_RECOGNITION_zscore_dummy', 'NUMERICAL_REASONING_zscore_dummy', \
'OBJECTIVITY_zscore_dummy', 'OPTIMISM_zscore_dummy', 'ORGANIZATIONAL_STRUCTURE_zscore_dummy', \
'ORGANIZATIONAL_TENDENCY_zscore_dummy', 'PACE_zscore_dummy', 'PEOPLE_ORIENTATION_zscore_dummy', \
'PRACTICAL_zscore_dummy', 'REALISTIC_THINKING_zscore_dummy', 'REFLECTIVE_zscore_dummy', 'RISK_TAKER_zscore_dummy', \
'SELF_RELIANCE_zscore_dummy', 'SOCIABILITY_zscore_dummy', 'SOCIAL_CONTACT_zscore_dummy', 'STRESS_TOLERANCE_zscore_dummy', \
'TEAM_ORIENTATION_zscore_dummy', 'TOUGH_MINDEDNESS_zscore_dummy', 'VERBAL_REASONING_zscore_dummy']

Population_Dummy_coded_Zscored_Headers_PPI2 = ['z_ACCEPTANCE_OF_AUTHORITY_dummy', 'z_AMBITION_dummy', 'z_ANALYTICAL_dummy',\
 'z_ASSERTIVENESS_dummy', 'z_ATTENTION_TO_DETAILS_dummy', 'z_BUSINESS_ATTITUDE_dummy', 'z_CHANGE_ORIENTATION_dummy', \
 'z_COMPETITIVE_FIERCENESS_dummy', 'z_CONFIDENCE_dummy', 'z_CONSCIENTIOUSNESS_dummy', 'z_COOPERATIVENESS_dummy', \
 'z_CREATIVITY_dummy', 'z_DISCIPLINE_dummy', 'z_EMOTIONAL_CONSISTENCY_dummy', 'z_ENERGY_dummy', 'z_FLEXIBILITY_dummy', \
 'z_INSIGHT_INTO_OTHERS_dummy', 'z_JOB_ATMOSPHERE_dummy', 'z_LEADERSHIP_IMPACT_dummy', 'z_MENTAL_FLEXIBILITY_dummy', \
 'z_NEED_FOR_RECOGNITION_dummy', 'z_NUMERICAL_REASONING_dummy', 'z_OBJECTIVITY_dummy', 'z_OPTIMISM_dummy', \
 'z_ORGANIZATIONAL_STRUCTURE_dummy', 'z_ORGANIZATIONAL_TENDENCY_dummy', 'z_PACE_dummy', 'z_PEOPLE_ORIENTATION_dummy', \
 'z_PRACTICAL_dummy', 'z_REALISTIC_THINKING_dummy', 'z_REFLECTIVE_dummy', 'z_RISK_TAKER_dummy', 'z_SELF_RELIANCE_dummy', \
 'z_SOCIABILITY_dummy', 'z_SOCIAL_CONTACT_dummy', 'z_STRESS_TOLERANCE_dummy', 'z_TEAM_ORIENTATION_dummy', 'z_TOUGH_MINDEDNESS_dummy',\
 'z_VERBAL_REASONING_dummy']

Assessment_Dummy_column = lambda x: x > 4

Sample_Assessment_outlier_dummy_code_columns['Sum of Sample Dependent Dimension Outliers'] = Series(Sample_Assessment_outlier_dummy_code_columns.ix[:,Sample_Dummy_coded_Zscored_Headers_PPI2].sum(axis = 1))
Sample_Assessment_outlier_dummy_code_columns['Removal - Sample Dependent Assessment Outlier = 1'] = Series(Sample_Assessment_outlier_dummy_code_columns.ix[:,Sample_Dummy_coded_Zscored_Headers_PPI2].sum(axis = 1).apply(Assessment_Dummy_column),dtype=np.int32)

Population_Assessment_outlier_dummy_code_columns['Sum of Population Dependent Dimension Outliers'] = Series(Population_Assessment_outlier_dummy_code_columns.ix[:,Population_Dummy_coded_Zscored_Headers_PPI2].sum(axis = 1))
Population_Assessment_outlier_dummy_code_columns['Removal - Population Dependent Assessment Outlier = 1'] = Series(Population_Assessment_outlier_dummy_code_columns.ix[:,Population_Dummy_coded_Zscored_Headers_PPI2].sum(axis = 1).apply(Assessment_Dummy_column),dtype=np.int32)

data_match = data_match.join(Sample_Assessment_outlier_dummy_code_columns[['Sum of Sample Dependent Dimension Outliers', 'Removal - Sample Dependent Assessment Outlier = 1']], rsuffix= '_y')
data_match = data_match.join(Population_Assessment_outlier_dummy_code_columns[['Sum of Population Dependent Dimension Outliers', 'Removal - Population Dependent Assessment Outlier = 1']], rsuffix = '_y')




#Export to file to merge with Data file into a Master file
#******************************************************************************

full_data = spin_master.join(data_match, rsuffix='_y')
full_data['Performance Type'] = ''
full_data['Set Type'] = 'Training'
full_data['Overall Performance Rating'] = full_data['Item Avg']
full_data = full_data.rename(columns={'FIRST_NAME': 'Match First', \
'LAST_NAME': 'Match Last', 'LAST_FOUR_SSN': 'Match SSN'})


master_tab_headers = ['PA ID', 'Match First', 'Match Last', 'Match Email', 'Match SSN', 'Performance Type', \
'Set Type', tenure_label, 'Overall Performance Rating']

data5 = full_data.ix[:,master_tab_headers]


end_item_location = spin_master.columns.get_loc('Item Avg')
data6 = data5.join(spin_master.ix[:,19:(end_item_location+6)], rsuffix= '_y')
#data 6 is to append from item 1 through any supplemental items.
#this is why we had to work backwards from the location of Item Avg
#Item Avg should not be displayed here though. its just a reference point


master_tab_headers2 = ['Job Title', 'Hire Date', 'Random ID', 'Ratee First Name', 'Ratee Last Name',\
'Ratee Unique ID', 'Survey Group', 'Geo Level 1', 'Geo Level 2', 'Geo Level 3', 'Geo Level 4', 'Rater First Name',\
'Rater Last Name', 'Rater Unique ID', 'Ratee Status']

data7 = data6.join(full_data.ix[:,master_tab_headers2], rsuffix= '_y')
data8 = data7.join(full_data.ix[:,Headers_PPI2], rsuffix= '_y')
data9 = data8.join(full_data.ix[:,Quality_control_headers], rsuffix= '_y')

data9['Removal - Not Matched = 1'] = Series(data9['PA ID'].isnull(),dtype=np.int32)

#this can be built to hold all the removal headers
master_tab_dummyheaders =['Removal - Quality Control = 1', 'Removal - Missing tenure data = 1', tenure_rem_label,\
'Removal - Completely missing performance data = 1', 'Removal - Partially missing performance data = 1', 'Removal - Lack of Rater Confidence = 1',\
 'Removal - Sample Dependent Assessment Outlier = 1', 'Removal - Population Dependent Assessment Outlier = 1']

data10 = data9.join(full_data.ix[:,master_tab_dummyheaders], rsuffix= '_y')
#Below we fill any misising cells with zeros
data10[master_tab_dummyheaders] = data10[master_tab_dummyheaders].fillna(0)

data11 = data10
final_removal_headers = ['Removal - Not Matched = 1'] + master_tab_dummyheaders
data11['Final Sample Code'] = data11.ix[:,final_removal_headers].sum(axis= 1)
Final_Data = data11[data11['Final Sample Code']==0]
Final_Data = Final_Data.drop('Final Sample Code', 1)
data10= data10.drop('Final Sample Code', 1)



'''MODELLING SECTION'''

from sklearn import linear_model


model_data = Final_Data

for i in range(0,39):
    z_name = '_'.join(('z',ppi2_pop_params.ix[i,'Dimension Name']))
    model_data[z_name] = (model_data[ppi2_pop_params.ix[i,'Dimension Name']] \
    - ppi2_pop_params[ppi2_pop_params['Dimension Name']==ppi2_pop_params.ix[i,'Dimension Name']].ix[:,'Mean.p'].values)\
    / ppi2_pop_params[ppi2_pop_params['Dimension Name']==ppi2_pop_params.ix[i,'Dimension Name']].ix[:,'Std.p'].values


model_headers = ['Overall Performance Rating'] + Headers_PPI2


z_ppi2_headers = ['z_ACCEPTANCE_OF_AUTHORITY', 'z_AMBITION',\
'z_ANALYTICAL', 'z_ASSERTIVENESS', 'z_ATTENTION_TO_DETAILS', 'z_BUSINESS_ATTITUDE', 'z_CHANGE_ORIENTATION','z_COMPETITIVE_FIERCENESS',\
'z_CONFIDENCE', 'z_CONSCIENTIOUSNESS', 'z_COOPERATIVENESS', 'z_CREATIVITY', 'z_DISCIPLINE', \
 'z_EMOTIONAL_CONSISTENCY', 'z_ENERGY', 'z_FLEXIBILITY', 'z_INSIGHT_INTO_OTHERS', 'z_JOB_ATMOSPHERE', \
 'z_LEADERSHIP_IMPACT', 'z_MENTAL_FLEXIBILITY', 'z_NEED_FOR_RECOGNITION', 'z_NUMERICAL_REASONING', 'z_OBJECTIVITY', \
 'z_OPTIMISM', 'z_ORGANIZATIONAL_STRUCTURE', 'z_ORGANIZATIONAL_TENDENCY', 'z_PACE', 'z_PEOPLE_ORIENTATION', \
 'z_PRACTICAL', 'z_REALISTIC_THINKING', 'z_REFLECTIVE', 'z_RISK_TAKER', 'z_SELF_RELIANCE', 'z_SOCIABILITY', \
 'z_SOCIAL_CONTACT', 'z_STRESS_TOLERANCE', 'z_TEAM_ORIENTATION', 'z_TOUGH_MINDEDNESS', 'z_VERBAL_REASONING']

z_model_headers = ['Overall Performance Rating'] + z_ppi2_headers

#array of predictor values
X = model_data[model_headers[1:]].values
#z_X = model_data[z_model_headers[1:]].values

#1d array of outcome values
Y = model_data[model_headers[0]].values



#instantiate the regression object
clf = linear_model.LinearRegression()

#the fit method regresses predictors array X on outcome array Y
clf.fit(X,Y)

#now the object has stored the coefficients
clf.coef_



#Add MVO optional add-on, it's cooks d and 4/n and leverage by case

'''END MODELLING SECTION'''


#******************************************************************************
outfile2 = "Final_Results"



in_name_list2 = string.split(filename,sep='/')
in_name2 = str(in_name_list2[-1].rstrip('.csv'))

outfilename2 = outfile2+"("+in_name2+")"


try:
    #this needed to be in .xlsx to align with the default excel writer engine on this computer
    xl_out2 = outfilename2+".xlsx"
    spin_master.to_excel(xl_out2,sheet_name='Spin Master',index=True,merge_cells=False,na_rep='blank')
    
    with ExcelWriter(xl_out2) as writer:
        #spin_master.to_excel(writer,sheet_name='Spin File',index=True,merge_cells=False,na_rep='blank')
        #Population_Assessment_outlier_dummy_code_columns.to_excel(writer,sheet_name='Population Values',index=True,merge_cells=False,na_rep='blank')
        #Assessment_outlier_dummy_code_columns.to_excel(writer,sheet_name='Dimensions',index=True,merge_cells=False,na_rep='blank')
        #full_data.to_excel(writer,sheet_name='Twister Master',index=True,merge_cells=False,na_rep='blank')
        #data.ix[:,18:(18+k)].to_excel(writer,sheet_name='Items',index=True,merge_cells=False,na_rep='blank')
        #data8.to_excel(writer, sheet_name='Master', index=False, merge_cells=False)
        data10.to_excel(writer, sheet_name='Master File', index=False, merge_cells=False)
        Final_Data.to_excel(writer,sheet_name='Final Data',index=False,merge_cells=False,na_rep='blank')

except:
    data10.to_csv(outfilename2+".csv",na_rep='blank')

print "\n\n>>>>>\n>>>>>"
print "\nAll done!\n"
print "Look in your working directory for a file named '{}.xlsx'".format(outfilename2)



print "\n\n>>>>>\n>>>>>"
print "\nSay, Did you want to try the new CSO add-on?\
\nYour client-sign off email will be written for you in a .docx file\
\nassuming no other analyses or removals are needed."

while True:
    try:
        userq_cso_try = raw_input("\nType Y or N: ").upper()
        if userq_cso_try == 'Y' or userq_cso_try == 'N':
            break
        else:
            print "What was that?"
    except:
        print "Please type Y or N:"
            
if userq_cso_try == 'Y':
    
    from docx import Document
    from docx.shared import Inches
   
    document = Document()
        
    bs_name = raw_input("\nWho is the sign-off to? (ex: Patricia): ")
    dear_bs = "Hi {},".format(bs_name)
    paragraph = document.add_paragraph(dear_bs)
    
     
    profile_name = raw_input("\nWhat is the proper name of the profile? \
    \n(ex: Justice League 7 Nov15 - Principal Architect 1.00): ")   
    cso_opener = "Here is the client sign-off information for the {} profile --".format(profile_name)
    paragraph = document.add_paragraph(cso_opener)
    
    
    start_sample = "Initial sample of {} incumbents".format(len(data))
    paragraph = document.add_paragraph(start_sample)    
    print "\nGreat!\
    \nWe are looking at a starting sample of {} incumbents".format(len(data))
       
              
    #just re-ordered from final_removal_headers   
    cso_removal_headers = ['Removal - Not Matched = 1',\
     'Removal - Completely missing performance data = 1',\
     'Removal - Partially missing performance data = 1',\
     'Removal - Missing tenure data = 1',\
     tenure_rem_label,\
     'Removal - Lack of Rater Confidence = 1',\
     'Removal - Quality Control = 1',\
     'Removal - Sample Dependent Assessment Outlier = 1',\
     'Removal - Population Dependent Assessment Outlier = 1']
   
    cso_removal_bullet_tags =  [' with missing assessment data',\
     ' with completely missing performance data',\
     ' with partially missing performance data',\
     ' with missing tenure data',\
     ' with insufficient tenure',\
     ' with poor rater confidence',\
     ' with QC issues in the assessment',\
     ' assessment outliers (sample)',\
     ' assessment outliers (population)']  
   
   
    remdata = data10[cso_removal_headers] 
    
    cso_actual_bullets = []
    
    remdata_temp = remdata
    
    for i in range(0,len(cso_removal_headers)):
        if remdata_temp[cso_removal_headers[i]].sum() == 0:
            next
        else:          
           cso_actual_bullets = \
           cso_actual_bullets + ['{} {}'.format(int(remdata_temp[cso_removal_headers[i]].sum()), cso_removal_bullet_tags[i])]
           
           remdata_temp = remdata_temp[remdata_temp[cso_removal_headers[0:i+1]].sum(axis=1)==0]
           
           
    print "Here are the removals as best I can tell:"
    for point in cso_actual_bullets:
        
        document.add_paragraph(point, style='List Bullet')
        print point
        #document.add_paragraph(point, style='ListBullet')
    
    
    final_sample = "Final sample size: {}".format(len(Final_Data))        
    print "\nThe final sample contains {} incumbents".format(len(Final_Data))
    fs_paragraph = document.add_paragraph()
    fs_paragraph.add_run(final_sample).bold = True

    document.add_heading("Performance Metric: Overall Performance Rating", level=1)
    
    metric_avg = "Average score: {}".format(("%.2f" % Final_Data['Item Avg'].mean()))
    document.add_paragraph(metric_avg)
    
    metric_min = "Min: {}".format(("%.2f" % Final_Data['Item Avg'].min()))
    document.add_paragraph(metric_min, style='List Bullet')
        
    metric_max = "Max: {}".format(("%.2f" % Final_Data['Item Avg'].max()))
    document.add_paragraph(metric_max, style='List Bullet')
    
    document.add_paragraph("Please let me know if I can provide any further information.")
    document.add_paragraph("Thanks,")
    document.add_paragraph("Twister")

    document.add_page_break()  
    document.add_heading('Spinsight(c)', level=1)
    spinsight_intro = "\nSpinsight is dedicated to helping you understand your performance data. Performance data analyses are based on the sample with only performance-side removals considered. This sample is typically larger than the final sample. Performance analyses are performed agnostic of the predictor data in order to establish more generalizable statistical estimates of performance data characteristics.\n"
    document.add_paragraph(spinsight_intro)    
    
    document.add_heading('Performance Instrument Summary', level=2)
    
    pis_body = "\nScale Items:    {}\
\nSample Size:    {}\
\nr-bar (average inter-item correlation):    {}\
\nCronbach's alpha (standardized):    {}\n".format(\
df_report.Values[0],\
df_report.Values[1],\
df_report.Values[2],\
df_report.Values[3])    
    
    document.add_paragraph(pis_body)
    
    document.add_heading('Standardized Performance Distribution', level=2)    
    document.add_picture("hist_show.png", width=Inches(5.0), height=Inches(4.0))

    document.add_page_break()      
    
    document.add_heading('Covariate Analyses', level=2)   
    cov_intro = "\nCovariate analyses indicate if significant mean differences in performance were observed across a known set of categorical variables. Ratio values indicate the proportion of observations to categories. ANOVA tests with ratios less than 10 are viewed as unreliable tests of mean differences. The table should be used to compare if the p_value is less than the p_cutoff; if so, this is a statistically significant ANOVA F-test.\n\n"
    document.add_paragraph(cov_intro)   
    
    table = document.add_table(7,6)    
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Variable Name'
    heading_cells[1].text = 'Ratio'
    heading_cells[2].text = 'ANOVA F_test'
    heading_cells[3].text = 'p_value'
    heading_cells[4].text = 'p_cutoff'
    heading_cells[5].text = 'Spinsight(c)'

    #Scale summary           
    anova_array = df_report.values[5:11,0:6]
    for i in range(0,6):
        anova_array[i,1] = anova_array[i,1].split('= ')[-1]
    
    
    jobtitle_row = table.rows[1].cells
    raters_row = table.rows[2].cells
    geo1_row = table.rows[3].cells
    geo2_row = table.rows[4].cells    
    geo3_row = table.rows[5].cells
    geo4_row = table.rows[6].cells
    
    
    for i in range(0,6):
        jobtitle_row[i].text = str(anova_array[0,i])
        raters_row[i].text = str(anova_array[1,i])
        geo1_row[i].text = str(anova_array[2,i])
        geo2_row[i].text = str(anova_array[3,i])
        geo3_row[i].text = str(anova_array[4,i])
        geo4_row[i].text = str(anova_array[5,i])
        
    
    table.style = 'Light Shading'
    #table.style = 'Light List Accent 1'
        
    document.save('cso_twisted({}).docx'.format(profile_name))
   
    print "\nCheck your Twister folder for a file named 'cso_twisted({}).docx'".format(profile_name)
        
else:
    pass


print "*"*60,"\n","*"*60,"\n","*"*60
print "\nTwister has been brought to you by: Derek Brown and Phillip Gilmore. \
\n\nPlease report any bugs or suggested improvements to: \
\nDerek Brown (derek.brown@infor.com)"


'''TO DO

-deal with duplicates
-rater policies should consider all the rater information--
prior to matching and all that.

'''